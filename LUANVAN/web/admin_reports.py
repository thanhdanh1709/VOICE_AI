"""
Admin report export — PDF, Word, CSV.
"""
import csv
import io
import os
from datetime import datetime, timedelta

REPORT_FORMATS = {
    'overview': {'pdf', 'docx'},
    'users': {'csv'},
    'conversions': {'csv'},
    'payments': {'csv'},
}

MAX_CSV_ROWS = 5000


class ReportExportError(Exception):
    def __init__(self, message, status_code=400):
        super().__init__(message)
        self.status_code = status_code


def _fmt_num(n):
    try:
        n = int(n or 0)
    except (TypeError, ValueError):
        n = 0
    return f'{n:,}'.replace(',', '.')


def _fmt_dt(v):
    if not v:
        return '—'
    if hasattr(v, 'strftime'):
        return v.strftime('%d/%m/%Y %H:%M')
    return str(v)


def _fmt_date(v):
    if not v:
        return '—'
    if hasattr(v, 'strftime'):
        return v.strftime('%d/%m/%Y')
    return str(v)


def resolve_report_period(period, date_from='', date_to=''):
    now = datetime.now()
    today_start = now.replace(hour=0, minute=0, second=0, microsecond=0)

    if period == 'today':
        start = today_start
        end = today_start + timedelta(days=1)
        label = 'Hôm nay'
    elif period == 'month':
        start = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
        end = now + timedelta(seconds=1)
        label = 'Tháng này'
    elif period == 'custom' and date_from:
        try:
            start = datetime.strptime(date_from, '%Y-%m-%d')
            end_day = date_to or date_from
            end = datetime.strptime(end_day, '%Y-%m-%d') + timedelta(days=1)
            label = f'{date_from} → {end_day}'
        except ValueError:
            raise ReportExportError('Ngày không hợp lệ (YYYY-MM-DD)')
    else:
        start = today_start - timedelta(days=now.weekday())
        end = now + timedelta(seconds=1)
        label = 'Tuần này'

    return start, end, label


def validate_report_request(report_type, fmt):
    report_type = (report_type or 'overview').lower()
    fmt = (fmt or 'pdf').lower()
    allowed = REPORT_FORMATS.get(report_type)
    if not allowed:
        raise ReportExportError('Loại báo cáo không hợp lệ')
    if fmt not in allowed:
        raise ReportExportError(f'Định dạng {fmt.upper()} không hỗ trợ cho loại báo cáo này')
    return report_type, fmt


def fetch_overview_data(conn, period, date_from='', date_to=''):
    start, end, period_label = resolve_report_period(period, date_from, date_to)
    exported_at = datetime.now().strftime('%d/%m/%Y %H:%M')
    today_start = datetime.now().replace(hour=0, minute=0, second=0, microsecond=0)

    with conn.cursor() as cursor:
        cursor.execute("SELECT COUNT(*) AS total FROM users")
        total_users = cursor.fetchone()['total'] or 0

        cursor.execute("SELECT COUNT(*) AS total FROM conversions")
        total_conversions = cursor.fetchone()['total'] or 0

        cursor.execute("SELECT COUNT(*) AS total FROM voices")
        total_voices = cursor.fetchone()['total'] or 0

        cursor.execute(
            "SELECT SUM(text_length) AS total FROM conversions WHERE text_length IS NOT NULL"
        )
        total_characters = cursor.fetchone()['total'] or 0

        cursor.execute(
            "SELECT AVG(text_length) AS avg FROM conversions WHERE text_length IS NOT NULL"
        )
        avg_text_length = round(cursor.fetchone()['avg'] or 0, 0)

        cursor.execute("SELECT COUNT(*) AS total FROM conversions WHERE status = 'completed'")
        completed = cursor.fetchone()['total'] or 0
        success_rate = round((completed / total_conversions * 100) if total_conversions else 0, 1)

        cursor.execute("SELECT COUNT(DISTINCT user_id) AS total FROM conversions")
        active_users = cursor.fetchone()['total'] or 0

        cursor.execute(
            """
            SELECT COUNT(*) AS conversions,
                   COALESCE(SUM(text_length), 0) AS characters,
                   COALESCE(SUM(audio_file_size), 0) AS audio_size,
                   COALESCE(SUM(duration_seconds), 0) AS duration
            FROM conversions
            WHERE created_at >= %s AND created_at < %s
            """,
            (start, end),
        )
        period_row = cursor.fetchone()

        chart_data = []
        for i in range(6, -1, -1):
            day_start = today_start - timedelta(days=i)
            day_end = day_start + timedelta(days=1)
            cursor.execute(
                """
                SELECT COUNT(*) AS count FROM conversions
                WHERE created_at >= %s AND created_at < %s
                """,
                (day_start, day_end),
            )
            count = cursor.fetchone()['count'] or 0
            chart_data.append({
                'label': day_start.strftime('%d/%m'),
                'date': day_start.strftime('%Y-%m-%d'),
                'conversions': count,
            })

        cursor.execute(
            """
            SELECT u.username, u.full_name,
                   COUNT(c.id) AS conversion_count,
                   COALESCE(SUM(c.text_length), 0) AS total_characters
            FROM users u
            LEFT JOIN conversions c ON u.id = c.user_id
            GROUP BY u.id, u.username, u.full_name
            ORDER BY conversion_count DESC
            LIMIT 5
            """
        )
        top_users = cursor.fetchall()

        cursor.execute(
            """
            SELECT voice_name, COUNT(*) AS usage_count,
                   COALESCE(SUM(text_length), 0) AS total_characters
            FROM conversions
            WHERE voice_name IS NOT NULL AND voice_name != '' AND voice_name != 'null'
            GROUP BY voice_name
            ORDER BY usage_count DESC
            LIMIT 5
            """
        )
        top_voices = cursor.fetchall()

        cursor.execute(
            """
            SELECT voice_name, COUNT(*) AS count
            FROM conversions
            WHERE voice_name IS NOT NULL AND voice_name != '' AND voice_name != 'null'
            GROUP BY voice_name
            HAVING COUNT(*) > 0
            ORDER BY count DESC
            LIMIT 12
            """
        )
        voice_distribution = cursor.fetchall()

    return {
        'exported_at': exported_at,
        'period_label': period_label,
        'period_range': f'{start.strftime("%d/%m/%Y")} — { (end - timedelta(seconds=1)).strftime("%d/%m/%Y") }',
        'global': {
            'total_users': total_users,
            'active_users': active_users,
            'total_conversions': total_conversions,
            'total_voices': total_voices,
            'total_characters': int(total_characters or 0),
            'avg_text_length': int(avg_text_length),
            'success_rate': success_rate,
        },
        'period_stats': {
            'conversions': period_row['conversions'] or 0,
            'characters': int(period_row['characters'] or 0),
            'audio_size_mb': round((period_row['audio_size'] or 0) / (1024 * 1024), 2),
            'duration_hours': round((period_row['duration'] or 0) / 3600, 2),
        },
        'chart_data': chart_data,
        'top_users': top_users,
        'top_voices': top_voices,
        'voice_distribution': voice_distribution,
    }


def _pdf_register_viet_fonts():
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    font_reg, font_bold = 'Helvetica', 'Helvetica-Bold'
    win_fonts = [
        ('C:/Windows/Fonts/arial.ttf', 'C:/Windows/Fonts/arialbd.ttf'),
        ('C:/Windows/Fonts/verdana.ttf', 'C:/Windows/Fonts/verdanab.ttf'),
    ]
    for reg_path, bold_path in win_fonts:
        if os.path.exists(reg_path):
            try:
                pdfmetrics.registerFont(TTFont('VV_Reg', reg_path))
                font_reg = 'VV_Reg'
                if os.path.exists(bold_path):
                    pdfmetrics.registerFont(TTFont('VV_Bold', bold_path))
                    font_bold = 'VV_Bold'
                break
            except Exception:
                pass
    return font_reg, font_bold


def build_overview_pdf(data):
    from io import BytesIO
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    buffer = BytesIO()
    font_reg, font_bold = _pdf_register_viet_fonts()
    page_w, page_h = A4
    margin = 14 * mm

    C_NAVY = colors.HexColor('#051424')
    C_PRIMARY = colors.HexColor('#7078ff')
    C_PRIMARY_LIGHT = colors.HexColor('#eef2ff')
    C_TEXT = colors.HexColor('#1e293b')
    C_MUTED = colors.HexColor('#64748b')
    C_BORDER = colors.HexColor('#e2e8f0')
    C_WHITE = colors.white
    C_CYAN = colors.HexColor('#0ea5e9')
    C_GREEN = colors.HexColor('#10b981')
    C_AMBER = colors.HexColor('#f59e0b')

    styles = {
        'section': ParagraphStyle(
            'sec', fontName=font_bold, fontSize=12, leading=16,
            textColor=C_NAVY, spaceBefore=14, spaceAfter=8,
        ),
        'cell_label': ParagraphStyle('cl', fontName=font_bold, fontSize=9, leading=12, textColor=C_TEXT),
        'cell_value': ParagraphStyle('cv', fontName=font_reg, fontSize=9, leading=12, textColor=C_TEXT),
        'kpi_label': ParagraphStyle('kl', fontName=font_bold, fontSize=8, leading=10, textColor=C_MUTED),
        'kpi_value': ParagraphStyle('kv', fontName=font_bold, fontSize=16, leading=18, textColor=C_NAVY),
        'note': ParagraphStyle('note', fontName=font_reg, fontSize=8, leading=11, textColor=C_MUTED),
        'footer': ParagraphStyle('foot', fontName=font_reg, fontSize=8, textColor=C_MUTED, alignment=TA_CENTER),
    }

    g = data['global']
    ps = data['period_stats']

    def on_page(canvas, doc):
        canvas.saveState()
        canvas.setStrokeColor(C_BORDER)
        canvas.line(margin, 22, page_w - margin, 22)
        canvas.setFont(font_reg, 8)
        canvas.setFillColor(C_MUTED)
        canvas.drawString(margin, 10, 'VietVoice AI — Báo cáo tổng quan hệ thống')
        canvas.drawRightString(page_w - margin, 10, f'Trang {doc.page}')
        canvas.restoreState()

    def on_first_page(canvas, doc):
        canvas.saveState()
        canvas.setFillColor(C_NAVY)
        canvas.rect(0, page_h - 34 * mm, page_w, 34 * mm, fill=1, stroke=0)
        canvas.setFillColor(C_PRIMARY)
        canvas.setFont(font_bold, 20)
        canvas.drawString(margin, page_h - 18 * mm, 'VietVoice AI')
        canvas.setFillColor(C_WHITE)
        canvas.setFont(font_bold, 13)
        canvas.drawRightString(page_w - margin, page_h - 16 * mm, 'BÁO CÁO TỔNG QUAN HỆ THỐNG')
        canvas.setFont(font_reg, 9)
        canvas.setFillColor(colors.HexColor('#94a3b8'))
        canvas.drawRightString(page_w - margin, page_h - 22 * mm, f'Xuất: {data["exported_at"]}')
        canvas.drawRightString(page_w - margin, page_h - 27 * mm, f'Kỳ: {data["period_label"]} ({data["period_range"]})')
        on_page(canvas, doc)
        canvas.restoreState()

    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        leftMargin=margin, rightMargin=margin,
        topMargin=40 * mm, bottomMargin=28 * mm,
        title='VietVoice — Báo cáo tổng quan',
    )

    def kpi_table():
        kpis = [
            ('Tổng người dùng', _fmt_num(g['total_users']), C_PRIMARY_LIGHT, C_PRIMARY),
            ('Tổng chuyển đổi', _fmt_num(g['total_conversions']), colors.HexColor('#e0f2fe'), C_CYAN),
            ('Số giọng đọc', _fmt_num(g['total_voices']), colors.HexColor('#d1fae5'), C_GREEN),
            ('Tổng ký tự', _fmt_num(g['total_characters']), colors.HexColor('#fef3c7'), C_AMBER),
        ]
        cells = []
        for label, value, bg, accent in kpis:
            inner = Table([
                [Paragraph(label, styles['kpi_label'])],
                [Paragraph(value, styles['kpi_value'])],
            ], colWidths=[38 * mm])
            inner.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), bg),
                ('BOX', (0, 0), (-1, -1), 1, accent),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                ('LEFTPADDING', (0, 0), (-1, -1), 10),
            ]))
            cells.append(inner)
        row = Table([cells], colWidths=[40 * mm, 40 * mm, 40 * mm, 40 * mm])
        row.setStyle(TableStyle([
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0, 0), (-1, -1), 0),
            ('RIGHTPADDING', (0, 0), (-1, -1), 4),
        ]))
        return row

    def kv_table(rows):
        body = [
            [Paragraph(label, styles['cell_label']), Paragraph(str(val), styles['cell_value'])]
            for label, val in rows
        ]
        tbl = Table(body, colWidths=[42 * mm, 118 * mm])
        tbl.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), C_PRIMARY_LIGHT),
            ('BOX', (0, 0), (-1, -1), 0.6, C_BORDER),
            ('INNERGRID', (0, 0), (-1, -1), 0.4, C_BORDER),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('LEFTPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 7),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 7),
        ]))
        return tbl

    def data_table(headers, rows, col_widths):
        header_cells = [
            Paragraph(f'<b>{h}</b>', ParagraphStyle('th', fontName=font_bold, fontSize=9, textColor=C_WHITE))
            for h in headers
        ]
        body = [header_cells]
        for row in rows:
            body.append([
                Paragraph(str(c), styles['cell_value']) for c in row
            ])
        tbl = Table(body, colWidths=col_widths, repeatRows=1)
        style_cmds = [
            ('BACKGROUND', (0, 0), (-1, 0), C_NAVY),
            ('BOX', (0, 0), (-1, -1), 0.6, C_BORDER),
            ('INNERGRID', (0, 0), (-1, -1), 0.4, C_BORDER),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('LEFTPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ]
        for ri in range(1, len(body)):
            if ri % 2 == 0:
                style_cmds.append(('BACKGROUND', (0, ri), (-1, ri), colors.HexColor('#f8fafc')))
        tbl.setStyle(TableStyle(style_cmds))
        return tbl

    story = [
        Spacer(1, 4 * mm),
        kpi_table(),
        Spacer(1, 6 * mm),
        Paragraph('Chỉ số trong kỳ báo cáo', styles['section']),
        kv_table([
            ('Chuyển đổi', _fmt_num(ps['conversions'])),
            ('Ký tự xử lý', _fmt_num(ps['characters'])),
            ('Dung lượng audio', f'{ps["audio_size_mb"]} MB'),
            ('Thời lượng', f'{ps["duration_hours"]} giờ'),
            ('Tỷ lệ thành công (toàn hệ thống)', f'{g["success_rate"]}%'),
            ('TB ký tự / lượt', _fmt_num(g['avg_text_length'])),
            ('Người dùng có chuyển đổi', _fmt_num(g['active_users'])),
        ]),
        Paragraph('Xu hướng 7 ngày qua', styles['section']),
        data_table(
            ['Ngày', 'Chuyển đổi'],
            [(d['label'], _fmt_num(d['conversions'])) for d in data['chart_data']],
            [60 * mm, 120 * mm],
        ),
        Paragraph('Top 5 người dùng', styles['section']),
        data_table(
            ['#', 'Người dùng', 'Chuyển đổi', 'Ký tự'],
            [
                (i + 1, u['username'] or '—', _fmt_num(u['conversion_count']), _fmt_num(u['total_characters']))
                for i, u in enumerate(data['top_users'])
            ],
            [12 * mm, 58 * mm, 35 * mm, 55 * mm],
        ),
        Paragraph('Top 5 giọng đọc', styles['section']),
        data_table(
            ['#', 'Giọng', 'Lượt dùng', 'Ký tự'],
            [
                (i + 1, v['voice_name'] or '—', _fmt_num(v['usage_count']), _fmt_num(v['total_characters']))
                for i, v in enumerate(data['top_voices'])
            ],
            [12 * mm, 58 * mm, 35 * mm, 55 * mm],
        ),
        Paragraph('Phân bổ giọng đọc', styles['section']),
        data_table(
            ['Giọng', 'Lượt sử dụng'],
            [(v['voice_name'] or '—', _fmt_num(v['count'])) for v in data['voice_distribution']],
            [100 * mm, 60 * mm],
        ),
        Spacer(1, 8 * mm),
        Paragraph(
            'Báo cáo được tạo tự động từ VietVoice AI Admin. Dữ liệu thời gian thực tại thời điểm xuất.',
            styles['note'],
        ),
    ]

    doc.build(story, onFirstPage=on_first_page, onLaterPages=on_page)
    buffer.seek(0)
    return buffer


def _docx_set_cell_shading(cell, fill_hex):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill_hex)
    shd.set(qn('w:val'), 'clear')
    tc_pr.append(shd)


def _docx_write_cell(cell, text, bold=False, size=10, color_rgb=(0x1e, 0x29, 0x3b), center=False):
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    cell.text = ''
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(*color_rgb)


def _docx_kv_table(doc, rows):
    from docx.shared import Cm
    table = doc.add_table(rows=len(rows), cols=2)
    table.autofit = False
    for row_idx, (label, value) in enumerate(rows):
        c0, c1 = table.rows[row_idx].cells
        c0.width = Cm(4.5)
        c1.width = Cm(12)
        _docx_set_cell_shading(c0, 'EEF2FF')
        _docx_write_cell(c0, label, bold=True, size=10, color_rgb=(0x37, 0x41, 0x51))
        _docx_write_cell(c1, value, size=10)
    return table


def _docx_data_table(doc, headers, rows):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.autofit = False
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        _docx_set_cell_shading(hdr[i], '051424')
        _docx_write_cell(hdr[i], h, bold=True, size=9, color_rgb=(0xff, 0xff, 0xff), center=True)
    for ri, row_data in enumerate(rows):
        cells = table.rows[ri + 1].cells
        bg = 'FFFFFF' if ri % 2 == 0 else 'F8FAFC'
        for ci, val in enumerate(row_data):
            _docx_set_cell_shading(cells[ci], bg)
            _docx_write_cell(cells[ci], val, size=9)
    return table


def build_overview_docx(data):
    from io import BytesIO
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()
    g = data['global']
    ps = data['period_stats']

    for section in doc.sections:
        section.top_margin = Cm(1.2)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    normal = doc.styles['Normal']
    normal.font.name = 'Arial'
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    normal.font.size = Pt(10)

    banner = doc.add_table(rows=1, cols=1)
    banner_cell = banner.rows[0].cells[0]
    _docx_set_cell_shading(banner_cell, '051424')
    p_brand = banner_cell.paragraphs[0]
    r_brand = p_brand.add_run('VietVoice AI')
    r_brand.bold = True
    r_brand.font.size = Pt(22)
    r_brand.font.color.rgb = RGBColor(0xa5, 0xb4, 0xfc)
    p_title = banner_cell.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_title = p_title.add_run('BÁO CÁO TỔNG QUAN HỆ THỐNG')
    r_title.bold = True
    r_title.font.size = Pt(14)
    r_title.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
    p_meta = banner_cell.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_meta = p_meta.add_run(
        f'Xuất: {data["exported_at"]}  |  Kỳ: {data["period_label"]} ({data["period_range"]})'
    )
    r_meta.font.size = Pt(9)
    r_meta.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

    doc.add_paragraph()

    def section_title(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x05, 0x14, 0x24)
        bar = doc.add_table(rows=1, cols=1)
        bar_cell = bar.rows[0].cells[0]
        _docx_set_cell_shading(bar_cell, '7078FF')
        bar_cell.height = Cm(0.08)

    kpi_table = doc.add_table(rows=1, cols=4)
    kpi_table.autofit = False
    kpis = [
        ('Người dùng', _fmt_num(g['total_users']), 'EEF2FF'),
        ('Chuyển đổi', _fmt_num(g['total_conversions']), 'E0F2FE'),
        ('Giọng đọc', _fmt_num(g['total_voices']), 'D1FAE5'),
        ('Ký tự', _fmt_num(g['total_characters']), 'FEF3C7'),
    ]
    for i, (label, val, bg) in enumerate(kpis):
        cell = kpi_table.rows[0].cells[i]
        _docx_set_cell_shading(cell, bg)
        _docx_write_cell(cell, label, bold=True, size=9, color_rgb=(0x64, 0x74, 0x8b))
        p2 = cell.add_paragraph()
        r2 = p2.add_run(val)
        r2.bold = True
        r2.font.size = Pt(16)
        r2.font.color.rgb = RGBColor(0x05, 0x14, 0x24)

    section_title('Chỉ số trong kỳ báo cáo')
    _docx_kv_table(doc, [
        ('Chuyển đổi', _fmt_num(ps['conversions'])),
        ('Ký tự xử lý', _fmt_num(ps['characters'])),
        ('Dung lượng audio', f'{ps["audio_size_mb"]} MB'),
        ('Thời lượng', f'{ps["duration_hours"]} giờ'),
        ('Tỷ lệ thành công', f'{g["success_rate"]}%'),
        ('TB ký tự / lượt', _fmt_num(g['avg_text_length'])),
    ])

    section_title('Xu hướng 7 ngày qua')
    _docx_data_table(
        doc,
        ['Ngày', 'Chuyển đổi'],
        [(d['label'], _fmt_num(d['conversions'])) for d in data['chart_data']],
    )

    section_title('Top 5 người dùng')
    _docx_data_table(
        doc,
        ['#', 'Người dùng', 'Chuyển đổi', 'Ký tự'],
        [
            (i + 1, u['username'] or '—', _fmt_num(u['conversion_count']), _fmt_num(u['total_characters']))
            for i, u in enumerate(data['top_users'])
        ],
    )

    section_title('Top 5 giọng đọc')
    _docx_data_table(
        doc,
        ['#', 'Giọng', 'Lượt dùng', 'Ký tự'],
        [
            (i + 1, v['voice_name'] or '—', _fmt_num(v['usage_count']), _fmt_num(v['total_characters']))
            for i, v in enumerate(data['top_voices'])
        ],
    )

    section_title('Phân bổ giọng đọc')
    _docx_data_table(
        doc,
        ['Giọng', 'Lượt sử dụng'],
        [(v['voice_name'] or '—', _fmt_num(v['count'])) for v in data['voice_distribution']],
    )

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(12)
    r_note = note.add_run(
        'Báo cáo được tạo tự động từ VietVoice AI Admin. Dữ liệu thời gian thực tại thời điểm xuất.'
    )
    r_note.font.size = Pt(8)
    r_note.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def build_users_csv(conn):
    with conn.cursor() as cursor:
        cursor.execute(
            """
            SELECT u.id, u.username, u.email, u.full_name, u.role,
                   u.is_active, u.created_at,
                   sp.package_name, us.characters_limit, us.characters_used, us.end_date,
                   (SELECT COUNT(*) FROM conversions c WHERE c.user_id = u.id) AS total_conversions
            FROM users u
            LEFT JOIN user_subscriptions us ON us.user_id = u.id AND us.is_active = 1
            LEFT JOIN subscription_packages sp ON us.package_id = sp.id
            ORDER BY u.id ASC
            LIMIT %s
            """,
            (MAX_CSV_ROWS,),
        )
        rows = cursor.fetchall()

    output = io.StringIO()
    output.write('\ufeff')
    writer = csv.writer(output)
    writer.writerow([
        'ID', 'Username', 'Email', 'Họ tên', 'Vai trò', 'Trạng thái',
        'Gói', 'Hạn mức', 'Đã dùng', 'Hết hạn', 'Chuyển đổi', 'Ngày tạo',
    ])
    for r in rows:
        writer.writerow([
            r['id'],
            r.get('username') or '',
            r.get('email') or '',
            r.get('full_name') or '',
            r.get('role') or '',
            'Hoạt động' if r.get('is_active') else 'Khóa',
            r.get('package_name') or '',
            r.get('characters_limit') or '',
            r.get('characters_used') or '',
            _fmt_date(r.get('end_date')),
            r.get('total_conversions') or 0,
            _fmt_dt(r.get('created_at')),
        ])
    return output.getvalue()


def build_conversions_csv(conn, period, date_from='', date_to='', status=''):
    start, end, _ = resolve_report_period(period, date_from, date_to)
    where = ["c.created_at >= %s", "c.created_at < %s"]
    params = [start, end]
    if status and status != 'all':
        where.append("c.status = %s")
        params.append(status)

    where_sql = " AND ".join(where)
    with conn.cursor() as cursor:
        cursor.execute(
            f"""
            SELECT c.id, u.username, c.voice_name, c.text_length, c.status,
                   c.duration_seconds, LEFT(c.text_input, 200) AS text_preview, c.created_at
            FROM conversions c
            LEFT JOIN users u ON c.user_id = u.id
            WHERE {where_sql}
            ORDER BY c.created_at DESC
            LIMIT %s
            """,
            params + [MAX_CSV_ROWS],
        )
        rows = cursor.fetchall()

    output = io.StringIO()
    output.write('\ufeff')
    writer = csv.writer(output)
    writer.writerow([
        'ID', 'User', 'Giọng', 'Ký tự', 'Trạng thái', 'Thời lượng (s)', 'Nội dung (rút gọn)', 'Thời gian',
    ])
    for r in rows:
        writer.writerow([
            r['id'],
            r.get('username') or '',
            r.get('voice_name') or '',
            r.get('text_length') or 0,
            r.get('status') or '',
            r.get('duration_seconds') or '',
            (r.get('text_preview') or '').replace('\n', ' '),
            _fmt_dt(r.get('created_at')),
        ])
    return output.getvalue()


def build_payments_csv(conn, period, date_from='', date_to='', status=''):
    where = ["1=1"]
    params = []
    if status and status != 'all':
        where.append("p.payment_status = %s")
        params.append(status)

    if period == 'custom' and date_from:
        where.append("DATE(p.created_at) >= %s")
        params.append(date_from)
        if date_to:
            where.append("DATE(p.created_at) <= %s")
            params.append(date_to)
    elif period in ('today', 'week', 'month'):
        start, end, _ = resolve_report_period(period)
        where.append("p.created_at >= %s AND p.created_at < %s")
        params.extend([start, end])

    where_sql = " AND ".join(where)
    with conn.cursor() as cursor:
        cursor.execute(
            f"""
            SELECT p.id, p.transaction_id, u.username, sp.package_name,
                   p.amount_vnd, p.payment_method, p.payment_status,
                   p.created_at, p.completed_at
            FROM payments p
            LEFT JOIN users u ON p.user_id = u.id
            LEFT JOIN subscription_packages sp ON p.package_id = sp.id
            WHERE {where_sql}
            ORDER BY p.created_at DESC
            LIMIT %s
            """,
            params + [MAX_CSV_ROWS],
        )
        rows = cursor.fetchall()

    output = io.StringIO()
    output.write('\ufeff')
    writer = csv.writer(output)
    writer.writerow([
        'ID', 'Mã GD', 'User', 'Gói', 'Số tiền VND', 'Phương thức', 'Trạng thái', 'Tạo', 'Hoàn tất',
    ])
    for r in rows:
        writer.writerow([
            r['id'],
            r.get('transaction_id') or '',
            r.get('username') or '',
            r.get('package_name') or '',
            r.get('amount_vnd') or 0,
            r.get('payment_method') or '',
            r.get('payment_status') or '',
            _fmt_dt(r.get('created_at')),
            _fmt_dt(r.get('completed_at')),
        ])
    return output.getvalue()


def make_filename(report_type, fmt):
    stamp = datetime.now().strftime('%Y%m%d')
    ext = fmt if fmt != 'docx' else 'docx'
    if fmt == 'pdf':
        ext = 'pdf'
    elif fmt == 'csv':
        ext = 'csv'
    return f'vietvoice-report-{report_type}-{stamp}.{ext}'


def render_admin_report(conn, report_type, fmt, period='week', date_from='', date_to='', status=''):
    report_type, fmt = validate_report_request(report_type, fmt)

    if report_type == 'overview':
        data = fetch_overview_data(conn, period, date_from, date_to)
        if fmt == 'pdf':
            try:
                buf = build_overview_pdf(data)
            except ImportError:
                raise ReportExportError(
                    'Thư viện reportlab chưa được cài. Chạy: pip install reportlab', 503
                )
            from flask import Response
            resp = Response(buf.getvalue(), mimetype='application/pdf')
            resp.headers['Content-Disposition'] = f'attachment; filename={make_filename(report_type, fmt)}'
            return resp
        if fmt == 'docx':
            try:
                buf = build_overview_docx(data)
            except ImportError:
                raise ReportExportError(
                    'Thư viện python-docx chưa được cài. Chạy: pip install python-docx', 503
                )
            from flask import Response
            resp = Response(buf.getvalue(), mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            resp.headers['Content-Disposition'] = f'attachment; filename={make_filename(report_type, fmt)}'
            return resp

    if report_type == 'users':
        content = build_users_csv(conn)
        from flask import Response
        resp = Response(content, mimetype='text/csv; charset=utf-8')
        resp.headers['Content-Disposition'] = f'attachment; filename={make_filename(report_type, fmt)}'
        return resp

    if report_type == 'conversions':
        content = build_conversions_csv(conn, period, date_from, date_to, status)
        from flask import Response
        resp = Response(content, mimetype='text/csv; charset=utf-8')
        resp.headers['Content-Disposition'] = f'attachment; filename={make_filename(report_type, fmt)}'
        return resp

    if report_type == 'payments':
        content = build_payments_csv(conn, period, date_from, date_to, status)
        from flask import Response
        resp = Response(content, mimetype='text/csv; charset=utf-8')
        resp.headers['Content-Disposition'] = f'attachment; filename={make_filename(report_type, fmt)}'
        return resp

    raise ReportExportError('Loại báo cáo không hỗ trợ')
