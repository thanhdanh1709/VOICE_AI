"""
TTS Web Application - Flask Backend
Ung dung web chuyen van ban thanh giong noi su dung Flask
"""
import os
import sys
import traceback
import time

# Load .env.local for local development (not pushed to GitHub)
_env_file = os.path.join(os.path.dirname(__file__), '.env.local')
if os.path.exists(_env_file):
    with open(_env_file, encoding='utf-8') as _f:
        for _line in _f:
            _line = _line.strip()
            if _line and not _line.startswith('#') and '=' in _line:
                _k, _v = _line.split('=', 1)
                os.environ.setdefault(_k.strip(), _v.strip())

# Set UTF-8 encoding for console output on Windows
if sys.platform == 'win32':
    import codecs
    sys.stdout = codecs.getwriter('utf-8')(sys.stdout.buffer, 'strict')
    sys.stderr = codecs.getwriter('utf-8')(sys.stderr.buffer, 'strict')
from pathlib import Path
from flask import Flask, render_template, request, jsonify, session, redirect, url_for, send_file, abort, Response
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
import pymysql
import uuid
import secrets

# Thêm đường dẫn VieNeu-TTS-main vào sys.path để import Vieneu
BASE_DIR = Path(__file__).resolve().parent.parent
WEB_DIR = Path(__file__).resolve().parent   # d:/.../LUANVAN/web/
VieNeu_TTS_DIR = BASE_DIR / 'VieNeu-TTS-main'
if str(VieNeu_TTS_DIR) not in sys.path:
    sys.path.insert(0, str(VieNeu_TTS_DIR))

def resolve_audio_path(path: str) -> str:
    """Convert relative sample_audio_path stored in DB to absolute path.
    Tries multiple base directories to handle files uploaded from different working directories.
    """
    if not path:
        return path
    p = Path(path)
    if p.is_absolute():
        return str(p)
    # Try 1: resolve from WEB_DIR (app.py directory) — the canonical location
    candidate1 = WEB_DIR / p
    if candidate1.exists():
        return str(candidate1)
    # Try 2: resolve from WEB_DIR's parent (LUANVAN/) — app may have been run from there
    candidate2 = WEB_DIR.parent / p
    if candidate2.exists():
        return str(candidate2)
    # Try 3: current working directory
    candidate3 = Path.cwd() / p
    if candidate3.exists():
        return str(candidate3)
    # Default: return WEB_DIR-based path (caller will check existence and report error)
    return str(candidate1)

from config import DB_CONFIG, UPLOAD_DIR, AUDIO_OUTPUT_DIR, BANK_NAME, BANK_ACCOUNT_NUMBER, BANK_ACCOUNT_NAME, BANK_BRANCH
from config import SEPAY_API_URL, SEPAY_TOKEN, SEPAY_ACCOUNT_NUMBER, SEPAY_BANK_ID, SEPAY_TIMEOUT, SEPAY_QR_API
from config import GOOGLE_CLIENT_ID, GOOGLE_CLIENT_SECRET, SECRET_KEY
from config import SMTP_HOST, SMTP_PORT, SMTP_USER, SMTP_PASSWORD, SMTP_FROM, APP_BASE_URL, DEBUG, SMTP_USE_SSL, ADMIN_EMAIL
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from audio_export import export_audio, ffmpeg_available, SUPPORTED_FORMATS, ALLOWED_BITRATES

# Deep link scheme cho Flutter mobile OAuth callback (phải khớp AndroidManifest & config.dart)
MOBILE_CALLBACK_SCHEME = 'petai'

# Số ngày vô hiệu hóa trước khi xóa vĩnh viễn sau admin duyệt xóa tài khoản
ACCOUNT_DELETION_GRACE_DAYS = 30

ALLOWED_AVATAR_EXTENSIONS = {'jpg', 'jpeg', 'png', 'webp'}
MAX_AVATAR_BYTES = 2 * 1024 * 1024
AVATAR_UPLOAD_DIR = UPLOAD_DIR / 'avatars'
AVATAR_UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
from authlib.integrations.flask_client import OAuth
import qrcode
import io
import base64
import requests
import json
import hashlib
import zipfile
import hmac
import re
from html import unescape
from datetime import datetime, timedelta

# Import RVC wrapper for voice conversion
try:
    from rvc_wrapper import get_rvc_processor
    RVC_AVAILABLE = True
    print("[INFO] RVC voice conversion is available")
except ImportError as e:
    RVC_AVAILABLE = False
    print(f"[WARNING] RVC not available: {e}")

# Import viXTTS Emotional TTS
VIXTTS_EMOTIONAL_AVAILABLE = False
VIXTTS_INSTANCE = None  # Store the instance globally
try:
    from emotional_tts_vixtts import get_vixtts_emotional_instance
    VIXTTS_EMOTIONAL_AVAILABLE = True
    print("[INFO] viXTTS Emotional TTS is available")
except ImportError as e:
    VIXTTS_EMOTIONAL_AVAILABLE = False
    print(f"[WARNING] viXTTS Emotional TTS not available: {e}")

# Import custom voice training modules
try:
    from voice_training import get_training_service
    from audio_processor import get_audio_processor
    from background_worker import start_worker, get_worker_status
    CUSTOM_VOICE_AVAILABLE = True
    print("[INFO] Custom voice training is available")
except ImportError as e:
    CUSTOM_VOICE_AVAILABLE = False
    print(f"[WARNING] Custom voice training not available: {e}")

# Flask app configuration
app = Flask(__name__, 
            static_folder='static',
            static_url_path='/static',
            template_folder='templates')
app.secret_key = SECRET_KEY


def _strip_html_text(html):
    """Bỏ tag HTML, trả về text thuần (dùng kiểm tra nội dung rỗng)."""
    if not html:
        return ''
    text = re.sub(r'<[^>]+>', '', unescape(str(html)))
    return text.strip()


def legal_section_has_content(section):
    """Mục pháp lý có tiêu đề hoặc nội dung thực sự."""
    if not section:
        return False
    if (section.get('title') or '').strip():
        return True
    return bool(_strip_html_text(section.get('content') or ''))


def legal_page_has_custom_content(page):
    """Trang pháp lý có nội dung tùy chỉnh (không phải placeholder rỗng)."""
    if not page:
        return False
    body = (page.get('body_html') or '').strip()
    if body and _strip_html_text(body):
        return True
    sections = page.get('sections') or []
    return any(legal_section_has_content(s) for s in sections)


@app.template_filter('legal_has_custom_content')
def legal_has_custom_content_filter(page):
    return legal_page_has_custom_content(page)

# Hỗ trợ reverse proxy (ngrok, nginx...) - đọc X-Forwarded headers
from werkzeug.middleware.proxy_fix import ProxyFix
app.wsgi_app = ProxyFix(app.wsgi_app, x_for=1, x_proto=1, x_host=1, x_prefix=1)

# ── Google OAuth ──────────────────────────────────────────
oauth = OAuth(app)
google = oauth.register(
    name='google',
    client_id=GOOGLE_CLIENT_ID,
    client_secret=GOOGLE_CLIENT_SECRET,
    server_metadata_url='https://accounts.google.com/.well-known/openid-configuration',
    client_kwargs={'scope': 'openid email profile'},
)

# ── Mobile Auth Token Store ────────────────────────────────
# Token một lần (5 phút) dùng để kết nối CCT OAuth → WebView session.
# Tạo sau khi Google OAuth thành công trong CCT, tiêu thụ khi WebView gọi /auth/mobile/callback.
import threading as _threading

_mobile_tokens: dict = {}
_mobile_tokens_lock = _threading.Lock()

def _create_mobile_token(user_id: int) -> str:
    """Tạo token một lần (5 phút) cho mobile auth."""
    token = str(uuid.uuid4()).replace('-', '')
    expires = time.time() + 300
    with _mobile_tokens_lock:
        expired_keys = [k for k, v in _mobile_tokens.items() if v['expires'] < time.time()]
        for k in expired_keys:
            del _mobile_tokens[k]
        _mobile_tokens[token] = {'user_id': user_id, 'expires': expires}
    return token

def _consume_mobile_token(token: str):
    """Xác thực và tiêu thụ mobile token (dùng một lần). Trả về user_id hoặc None."""
    with _mobile_tokens_lock:
        data = _mobile_tokens.pop(token, None)
    if data and data['expires'] > time.time():
        return data['user_id']
    return None

PASSWORD_RESET_EXPIRY_HOURS = 1

_password_reset_waits: dict = {}
_password_reset_waits_lock = _threading.Lock()

def _create_password_reset_wait(token):
    """Phiên chờ laptop — điện thoại xác nhận qua email."""
    wait_id = secrets.token_urlsafe(16)
    expires = time.time() + PASSWORD_RESET_EXPIRY_HOURS * 3600
    with _password_reset_waits_lock:
        expired = [k for k, v in _password_reset_waits.items() if v['expires'] < time.time()]
        for k in expired:
            del _password_reset_waits[k]
        _password_reset_waits[wait_id] = {
            'token': token,
            'confirmed': False,
            'expires': expires,
        }
    return wait_id

def _confirm_password_reset_on_phone(token):
    with _password_reset_waits_lock:
        for data in _password_reset_waits.values():
            if data['token'] == token and data['expires'] > time.time():
                data['confirmed'] = True
                return True
    return False

def _get_password_reset_wait_status(wait_id):
    with _password_reset_waits_lock:
        data = _password_reset_waits.get(wait_id)
        if not data:
            return {'status': 'not_found'}
        if data['expires'] < time.time():
            del _password_reset_waits[wait_id]
            return {'status': 'expired'}
        if data['confirmed']:
            return {'status': 'confirmed', 'token': data['token']}
        return {'status': 'pending'}

def ensure_password_reset_table():
    """Tạo bảng password_reset_tokens nếu chưa có."""
    conn = get_db_connection()
    if not conn:
        return False
    try:
        with conn.cursor() as cursor:
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS password_reset_tokens (
                    id INT AUTO_INCREMENT PRIMARY KEY,
                    user_id INT NOT NULL,
                    token VARCHAR(64) NOT NULL,
                    expires_at DATETIME NOT NULL,
                    used TINYINT(1) NOT NULL DEFAULT 0,
                    created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                    UNIQUE KEY uq_token (token),
                    KEY idx_user_id (user_id),
                    KEY idx_expires (expires_at)
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)
            conn.commit()
        return True
    except Exception as e:
        print(f"[ERROR] ensure_password_reset_table: {e}")
        return False
    finally:
        conn.close()

_user_notify_throttle = {}
_user_notify_throttle_lock = _threading.Lock()

DEFAULT_USER_SETTINGS = {
    'default_voice_id': '',
    'default_emotional_voice_id': '',
    'default_pitch': 0,
    'default_speed': 1.0,
    'default_export_format': 'wav',
    'default_export_bitrate': 192,
    'default_language': 'vi',
    'notify_chars_low': True,
    'notify_payment': True,
    'notify_plan_expiry': True,
    'notify_marketing': False,
}

def ensure_user_settings_table():
    """Tạo bảng user_settings nếu chưa có."""
    conn = get_db_connection()
    if not conn:
        return False
    try:
        with conn.cursor() as cursor:
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS user_settings (
                    user_id INT NOT NULL PRIMARY KEY,
                    default_voice_id VARCHAR(64) DEFAULT NULL,
                    default_pitch INT NOT NULL DEFAULT 0,
                    default_speed DECIMAL(4,2) NOT NULL DEFAULT 1.00,
                    default_export_format VARCHAR(10) NOT NULL DEFAULT 'wav',
                    default_export_bitrate INT NOT NULL DEFAULT 192,
                    default_language VARCHAR(10) NOT NULL DEFAULT 'vi',
                    notify_chars_low TINYINT(1) NOT NULL DEFAULT 1,
                    notify_payment TINYINT(1) NOT NULL DEFAULT 1,
                    notify_plan_expiry TINYINT(1) NOT NULL DEFAULT 1,
                    notify_marketing TINYINT(1) NOT NULL DEFAULT 0,
                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
            """)
            conn.commit()
            try:
                cursor.execute(
                    "SHOW COLUMNS FROM user_settings LIKE 'default_emotional_voice_id'"
                )
                if not cursor.fetchone():
                    cursor.execute(
                        "ALTER TABLE user_settings "
                        "ADD COLUMN default_emotional_voice_id INT DEFAULT NULL "
                        "AFTER default_voice_id"
                    )
                    conn.commit()
            except Exception as mig_e:
                print(f"[WARN] user_settings migration: {mig_e}")
        return True
    except Exception as e:
        print(f"[ERROR] ensure_user_settings_table: {e}")
        return False
    finally:
        conn.close()

def _normalize_user_settings_row(row):
    if not row:
        return dict(DEFAULT_USER_SETTINGS)
    return {
        'default_voice_id': (row.get('default_voice_id') or '').strip(),
        'default_emotional_voice_id': (
            str(row.get('default_emotional_voice_id') or '').strip()
            if row.get('default_emotional_voice_id') else ''
        ),
        'default_pitch': int(row.get('default_pitch') or 0),
        'default_speed': float(row.get('default_speed') or 1.0),
        'default_export_format': (row.get('default_export_format') or 'wav').lower(),
        'default_export_bitrate': int(row.get('default_export_bitrate') or 192),
        'default_language': (row.get('default_language') or 'vi').lower(),
        'notify_chars_low': bool(row.get('notify_chars_low')),
        'notify_payment': bool(row.get('notify_payment')),
        'notify_plan_expiry': bool(row.get('notify_plan_expiry')),
        'notify_marketing': bool(row.get('notify_marketing')),
    }

def get_user_settings(user_id):
    ensure_user_settings_table()
    conn = get_db_connection()
    if not conn:
        return dict(DEFAULT_USER_SETTINGS)
    try:
        with conn.cursor() as cursor:
            cursor.execute("SELECT * FROM user_settings WHERE user_id = %s", (user_id,))
            row = cursor.fetchone()
            if row:
                return _normalize_user_settings_row(row)
            cursor.execute(
                """
                INSERT INTO user_settings (user_id) VALUES (%s)
                """,
                (user_id,),
            )
            conn.commit()
            return dict(DEFAULT_USER_SETTINGS)
    except Exception as e:
        print(f"[ERROR] get_user_settings: {e}")
        return dict(DEFAULT_USER_SETTINGS)
    finally:
        conn.close()

def save_user_settings(user_id, data):
    ensure_user_settings_table()
    settings = _normalize_user_settings_row(data)
    fmt = settings['default_export_format']
    if fmt not in SUPPORTED_FORMATS:
        fmt = 'wav'
    bitrate = settings['default_export_bitrate']
    if bitrate not in ALLOWED_BITRATES:
        bitrate = 192
    pitch = max(-12, min(12, settings['default_pitch']))
    speed = max(0.5, min(2.0, settings['default_speed']))
    lang = settings['default_language'] if settings['default_language'] in ('vi', 'en') else 'vi'
    voice_id = (settings.get('default_voice_id') or '').strip()[:64] or None
    emotional_raw = (settings.get('default_emotional_voice_id') or '').strip()
    emotional_voice_id = int(emotional_raw) if emotional_raw.isdigit() else None

    conn = get_db_connection()
    if not conn:
        return False, 'Database connection failed'
    try:
        with conn.cursor() as cursor:
            cursor.execute(
                """
                INSERT INTO user_settings (
                    user_id, default_voice_id, default_emotional_voice_id, default_pitch, default_speed,
                    default_export_format, default_export_bitrate, default_language,
                    notify_chars_low, notify_payment, notify_plan_expiry, notify_marketing
                ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                ON DUPLICATE KEY UPDATE
                    default_voice_id = VALUES(default_voice_id),
                    default_emotional_voice_id = VALUES(default_emotional_voice_id),
                    default_pitch = VALUES(default_pitch),
                    default_speed = VALUES(default_speed),
                    default_export_format = VALUES(default_export_format),
                    default_export_bitrate = VALUES(default_export_bitrate),
                    default_language = VALUES(default_language),
                    notify_chars_low = VALUES(notify_chars_low),
                    notify_payment = VALUES(notify_payment),
                    notify_plan_expiry = VALUES(notify_plan_expiry),
                    notify_marketing = VALUES(notify_marketing)
                """,
                (
                    user_id, voice_id, emotional_voice_id, pitch, speed, fmt, bitrate, lang,
                    1 if settings['notify_chars_low'] else 0,
                    1 if settings['notify_payment'] else 0,
                    1 if settings['notify_plan_expiry'] else 0,
                    1 if settings['notify_marketing'] else 0,
                ),
            )
            conn.commit()
        return True, get_user_settings(user_id)
    except Exception as e:
        conn.rollback()
        print(f"[ERROR] save_user_settings: {e}")
        return False, str(e)
    finally:
        conn.close()

def _should_send_user_notification(user_id, kind):
    """kind: chars_low | payment | plan_expiry | marketing"""
    settings = get_user_settings(user_id)
    mapping = {
        'chars_low': 'notify_chars_low',
        'payment': 'notify_payment',
        'plan_expiry': 'notify_plan_expiry',
        'marketing': 'notify_marketing',
    }
    key = mapping.get(kind)
    return bool(settings.get(key)) if key else False

def _throttle_user_notification(user_id, kind, hours=24):
    throttle_key = f"{user_id}:{kind}"
    now = time.time()
    with _user_notify_throttle_lock:
        last = _user_notify_throttle.get(throttle_key)
        if last and now - last < hours * 3600:
            return False
        _user_notify_throttle[throttle_key] = now
    return True

def _send_user_notification_email(user_id, kind, subject, text_body, html_body=None, skip_throttle=False):
    if not _should_send_user_notification(user_id, kind):
        return False
    if not skip_throttle and not _throttle_user_notification(user_id, kind):
        return False
    conn = get_db_connection()
    if not conn:
        return False
    try:
        with conn.cursor() as cursor:
            cursor.execute("SELECT email, full_name, username FROM users WHERE id = %s", (user_id,))
            user = cursor.fetchone()
        if not user or not user.get('email'):
            return False
        html = html_body or f"<p>{text_body}</p>"
        return send_email(user['email'], subject, html, text_body)
    except Exception as e:
        print(f"[ERROR] _send_user_notification_email: {e}")
        return False
    finally:
        conn.close()

def _notify_chars_low_if_needed(user_id):
    limit_info = get_user_characters_limit(user_id)
    if not limit_info or not limit_info.get('limit'):
        return
    limit = int(limit_info['limit'])
    remaining = int(limit_info['remaining'])
    if limit <= 0:
        return
    if remaining > limit * 0.1:
        return
    subject = '[VietVoice] Ký tự sắp hết'
    text = (
        f'Bạn còn {remaining:,} ký tự trong gói hiện tại (dưới 10% hạn mức).\n'
        f'Hãy nâng cấp hoặc mua thêm ký tự tại trang Bảng giá.'
    )
    _send_user_notification_email(user_id, 'chars_low', subject, text)

def _build_branded_email_html(title, greeting_name, intro_html, info_rows, extra_html='', cta_url=None, cta_label=None):
    """HTML email template VietVoice (dark theme, bảng chi tiết)."""
    rows = ''.join(
        f'<tr>'
        f'<td style="padding:10px 14px;color:#94a3b8;font-size:14px;border-bottom:1px solid #1e293b;width:42%">{label}</td>'
        f'<td style="padding:10px 14px;color:#e2e8f0;font-size:14px;border-bottom:1px solid #1e293b;font-weight:600">{value}</td>'
        f'</tr>'
        for label, value in info_rows
    )
    cta_block = ''
    if cta_url and cta_label:
        cta_block = (
            f'<p style="text-align:center;margin:28px 0 8px">'
            f'<a href="{cta_url}" style="display:inline-block;padding:14px 28px;'
            f'background:linear-gradient(135deg,#7c3aed,#6366f1);color:#fff;text-decoration:none;'
            f'border-radius:10px;font-weight:700;font-size:15px">{cta_label}</a></p>'
        )
    return f"""
    <div style="font-family:Arial,Helvetica,sans-serif;background:#051424;padding:24px">
      <div style="max-width:560px;margin:0 auto;background:#0d1c2d;border:1px solid #1e293b;border-radius:16px;padding:28px">
        <div style="margin-bottom:20px">
          <span style="font-size:22px;font-weight:800;color:#a5b4fc">VietVoice AI</span>
        </div>
        <h1 style="color:#f1f5f9;font-size:20px;margin:0 0 16px;line-height:1.35">{title}</h1>
        <p style="color:#cbd5e1;font-size:15px;line-height:1.6;margin:0 0 20px">Xin chào <strong style="color:#e2e8f0">{greeting_name}</strong>,</p>
        {intro_html}
        <table style="width:100%;border-collapse:collapse;margin:20px 0;background:#122131;border-radius:12px;overflow:hidden">
          {rows}
        </table>
        {extra_html}
        {cta_block}
        <p style="color:#64748b;font-size:12px;line-height:1.5;margin:24px 0 0;border-top:1px solid #1e293b;padding-top:16px">
          Email tự động từ VietVoice — không trả lời email này.<br>
          Hỗ trợ: support@vietvoice.ai
        </p>
      </div>
    </div>
    """


def _notify_payment_result(
    user_id,
    success,
    package_name='',
    amount_vnd=0,
    message='',
    transaction_id='',
    characters_added=0,
    duration_days=0,
    characters_limit=None,
    characters_remaining=None,
    end_date=None,
    payment_method='',
):
    kind = 'payment'
    app_url = _get_app_base_url()
    if success:
        subject = '[VietVoice] Thanh toán thành công — Gói đã được kích hoạt'
        pkg = package_name or 'Gói dịch vụ'
        amt = _fmt_vnd(amount_vnd)
        txn = transaction_id or '—'
        chars_added = _fmt_num(characters_added) if characters_added else '—'
        limit_s = _fmt_num(characters_limit) if characters_limit else '—'
        remain_s = _fmt_num(characters_remaining) if characters_remaining else '—'
        if end_date:
            if hasattr(end_date, 'strftime'):
                end_s = end_date.strftime('%d/%m/%Y')
            else:
                end_s = str(end_date)
        else:
            end_s = '—'
        dur_s = f'{duration_days} ngày' if duration_days else '—'
        method_s = (payment_method or 'bank_qr').replace('_', ' ').upper()
        text = (
            f'Thanh toán thành công!\n\n'
            f'Gói dịch vụ: {pkg}\n'
            f'Số tiền: {amt}\n'
            f'Mã giao dịch: {txn}\n'
            f'Phương thức: {method_s}\n'
            f'Ký tự được cộng: +{chars_added}\n'
            f'Hạn mức hiện tại: {limit_s} ký tự\n'
            f'Ký tự còn lại: {remain_s}\n'
            f'Thời hạn gói đến: {end_s} (+{dur_s})\n\n'
            f'Bạn có thể sử dụng Workspace ngay: {app_url}/'
        )
        intro = (
            '<p style="color:#cbd5e1;font-size:15px;line-height:1.6;margin:0 0 8px">'
            '🎉 <strong style="color:#34d399">Thanh toán đã được xác nhận thành công.</strong> '
            'Gói dịch vụ của bạn đã được kích hoạt và cộng ký tự vào tài khoản.</p>'
        )
        info_rows = [
            ('Gói dịch vụ', pkg),
            ('Số tiền thanh toán', amt),
            ('Mã giao dịch', txn),
            ('Phương thức', method_s),
            ('Ký tự được cộng', f'+{chars_added}'),
            ('Hạn mức hiện tại', f'{limit_s} ký tự'),
            ('Ký tự còn lại', f'{remain_s} ký tự'),
            ('Gói có hiệu lực đến', end_s),
            ('Thời hạn thêm', dur_s),
        ]
        extra = (
            '<p style="color:#94a3b8;font-size:13px;line-height:1.5;margin:8px 0 0">'
            '💡 Mở Workspace để bắt đầu chuyển đổi văn bản thành giọng nói ngay.</p>'
        )
        html = _build_branded_email_html(
            'Thanh toán thành công',
            'bạn',
            intro,
            info_rows,
            extra_html=extra,
            cta_url=f'{app_url}/',
            cta_label='Mở Workspace',
        )
        _send_user_notification_email(user_id, kind, subject, text, html, skip_throttle=True)
    else:
        subject = '[VietVoice] Thanh toán không thành công'
        text = message or 'Giao dịch thanh toán không được xác nhận. Vui lòng thử lại hoặc liên hệ hỗ trợ.'
        intro = (
            '<p style="color:#fca5a5;font-size:15px;line-height:1.6;margin:0">'
            'Giao dịch thanh toán không được xác nhận. Vui lòng thử lại hoặc liên hệ hỗ trợ.</p>'
        )
        html = _build_branded_email_html(
            'Thanh toán không thành công',
            'bạn',
            intro,
            [('Chi tiết', message or text)],
            cta_url=f'{app_url}/pricing',
            cta_label='Thử lại / Nâng cấp',
        )
        _send_user_notification_email(user_id, kind, subject, text, html)


def _notify_payment_success_from_row(payment_row):
    """Gửi email và trả thông tin subscription sau khi payment đã completed."""
    user_id = payment_row['user_id']
    limit_info = get_user_characters_limit(user_id)
    end_date = limit_info.get('end_date')
    _notify_payment_result(
        user_id,
        True,
        package_name=payment_row.get('package_name') or '',
        amount_vnd=int(payment_row.get('amount_vnd') or 0),
        transaction_id=payment_row.get('transaction_id') or '',
        characters_added=int(payment_row.get('characters_limit') or 0),
        duration_days=int(payment_row.get('duration_days') or 0),
        characters_limit=limit_info.get('limit'),
        characters_remaining=limit_info.get('remaining'),
        end_date=end_date,
        payment_method=payment_row.get('payment_method') or '',
    )

def _format_datetime_vn(dt):
    if not dt:
        return '—'
    if isinstance(dt, datetime):
        return dt.strftime('%d/%m/%Y %H:%M')
    return str(dt)


def _fmt_vnd(n):
    try:
        return f'{int(n):,}'.replace(',', '.') + ' VND'
    except (TypeError, ValueError):
        return '—'


def _fmt_num(n):
    try:
        return f'{int(n):,}'.replace(',', '.')
    except (TypeError, ValueError):
        return '—'


def _payment_status_vn(status):
    m = {
        'completed': 'Đã thanh toán',
        'pending': 'Chờ xử lý',
        'failed': 'Thất bại',
        'cancelled': 'Đã hủy',
    }
    return m.get((status or '').lower(), status or '—')


def _is_account_deleted(user):
    """Tài khoản đã xóa vĩnh viễn (hết thời gian chờ 30 ngày)."""
    if not user:
        return False
    status = (user.get('status') or 'active').lower()
    if status == 'deleted':
        return True
    effective = user.get('deletion_effective_at')
    if effective and isinstance(effective, datetime) and effective <= datetime.now():
        return True
    return False

def _is_in_deletion_grace(user):
    """Tài khoản đang trong thời gian vô hiệu hóa 30 ngày (có thể khôi phục)."""
    if not user:
        return False
    if _is_account_deleted(user):
        return False
    delete_status = (user.get('delete_status') or 'none').lower()
    status = (user.get('status') or 'active').lower()
    if delete_status != 'approved':
        return False
    if status not in ('deactivated', 'deleted'):
        return False
    effective = user.get('deletion_effective_at')
    if effective and isinstance(effective, datetime):
        return effective > datetime.now()
    # deleted_at set nhưng chưa có deletion_effective_at — coi như còn trong grace
    if user.get('deleted_at'):
        return True
    return False

def _login_block_response(user):
    """Phản hồi lỗi khi không cho đăng nhập (có error_code cho i18n frontend)."""
    if _is_account_deleted(user):
        return {
            'error_code': 'account_deleted',
            'message': 'Tài khoản của bạn đã bị xóa hoặc vô hiệu hóa.',
        }
    if _is_in_deletion_grace(user):
        until = _format_datetime_vn(user.get('deletion_effective_at'))
        return {
            'error_code': 'account_deactivated_grace',
            'message': (
                f'Tài khoản của bạn đã bị vô hiệu hóa. Bạn có thể yêu cầu khôi phục trong vòng 30 ngày '
                f'(đến {until}). Vui lòng dùng chức năng "Khôi phục tài khoản" trên trang đăng nhập.'
            ),
            'error_vars': {'until': until},
        }
    return None

def _login_block_message(user):
    """Thông báo khi không cho đăng nhập vì trạng thái tài khoản."""
    block = _login_block_response(user)
    return block['message'] if block else None

def _finalize_expired_account_deletions(cursor):
    """Chuyển tài khoản hết hạn grace period sang xóa vĩnh viễn."""
    cursor.execute("""
        UPDATE users SET status = 'deleted'
        WHERE delete_status = 'approved'
          AND status = 'deactivated'
          AND deletion_effective_at IS NOT NULL
          AND deletion_effective_at <= NOW()
    """)

def _maybe_finalize_user_deletion(conn, user_id):
    """Finalize một user nếu đã hết grace period."""
    try:
        with conn.cursor() as cursor:
            _finalize_expired_account_deletions(cursor)
            conn.commit()
            cursor.execute("SELECT * FROM users WHERE id = %s", (user_id,))
            return cursor.fetchone()
    except Exception as e:
        print(f"[ERROR] _maybe_finalize_user_deletion: {e}")
        conn.rollback()
        return None

def _get_admin_notification_email():
    admin = (ADMIN_EMAIL or '').strip()
    if admin:
        return admin
    return get_support_email()

def _send_admin_delete_request_email(user):
    full_name = user.get('full_name') or user.get('username') or '—'
    requested_at = _format_datetime_vn(user.get('delete_requested_at'))
    reason = (user.get('delete_reason') or '').strip() or '—'
    subject = '[Yêu cầu xóa tài khoản] Người dùng yêu cầu xóa tài khoản'
    text_body = (
        'Xin chào Admin,\n\n'
        'Hệ thống vừa nhận được một yêu cầu xóa tài khoản từ người dùng.\n\n'
        'Thông tin tài khoản:\n'
        f'- ID người dùng: {user.get("id")}\n'
        f'- Họ tên: {full_name}\n'
        f'- Email: {user.get("email") or "—"}\n'
        f'- Ngày gửi yêu cầu: {requested_at}\n'
        f'- Lý do xóa tài khoản: {reason}\n\n'
        'Vui lòng đăng nhập vào trang quản trị để kiểm tra và xác nhận yêu cầu xóa tài khoản này.\n\n'
        'Trạng thái hiện tại: Đang chờ duyệt\n\n'
        'Trân trọng,\n'
        'Hệ thống chuyển văn bản thành giọng nói tiếng Việt'
    )
    html_body = '<pre style="font-family:Manrope,sans-serif;white-space:pre-wrap;">' + text_body + '</pre>'
    return send_email(_get_admin_notification_email(), subject, html_body, text_body)

def _send_user_delete_approved_email(user, requested_at, deleted_at, effective_at):
    to_email = (user.get('email') or '').strip()
    if not to_email:
        return False
    full_name = user.get('full_name') or user.get('username') or 'Người dùng'
    subject = '[Xác nhận xóa tài khoản] Tài khoản của bạn đã được vô hiệu hóa'
    effective_str = _format_datetime_vn(effective_at)
    text_body = (
        f'Xin chào {full_name},\n\n'
        'Yêu cầu xóa tài khoản của bạn đã được admin xác nhận.\n\n'
        'Thông tin tài khoản:\n'
        f'- Họ tên: {full_name}\n'
        f'- Email: {to_email}\n'
        f'- Thời gian gửi yêu cầu: {_format_datetime_vn(requested_at)}\n'
        f'- Thời gian được xác nhận: {_format_datetime_vn(deleted_at)}\n\n'
        f'Tài khoản của bạn hiện đã bị vô hiệu hóa trong vòng {ACCOUNT_DELETION_GRACE_DAYS} ngày '
        f'(đến {effective_str}). Trong thời gian này bạn không thể đăng nhập.\n'
        'Nếu muốn khôi phục, hãy gửi yêu cầu khôi phục trên trang đăng nhập hoặc liên hệ quản trị viên.\n\n'
        f'Sau {effective_str}, tài khoản sẽ bị xóa vĩnh viễn và không thể khôi phục.\n\n'
        'Trân trọng,\n'
        'Hệ thống chuyển văn bản thành giọng nói tiếng Việt'
    )
    html_body = '<pre style="font-family:Manrope,sans-serif;white-space:pre-wrap;">' + text_body + '</pre>'
    return send_email(to_email, subject, html_body, text_body)

def _send_admin_restore_request_email(user):
    full_name = user.get('full_name') or user.get('username') or '—'
    requested_at = _format_datetime_vn(user.get('restore_requested_at'))
    effective = _format_datetime_vn(user.get('deletion_effective_at'))
    subject = '[Yêu cầu khôi phục tài khoản] Người dùng yêu cầu khôi phục'
    text_body = (
        'Xin chào Admin,\n\n'
        'Hệ thống vừa nhận được yêu cầu khôi phục tài khoản.\n\n'
        'Thông tin tài khoản:\n'
        f'- ID người dùng: {user.get("id")}\n'
        f'- Họ tên: {full_name}\n'
        f'- Email: {user.get("email") or "—"}\n'
        f'- Thời gian yêu cầu khôi phục: {requested_at}\n'
        f'- Hạn khôi phục (trước khi xóa vĩnh viễn): {effective}\n\n'
        'Vui lòng đăng nhập trang quản trị để xử lý yêu cầu khôi phục.\n\n'
        'Trân trọng,\n'
        'Hệ thống chuyển văn bản thành giọng nói tiếng Việt'
    )
    html_body = '<pre style="font-family:Manrope,sans-serif;white-space:pre-wrap;">' + text_body + '</pre>'
    return send_email(_get_admin_notification_email(), subject, html_body, text_body)

def _send_user_restored_email(user):
    to_email = (user.get('email') or '').strip()
    if not to_email:
        return False
    full_name = user.get('full_name') or user.get('username') or 'Người dùng'
    subject = '[Khôi phục tài khoản] Tài khoản của bạn đã được khôi phục'
    text_body = (
        f'Xin chào {full_name},\n\n'
        'Tài khoản của bạn đã được khôi phục thành công. Bạn có thể đăng nhập và sử dụng dịch vụ bình thường.\n\n'
        'Trân trọng,\n'
        'Hệ thống chuyển văn bản thành giọng nói tiếng Việt'
    )
    html_body = '<pre style="font-family:Manrope,sans-serif;white-space:pre-wrap;">' + text_body + '</pre>'
    return send_email(to_email, subject, html_body, text_body)

def _send_user_delete_rejected_email(user, note):
    to_email = (user.get('email') or '').strip()
    if not to_email:
        return False
    full_name = user.get('full_name') or user.get('username') or 'Người dùng'
    reason = (note or '').strip() or '—'
    subject = '[Yêu cầu xóa tài khoản] Yêu cầu chưa được chấp nhận'
    text_body = (
        f'Xin chào {full_name},\n\n'
        'Yêu cầu xóa tài khoản của bạn chưa được chấp nhận.\n\n'
        f'Lý do:\n{reason}\n\n'
        'Nếu cần hỗ trợ thêm, vui lòng liên hệ quản trị viên.\n\n'
        'Trân trọng,\n'
        'Hệ thống chuyển văn bản thành giọng nói tiếng Việt'
    )
    html_body = '<pre style="font-family:Manrope,sans-serif;white-space:pre-wrap;">' + text_body + '</pre>'
    return send_email(to_email, subject, html_body, text_body)

def _get_app_base_url():
    if APP_BASE_URL:
        return APP_BASE_URL.rstrip('/')
    try:
        return request.host_url.rstrip('/')
    except RuntimeError:
        return 'http://127.0.0.1:5000'

def _share_audio_url(token):
    """URL công khai để chia sẻ audio (ưu tiên APP_BASE_URL / ngrok)."""
    if not token:
        return None
    return f"{_get_app_base_url()}/audio/share/{token}"

def send_email(to_email, subject, html_body, text_body=None):
    """Gửi email qua SMTP. Trả về True nếu thành công."""
    if not to_email:
        return False

    if not SMTP_HOST or not SMTP_USER:
        print(f"[EMAIL] SMTP chưa cấu hình — gửi tới {to_email}: {subject}")
        if text_body:
            print(f"[EMAIL] {text_body}")
        elif html_body:
            print(f"[EMAIL] (html body omitted)")
        return DEBUG

    from_addr = SMTP_FROM or SMTP_USER
    try:
        msg = MIMEMultipart('alternative')
        msg['Subject'] = subject
        msg['From'] = from_addr
        msg['To'] = to_email
        if text_body:
            msg.attach(MIMEText(text_body, 'plain', 'utf-8'))
        msg.attach(MIMEText(html_body, 'html', 'utf-8'))
        payload = msg.as_string()

        use_ssl = SMTP_USE_SSL or SMTP_PORT == 465
        if use_ssl:
            with smtplib.SMTP_SSL(SMTP_HOST, SMTP_PORT, timeout=30) as server:
                if SMTP_PASSWORD:
                    server.login(SMTP_USER, SMTP_PASSWORD)
                server.sendmail(from_addr, [to_email], payload)
        else:
            with smtplib.SMTP(SMTP_HOST, SMTP_PORT, timeout=30) as server:
                server.ehlo()
                server.starttls()
                server.ehlo()
                if SMTP_PASSWORD:
                    server.login(SMTP_USER, SMTP_PASSWORD)
                server.sendmail(from_addr, [to_email], payload)

        print(f"[EMAIL] Đã gửi tới {to_email}: {subject}")
        return True
    except smtplib.SMTPAuthenticationError as e:
        print(f"[ERROR] send_email auth failed: {e}")
        print("[EMAIL] Gmail/Outlook: dùng App Password, không dùng mật khẩu đăng nhập thường.")
        return False
    except Exception as e:
        print(f"[ERROR] send_email failed: {e}")
        return False

def _create_password_reset_token(user_id):
    ensure_password_reset_table()
    token = secrets.token_urlsafe(32)
    expires_at = datetime.now() + timedelta(hours=PASSWORD_RESET_EXPIRY_HOURS)

    conn = get_db_connection()
    if not conn:
        return None
    try:
        with conn.cursor() as cursor:
            cursor.execute(
                "UPDATE password_reset_tokens SET used = 1 WHERE user_id = %s AND used = 0",
                (user_id,),
            )
            cursor.execute(
                "INSERT INTO password_reset_tokens (user_id, token, expires_at) VALUES (%s, %s, %s)",
                (user_id, token, expires_at),
            )
            conn.commit()
        return token
    except Exception as e:
        conn.rollback()
        print(f"[ERROR] _create_password_reset_token: {e}")
        return None
    finally:
        conn.close()

def _validate_password_reset_token(token):
    if not token:
        return None
    ensure_password_reset_table()
    conn = get_db_connection()
    if not conn:
        return None
    try:
        with conn.cursor() as cursor:
            cursor.execute("""
                SELECT user_id FROM password_reset_tokens
                WHERE token = %s AND used = 0 AND expires_at > NOW()
            """, (token,))
            row = cursor.fetchone()
            return row['user_id'] if row else None
    except Exception as e:
        print(f"[ERROR] _validate_password_reset_token: {e}")
        return None
    finally:
        conn.close()

def _mark_password_reset_token_used(token):
    conn = get_db_connection()
    if not conn:
        return
    try:
        with conn.cursor() as cursor:
            cursor.execute(
                "UPDATE password_reset_tokens SET used = 1 WHERE token = %s",
                (token,),
            )
            conn.commit()
    except Exception as e:
        print(f"[ERROR] _mark_password_reset_token_used: {e}")
    finally:
        conn.close()

def _send_password_reset_email(user, token):
    confirm_url = f"{_get_app_base_url()}/reset-password/confirm/{token}"
    name = user.get('full_name') or user.get('username') or 'bạn'
    subject = 'VietVoice — Xác nhận đặt lại mật khẩu'
    text_body = (
        f"Xin chào {name},\n\n"
        f"Bạn vừa yêu cầu đặt lại mật khẩu VietVoice.\n"
        f"Mở email này trên ĐIỆN THOẠI và nhấn xác nhận (hiệu lực {PASSWORD_RESET_EXPIRY_HOURS} giờ):\n{confirm_url}\n\n"
        f"Sau khi xác nhận, quay lại MÁY TÍNH để nhập mật khẩu mới.\n\n"
        f"Nếu bạn không yêu cầu, hãy bỏ qua email này.\n\n— VietVoice"
    )
    html_body = f"""
    <div style="font-family:Arial,sans-serif;max-width:520px;margin:0 auto;padding:24px">
      <h2 style="color:#7c3aed;margin:0 0 12px">VietVoice</h2>
      <p>Xin chào <strong>{name}</strong>,</p>
      <p>Bạn vừa yêu cầu đặt lại mật khẩu trên máy tính.</p>
      <p><strong>📱 Mở email này trên điện thoại</strong> và nhấn nút xác nhận bên dưới (hiệu lực {PASSWORD_RESET_EXPIRY_HOURS} giờ):</p>
      <p style="text-align:center;margin:28px 0">
        <a href="{confirm_url}" style="display:inline-block;padding:14px 28px;background:linear-gradient(135deg,#7c3aed,#3b82f6);color:#fff;text-decoration:none;border-radius:10px;font-weight:700;font-size:16px">
          ✅ Xác nhận trên điện thoại
        </a>
      </p>
      <p style="font-size:12px;color:#94a3b8;text-align:center">Sau khi xác nhận, quay lại <strong>máy tính</strong> — trang web sẽ tự chuyển sang đặt mật khẩu mới.</p>
      <p style="font-size:13px;color:#64748b">Nếu bạn không yêu cầu, hãy bỏ qua email này.</p>
    </div>
  """
    sent = send_email(user['email'], subject, html_body, text_body)
    if not sent and DEBUG:
        print(f"[DEV] Phone confirm link: {confirm_url}")
    return sent

def _mobile_cct_response(callback_url: str):
    """Trả về HTML tối giản để Chrome Custom Tab kích hoạt deep link và tự đóng.
    Page này user hầu như không thấy vì app sẽ tự kéo lên foreground.
    Thông báo đăng nhập thành công sẽ hiện trong Flutter app (SnackBar).
    """
    safe_url = callback_url.replace('"', '%22').replace("'", '%27')
    html = f'''<!DOCTYPE html>
<html>
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width,initial-scale=1">
  <title>VietVoice</title>
</head>
<body style="margin:0;background:#0f172a">
  <script>
  (function() {{
    var u = "{safe_url}";
    try {{
      var a = document.createElement("a");
      a.href = u; document.body.appendChild(a); a.click();
    }} catch(e) {{ window.location.replace(u); }}
    setTimeout(function() {{ window.close(); }}, 500);
  }})();
  </script>
</body>
</html>'''
    from flask import make_response
    resp = make_response(html, 200)
    resp.headers['Content-Type'] = 'text/html; charset=utf-8'
    return resp

# Request logging
@app.after_request
def add_ngrok_skip_warning(response):
    """Bỏ trang cảnh báo ngrok khi mở link từ Gmail/điện thoại."""
    response.headers['ngrok-skip-browser-warning'] = 'true'
    return response

@app.before_request
def sync_session_avatar():
    """Đồng bộ avatar_url vào session nếu thiếu (phiên đăng nhập cũ)."""
    if 'user_id' not in session or session.get('avatar_url'):
        return
    conn = get_db_connection()
    if not conn:
        return
    try:
        with conn.cursor() as cursor:
            cursor.execute("SELECT avatar_url FROM users WHERE id = %s", (session['user_id'],))
            row = cursor.fetchone()
            session['avatar_url'] = (row.get('avatar_url') or '') if row else ''
    except Exception as e:
        print(f'[WARN] sync_session_avatar: {e}')
    finally:
        conn.close()

@app.before_request
def log_request_info():
    """Log request information"""
    try:
        if hasattr(request, 'path') and request.path and request.path.startswith('/api/'):
            print(f"[REQUEST] {request.method} {request.path}")
            if request.is_json:
                try:
                    print(f"[REQUEST] JSON data: {request.get_json()}")
                except:
                    pass
    except:
        pass  # Ignore errors in logging

# Global error handler for unhandled exceptions (chỉ catch những exception chưa được handle)
@app.errorhandler(500)
def handle_500_error(e):
    """Handle 500 errors"""
    error_trace = traceback.format_exc()
    print(f"[ERROR] Unhandled 500 error: {e}")
    print(f"[ERROR] Traceback: {error_trace}")
    
    # Return JSON response for API routes
    try:
        if hasattr(request, 'path') and request.path and request.path.startswith('/api/'):
            return jsonify({
                'success': False,
                'message': f'Lỗi server nội bộ: {str(e)}'
            }), 500
    except:
        pass
    
    # For other routes, let Flask handle it
    return "Internal Server Error", 500

# Global TTS instance (khởi tạo một lần, tái sử dụng)
_tts_instance = None

def get_tts_instance():
    """Lấy TTS instance (khởi tạo một lần)"""
    global _tts_instance
    if _tts_instance is None:
        try:
            from vieneu import Vieneu
            print("[TTS] Initializing TTS engine...")
            _tts_instance = Vieneu()
            print("[TTS] TTS engine initialized successfully")
        except ImportError as e:
            print(f"[ERROR] Error importing Vieneu: {e}")
            raise Exception(f"Không thể import Vieneu: {str(e)}")
        except Exception as e:
            print(f"[ERROR] Error initializing TTS: {e}")
            raise Exception(f"Lỗi khởi tạo TTS: {str(e)}")
    return _tts_instance

# Database connection helper
def get_db_connection():
    """Kết nối database MySQL"""
    try:
        connection = pymysql.connect(
            host=DB_CONFIG['host'],
            user=DB_CONFIG['user'],
            password=DB_CONFIG['password'],
            database=DB_CONFIG['database'],
            charset=DB_CONFIG['charset'],
            cursorclass=pymysql.cursors.DictCursor,
            connect_timeout=5,
            read_timeout=10,
            write_timeout=10
        )
        return connection
    except Exception as e:
        print(f"[ERROR] Database connection error: {e}")
        return None

# Check if user is logged in
def is_logged_in():
    """Kiểm tra người dùng đã đăng nhập chưa"""
    return 'user_id' in session

def is_admin():
    """Kiểm tra người dùng có phải admin không"""
    return session.get('user_role') == 'admin'

# Login required decorator
def login_required(f):
    """Decorator to require login for routes"""
    from functools import wraps
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not is_logged_in():
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function

def get_user_characters_limit(user_id):
    """Lấy giới hạn ký tự của user"""
    conn = get_db_connection()
    if not conn:
        # Không kết nối được DB: trả về giới hạn mặc định để app vẫn chạy
        print("[WARNING] DB connection failed in get_user_characters_limit, using default limit")
        return {'limit': 100000, 'used': 0, 'remaining': 100000, 'end_date': None}
    
    try:
        with conn.cursor() as cursor:
            # Lấy subscription active hiện tại
            cursor.execute("""
                SELECT characters_limit, characters_used, end_date
                FROM user_subscriptions
                WHERE user_id = %s AND is_active = 1 AND end_date >= CURDATE()
                ORDER BY end_date DESC
                LIMIT 1
            """, (user_id,))
            subscription = cursor.fetchone()
            
            if subscription:
                return {
                    'limit': subscription['characters_limit'],
                    'used': subscription['characters_used'] or 0,
                    'remaining': subscription['characters_limit'] - (subscription['characters_used'] or 0),
                    'end_date': subscription['end_date']
                }
            else:
                # Nếu không có subscription, tạo free plan mặc định
                cursor.execute("""
                    INSERT INTO user_subscriptions (user_id, characters_limit, characters_used, start_date, end_date)
                    VALUES (%s, 100000, 0, CURDATE(), DATE_ADD(CURDATE(), INTERVAL 30 DAY))
                """, (user_id,))
                conn.commit()
                return {
                    'limit': 100000,
                    'used': 0,
                    'remaining': 100000,
                    'end_date': None
                }
    except Exception as e:
        print(f"[ERROR] Error getting user characters limit: {e}")
        # Fallback: trả về giới hạn mặc định khi DB lỗi hoặc thiếu bảng (sau khi tạo lại DB)
        return {
            'limit': 100000,
            'used': 0,
            'remaining': 100000,
            'end_date': None
        }
    finally:
        conn.close()

def check_characters_limit(user_id, text_length):
    """Kiểm tra xem user có đủ ký tự để convert không"""
    limit_info = get_user_characters_limit(user_id)
    if not limit_info:
        # Khi không kết nối được DB: dùng giới hạn mặc định để không chặn convert
        print(f"[WARNING] Using default character limit for user {user_id} (DB unavailable)")
        limit_info = {'remaining': 100000}
    
    if limit_info['remaining'] < text_length:
        return False, f"Bạn đã hết giới hạn ký tự. Còn lại: {limit_info['remaining']:,} ký tự. Vui lòng mua thêm gói để tiếp tục sử dụng."
    
    return True, None

def update_characters_used(user_id, text_length):
    """Cập nhật số ký tự đã sử dụng"""
    conn = get_db_connection()
    if not conn:
        return False

    success = False
    try:
        with conn.cursor() as cursor:
            cursor.execute("""
                UPDATE user_subscriptions
                SET characters_used = characters_used + %s,
                    updated_at = NOW()
                WHERE user_id = %s AND is_active = 1 AND end_date >= CURDATE()
                ORDER BY end_date DESC
                LIMIT 1
            """, (text_length, user_id))
            conn.commit()
            success = True
    except Exception as e:
        print(f"[ERROR] Error updating characters used: {e}")
        conn.rollback()
    finally:
        conn.close()

    if success:
        try:
            _notify_chars_low_if_needed(user_id)
        except Exception as e:
            print(f"[WARN] chars low notify: {e}")
    return success

# Routes
@app.route('/')
def index():
    """Trang chủ"""
    if not is_logged_in():
        return redirect(url_for('landing'))
    return render_template('index.html')

@app.route('/landing')
def landing():
    """Landing page - trang giới thiệu dành cho người chưa đăng nhập"""
    if is_logged_in():
        return redirect(url_for('index'))
    content = load_landing_content()
    return render_template('landing.html', lp=content)

@app.route('/login', methods=['GET', 'POST'])
def login():
    """Đăng nhập"""
    if request.method == 'POST':
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'message': 'Dữ liệu không hợp lệ'}), 400
        
        username = (data.get('username') or '').strip()
        password = data.get('password')
        
        if not username or not password:
            return jsonify({'success': False, 'message': 'Vui lòng nhập tên đăng nhập và mật khẩu'}), 400
        
        conn = get_db_connection()
        if not conn:
            return jsonify({'success': False, 'message': 'Database connection failed'}), 500
        
        try:
            with conn.cursor() as cursor:
                cursor.execute(
                    """SELECT * FROM users WHERE username = %s OR LOWER(TRIM(email)) = %s""",
                    (username, username.lower()),
                )
                user = cursor.fetchone()

                if user and check_password_hash(user['password'], password):
                    user = _maybe_finalize_user_deletion(conn, user['id']) or user
                    block = _login_block_response(user)
                    if block:
                        return jsonify({'success': False, **block}), 403
                    if not user.get('is_active'):
                        return jsonify({'success': False, 'message': 'Tên đăng nhập hoặc mật khẩu không đúng'}), 401

                    session['user_id'] = user['id']
                    session['username'] = user['username']
                    session['user_role'] = user['role']
                    session['full_name'] = user['full_name']
                    session['avatar_url'] = user.get('avatar_url') or ''
                    session.permanent = True
                    
                    return jsonify({
                        'success': True,
                        'message': 'Đăng nhập thành công',
                        'user': {
                            'id': user['id'],
                            'username': user['username'],
                            'role': user['role'],
                            'full_name': user['full_name']
                        }
                    })
                else:
                    return jsonify({'success': False, 'message': 'Tên đăng nhập hoặc mật khẩu không đúng'}), 401
        except Exception as e:
            print(f"[ERROR] Login error: {e}")
            return jsonify({'success': False, 'message': f'Lỗi: {str(e)}'}), 500
        finally:
            conn.close()
    
    return render_template('login.html', auth_page=True)

@app.route('/forgot-password', methods=['GET', 'POST'])
def forgot_password():
    """Quên mật khẩu — gửi link đặt lại qua email."""
    if request.method == 'POST':
        data = request.get_json() or {}
        identifier = (data.get('email') or '').strip().lower()
        generic_msg = 'Nếu email tồn tại trong hệ thống, chúng tôi đã gửi link đặt lại mật khẩu. Vui lòng kiểm tra hộp thư (cả thư mục spam).'

        if not identifier:
            return jsonify({'success': False, 'message': 'Vui lòng nhập email'}), 400

        conn = get_db_connection()
        if not conn:
            return jsonify({'success': False, 'message': 'Database connection failed'}), 500

        try:
            with conn.cursor() as cursor:
                cursor.execute(
                    """SELECT id, email, username, full_name FROM users
                       WHERE is_active = 1 AND (
                           LOWER(TRIM(email)) = %s
                           OR (username = %s AND %s NOT LIKE '%%@%%')
                       )""",
                    (identifier, identifier, identifier),
                )
                user = cursor.fetchone()

            if user and not user.get('email'):
                if DEBUG:
                    print(f"[FORGOT] User '{user.get('username')}' has no email on file — cannot send")
                user = None

            wait_id = None
            if user:
                token = _create_password_reset_token(user['id'])
                if token:
                    wait_id = _create_password_reset_wait(token)
                    sent = _send_password_reset_email(user, token)
                    if DEBUG:
                        print(f"[FORGOT] Reset email → {user['email']} (wait_id={wait_id}, sent={sent})")
                elif DEBUG:
                    print(f"[FORGOT] Failed to create token for user_id={user['id']}")
            elif DEBUG:
                print(f"[FORGOT] No account for '{identifier}' — no email sent (by design)")

            payload = {'success': True, 'message': generic_msg}
            if wait_id:
                payload['wait_id'] = wait_id
                payload['message'] = (
                    'Đã gửi email. Mở email trên điện thoại và nhấn Xác nhận. '
                    'Trang máy tính này sẽ tự chuyển sang đặt mật khẩu mới.'
                )
            return jsonify(payload)
        except Exception as e:
            print(f"[ERROR] forgot_password: {e}")
            return jsonify({'success': False, 'message': f'Lỗi: {str(e)}'}), 500
        finally:
            conn.close()

    return render_template('forgot_password.html', auth_page=True)

@app.route('/api/forgot-password/status/<wait_id>')
def forgot_password_status(wait_id):
    """Laptop poll — chờ điện thoại xác nhận qua email."""
    status = _get_password_reset_wait_status(wait_id)
    if status['status'] == 'not_found':
        return jsonify({'success': False, 'status': 'not_found'}), 404
    if status['status'] == 'expired':
        return jsonify({'success': False, 'status': 'expired', 'message': 'Phiên đặt lại mật khẩu đã hết hạn'})
    if status['status'] == 'confirmed':
        return jsonify({
            'success': True,
            'status': 'confirmed',
            'redirect': url_for('reset_password', token=status['token']),
        })
    return jsonify({'success': True, 'status': 'pending'})

@app.route('/reset-password/confirm/<token>')
def reset_password_confirm(token):
    """Điện thoại: chỉ xác nhận, không nhập mật khẩu tại đây."""
    user_id = _validate_password_reset_token(token)
    if not user_id:
        return render_template('reset_confirm.html', invalid=True, auth_page=True)
    _confirm_password_reset_on_phone(token)
    return render_template('reset_confirm.html', invalid=False, auth_page=True)

@app.route('/reset-password/<token>', methods=['GET', 'POST'])
def reset_password(token):
    """Đặt lại mật khẩu bằng token từ email."""
    user_id = _validate_password_reset_token(token)

    if request.method == 'POST':
        if not user_id:
            return jsonify({'success': False, 'message': 'Link đặt lại mật khẩu không hợp lệ hoặc đã hết hạn'}), 400

        data = request.get_json() or {}
        new_pw = data.get('new_password', '')
        confirm_pw = data.get('confirm_password', '')

        if not new_pw:
            return jsonify({'success': False, 'message': 'Vui lòng nhập mật khẩu mới'}), 400
        if new_pw != confirm_pw:
            return jsonify({'success': False, 'message': 'Mật khẩu xác nhận không khớp'}), 400
        if len(new_pw) < 6:
            return jsonify({'success': False, 'message': 'Mật khẩu phải có ít nhất 6 ký tự'}), 400

        conn = get_db_connection()
        if not conn:
            return jsonify({'success': False, 'message': 'Database connection failed'}), 500

        try:
            new_hash = generate_password_hash(new_pw)
            with conn.cursor() as cursor:
                cursor.execute(
                    "UPDATE users SET password = %s WHERE id = %s AND is_active = 1",
                    (new_hash, user_id),
                )
                if cursor.rowcount == 0:
                    return jsonify({'success': False, 'message': 'Người dùng không tồn tại'}), 404
                conn.commit()
            _mark_password_reset_token_used(token)
            return jsonify({'success': True, 'message': 'Đặt lại mật khẩu thành công. Bạn có thể đăng nhập ngay.'})
        except Exception as e:
            conn.rollback()
            print(f"[ERROR] reset_password: {e}")
            return jsonify({'success': False, 'message': f'Lỗi: {str(e)}'}), 500
        finally:
            conn.close()

    return render_template('reset_password.html', token=token, invalid=not user_id, auth_page=True)

@app.route('/reset-password', methods=['GET'])
def reset_password_link():
    """Hỗ trợ link dạng ?token=... (một số app email mở kiểu này)."""
    token = (request.args.get('token') or '').strip()
    if token:
        return redirect(url_for('reset_password', token=token))
    return redirect(url_for('forgot_password'))

@app.route('/register', methods=['GET', 'POST'])
def register():
    """Đăng ký"""
    if request.method == 'POST':
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'message': 'Dữ liệu không hợp lệ'}), 400
        
        username = data.get('username')
        email = data.get('email')
        password = data.get('password')
        full_name = data.get('full_name', '')
        accept_terms = data.get('accept_terms')
        if not accept_terms:
            return jsonify({'success': False, 'message': 'Vui lòng đồng ý với điều khoản sử dụng'}), 400
        if email:
            email = email.strip().lower()
        
        conn = get_db_connection()
        if not conn:
            return jsonify({'success': False, 'message': 'Database connection failed'}), 500
        
        try:
            with conn.cursor() as cursor:
                # Check if username exists
                cursor.execute("SELECT id FROM users WHERE username = %s", (username,))
                if cursor.fetchone():
                    return jsonify({'success': False, 'message': 'Tên đăng nhập đã tồn tại'}), 400
                
                # Check if email exists
                cursor.execute("SELECT id FROM users WHERE email = %s", (email,))
                if cursor.fetchone():
                    return jsonify({'success': False, 'message': 'Email đã tồn tại'}), 400
                
                # Create new user
                hashed_password = generate_password_hash(password)
                cursor.execute(
                    "INSERT INTO users (username, email, password, full_name) VALUES (%s, %s, %s, %s)",
                    (username, email, hashed_password, full_name)
                )
                user_id = cursor.lastrowid
                
                # Tạo free subscription cho user mới (100,000 ký tự/tháng)
                cursor.execute("""
                    INSERT INTO user_subscriptions (user_id, characters_limit, characters_used, start_date, end_date)
                    VALUES (%s, 100000, 0, CURDATE(), DATE_ADD(CURDATE(), INTERVAL 30 DAY))
                """, (user_id,))
                
                conn.commit()
                
                return jsonify({'success': True, 'message': 'Đăng ký thành công. Bạn có 100,000 ký tự miễn phí/tháng'})
        except Exception as e:
            conn.rollback()
            print(f"[ERROR] Register error: {e}")
            return jsonify({'success': False, 'message': f'Lỗi: {str(e)}'}), 500
        finally:
            conn.close()
    
    return render_template('register.html', auth_page=True)

@app.route('/logout')
def logout():
    """Đăng xuất"""
    session.clear()
    return redirect(url_for('login'))

# ══════════════════════════════════════════════════════════
# GOOGLE OAUTH
# ══════════════════════════════════════════════════════════

@app.route('/auth/google')
def google_login():
    """Bắt đầu luồng đăng nhập Google (Web)"""
    if not GOOGLE_CLIENT_ID or not GOOGLE_CLIENT_SECRET:
        return redirect(url_for('login') + '?error=google_not_configured')
    redirect_uri = url_for('google_callback', _external=True)
    return google.authorize_redirect(redirect_uri)


@app.route('/auth/google/login/flutter')
def google_login_flutter():
    """Bắt đầu luồng đăng nhập Google cho Flutter mobile qua Chrome Custom Tab.
    
    Flutter mở URL này bằng FlutterWebAuth2 (Chrome Custom Tab, không phải WebView).
    Sau khi Google OAuth xong, backend redirect về petai://callback?mobile_token=...
    Flutter nhận callback rồi load /auth/mobile/callback?mobile_token=... trong WebView.
    """
    if not GOOGLE_CLIENT_ID or not GOOGLE_CLIENT_SECRET:
        return _mobile_cct_response(f'{MOBILE_CALLBACK_SCHEME}://callback?error=google_not_configured')
    session['google_login_source'] = 'flutter'
    redirect_uri = url_for('google_callback', _external=True)
    return google.authorize_redirect(redirect_uri)


@app.route('/auth/google/callback')
def google_callback():
    """Callback sau khi Google xác thực (dùng chung cho Web và Flutter mobile)"""
    source = session.pop('google_login_source', 'web')

    if not GOOGLE_CLIENT_ID or not GOOGLE_CLIENT_SECRET:
        if source == 'flutter':
            return _mobile_cct_response(f'{MOBILE_CALLBACK_SCHEME}://callback?error=google_not_configured')
        return redirect(url_for('login') + '?error=google_not_configured')

    try:
        token = google.authorize_access_token()
        user_info = token.get('userinfo') or google.userinfo()
    except Exception as e:
        print(f"[GOOGLE AUTH] Error: {e}")
        if source == 'flutter':
            return _mobile_cct_response(f'{MOBILE_CALLBACK_SCHEME}://callback?error=google_auth_failed')
        return redirect(url_for('login') + '?error=google_auth_failed')

    google_id  = user_info.get('sub')
    email      = user_info.get('email', '')
    full_name  = user_info.get('name', '')
    avatar_url = user_info.get('picture', '')

    if not google_id or not email:
        if source == 'flutter':
            return _mobile_cct_response(f'{MOBILE_CALLBACK_SCHEME}://callback?error=google_no_email')
        return redirect(url_for('login') + '?error=google_no_email')

    conn = get_db_connection()
    if not conn:
        if source == 'flutter':
            return _mobile_cct_response(f'{MOBILE_CALLBACK_SCHEME}://callback?error=db_error')
        return redirect(url_for('login') + '?error=db_error')

    try:
        with conn.cursor() as cursor:
            # 1. Tìm user theo google_id
            cursor.execute("SELECT * FROM users WHERE google_id = %s", (google_id,))
            user = cursor.fetchone()

            # 2. Tìm theo email nếu chưa có google_id
            if not user:
                cursor.execute("SELECT * FROM users WHERE email = %s", (email,))
                user = cursor.fetchone()
                if user:
                    cursor.execute(
                        "UPDATE users SET google_id = %s, avatar_url = %s WHERE id = %s",
                        (google_id, avatar_url, user['id'])
                    )
                    conn.commit()

            # 3. Tạo tài khoản mới nếu chưa có
            if not user:
                base_username = email.split('@')[0].lower()
                base_username = ''.join(c for c in base_username if c.isalnum() or c == '_')[:30]
                username = base_username

                counter = 1
                while True:
                    cursor.execute("SELECT id FROM users WHERE username = %s", (username,))
                    if not cursor.fetchone():
                        break
                    username = f"{base_username}{counter}"
                    counter += 1

                cursor.execute(
                    """INSERT INTO users (username, email, password, full_name, google_id, avatar_url, is_active)
                       VALUES (%s, %s, %s, %s, %s, %s, 1)""",
                    (username, email, '', full_name, google_id, avatar_url)
                )
                user_id = cursor.lastrowid

                cursor.execute("""
                    INSERT INTO user_subscriptions
                    (user_id, characters_limit, characters_used, start_date, end_date)
                    VALUES (%s, 100000, 0, CURDATE(), DATE_ADD(CURDATE(), INTERVAL 30 DAY))
                """, (user_id,))
                conn.commit()

                cursor.execute("SELECT * FROM users WHERE id = %s", (user_id,))
                user = cursor.fetchone()

        if not user:
            if source == 'flutter':
                return _mobile_cct_response(f'{MOBILE_CALLBACK_SCHEME}://callback?error=user_not_found')
            return redirect(url_for('login') + '?error=user_not_found')

        user = _maybe_finalize_user_deletion(conn, user['id']) or user
        block_msg = _login_block_message(user)
        if block_msg:
            err = 'account_deactivated_grace' if _is_in_deletion_grace(user) else 'account_deleted'
            if source == 'flutter':
                return _mobile_cct_response(f'{MOBILE_CALLBACK_SCHEME}://callback?error={err}')
            return redirect(url_for('login') + f'?error={err}')

        if source == 'flutter':
            # Mobile: tạo one-time token, trả về HTML để CCT tự đóng
            mobile_token = _create_mobile_token(user['id'])
            print(f"[GOOGLE AUTH] Mobile login OK: {user['username']} ({email})")
            return _mobile_cct_response(f'{MOBILE_CALLBACK_SCHEME}://callback?mobile_token={mobile_token}')
        else:
            # Web: tạo Flask session bình thường
            session['user_id']   = user['id']
            session['username']  = user['username']
            session['user_role'] = user.get('role', 'user')
            session['full_name'] = user.get('full_name', '')
            session['avatar_url'] = user.get('avatar_url') or ''
            session.permanent    = True
            print(f"[GOOGLE AUTH] Web login OK: {user['username']} ({email})")
            return redirect(url_for('index'))

    except Exception as e:
        conn.rollback()
        print(f"[GOOGLE AUTH] DB error: {e}")
        print(traceback.format_exc())
        if source == 'flutter':
            return _mobile_cct_response(f'{MOBILE_CALLBACK_SCHEME}://callback?error=db_error')
        return redirect(url_for('login') + '?error=db_error')
    finally:
        conn.close()


@app.route('/auth/mobile/callback')
def mobile_auth_callback():
    """Flutter WebView: Đổi mobile_token → Flask session để WebView đăng nhập.
    
    Flutter load URL này trong WebView sau khi nhận được mobile_token từ deep link.
    Backend xác thực token, tạo session cho WebView, redirect về trang chủ.
    """
    mobile_token = request.args.get('mobile_token', '').strip()
    if not mobile_token:
        return redirect(url_for('login') + '?error=invalid_token')

    user_id = _consume_mobile_token(mobile_token)
    if not user_id:
        print(f"[MOBILE AUTH] Token không hợp lệ hoặc đã hết hạn")
        return redirect(url_for('login') + '?error=token_expired')

    conn = get_db_connection()
    if not conn:
        return redirect(url_for('login') + '?error=db_error')

    try:
        with conn.cursor() as cursor:
            cursor.execute("SELECT * FROM users WHERE id = %s", (user_id,))
            user = cursor.fetchone()

        if not user:
            return redirect(url_for('login') + '?error=user_not_found')

        user = _maybe_finalize_user_deletion(conn, user_id) or user
        block_msg = _login_block_message(user)
        if block_msg:
            err = 'account_deactivated_grace' if _is_in_deletion_grace(user) else 'account_deleted'
            return redirect(url_for('login') + f'?error={err}')

        if not user.get('is_active'):
            return redirect(url_for('login') + '?error=user_not_found')

        session['user_id']   = user['id']
        session['username']  = user['username']
        session['user_role'] = user.get('role', 'user')
        session['full_name'] = user.get('full_name', '')
        session['avatar_url'] = user.get('avatar_url') or ''
        session.permanent    = True

        print(f"[MOBILE AUTH] WebView session tạo thành công: {user['username']}")
        return redirect(url_for('index'))
    finally:
        conn.close()

@app.route('/api/translate', methods=['POST'])
def api_translate():
    """Dịch nội dung UI qua LLM (không dịch văn bản TTS của người dùng)."""
    data = request.get_json(silent=True) or {}
    text = data.get('text') or ''
    target = (data.get('target_language') or 'en').strip().lower()

    if not isinstance(text, str):
        return jsonify({'success': False, 'translated_text': '', 'message': 'Invalid text'}), 400

    if not text.strip():
        return jsonify({'success': True, 'translated_text': text})

    try:
        from translate_service import translate_text_safe
        translated = translate_text_safe(text, target)
        return jsonify({'success': True, 'translated_text': translated})
    except Exception as e:
        print(f'[ERROR] Translate API: {e}')
        return jsonify({
            'success': False,
            'translated_text': text,
            'message': 'Translation failed',
        })

LEGAL_DISPLAY_KEYS = {
    'terms': 'terms',
    'privacy': 'privacy',
    'data_deletion': 'data_deletion',
    'payment': 'payment',
    'user_guide': 'user_guide',
    'installation_guide': 'installation_guide',
}

LEGAL_CONTENT_PAGE_KEYS = tuple(LEGAL_DISPLAY_KEYS.values())
# Hướng dẫn: luôn đọc/ghi từ USER_GUIDE.md / INSTALLATION_GUIDE.md (không dùng legal_content.json)
GUIDE_MD_PAGE_KEYS = ('user_guide', 'installation_guide')


@app.route('/api/legal/display/<page_key>')
def api_legal_display(page_key):
    """Trả HTML nội dung pháp lý theo ngôn ngữ UI (vi/en)."""
    pk = LEGAL_DISPLAY_KEYS.get(page_key, page_key)
    if pk not in LEGAL_DISPLAY_KEYS:
        return jsonify({'success': False, 'message': 'Invalid page'}), 400
    lang = (request.args.get('lang') or 'vi').strip().lower()
    legal = get_legal_for_display(pk, lang=lang)
    html = render_template('partials/legal_dynamic_body.html', legal=legal)
    return jsonify({'success': True, 'html': html, 'lang': lang, 'page': pk})


@app.route('/api/support/display')
def api_support_display():
    """Trả HTML nội dung hỗ trợ/FAQ theo ngôn ngữ (vi/en)."""
    lang = (request.args.get('lang') or 'vi').strip().lower()
    support = get_support_for_display(lang)
    html = render_template('partials/support_dynamic_body.html', support=support)
    return jsonify({'success': True, 'html': html, 'lang': lang})

@app.route('/api/landing/display')
def api_landing_display():
    """Trả nội dung landing page theo ngôn ngữ (vi/en). EN dịch qua LLM, cache theo chuỗi."""
    lang = (request.args.get('lang') or 'vi').strip().lower()
    lp = get_landing_for_display(lang)
    return jsonify({'success': True, 'lp': lp, 'lang': lang})

@app.route('/api/voices')
def get_voices():
    """Lấy danh sách giọng nói"""
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': 'Database connection failed'}), 500
    
    try:
        with conn.cursor() as cursor:
            cursor.execute("SELECT * FROM voices WHERE is_active = 1 ORDER BY voice_id")
            voices = cursor.fetchall()
            
            # Add sample file path
            for voice in voices:
                sample_filename = f"{voice['voice_id']}_sample.wav"
                sample_path = BASE_DIR / 'web' / 'static' / 'voice-samples' / sample_filename
                voice['has_sample'] = sample_path.exists()
                voice['sample_url'] = f"/static/voice-samples/{sample_filename}" if voice['has_sample'] else None
            
            return jsonify({'success': True, 'voices': voices})
    except Exception as e:
        print(f"[ERROR] Get voices error: {e}")
        return jsonify({'success': False, 'message': f'Loi: {str(e)}'}), 500
    finally:
        conn.close()

@app.route('/api/admin/generate-voice-samples', methods=['POST'])
def generate_voice_samples():
    """Tao file mau cho tat ca giong doc (Admin only)"""
    if not is_logged_in() or not is_admin():
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401
    
    sample_text = "Xin chào, tôi là giọng đọc bản Việt. Đây là giọng mẫu để bạn thử nghe. Cảm ơn bạn đã sử dụng hệ thống của chúng tôi."
    
    # Create samples directory if not exists
    samples_dir = BASE_DIR / 'web' / 'static' / 'voice-samples'
    samples_dir.mkdir(parents=True, exist_ok=True)
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': 'Database connection failed'}), 500
    
    try:
        with conn.cursor() as cursor:
            cursor.execute("SELECT voice_id, voice_name FROM voices WHERE is_active = 1")
            voices = cursor.fetchall()
        
        # Get TTS instance
        tts = get_tts_instance()
        
        generated = []
        failed = []
        
        for voice in voices:
            try:
                voice_id = voice['voice_id']
                output_filename = f"{voice_id}_sample.wav"
                output_path = samples_dir / output_filename
                
                print(f"[SAMPLE] Generating sample for voice: {voice_id}")
                
                # Get voice data
                voice_data = tts.get_preset_voice(voice_id)
                
                # Generate audio
                audio = tts.infer(text=sample_text, voice=voice_data if voice_data else None)
                
                # Save audio
                tts.save(audio, str(output_path))
                
                generated.append(voice_id)
                print(f"[SAMPLE] Generated: {output_path}")
                
            except Exception as e:
                failed.append({'voice_id': voice_id, 'error': str(e)})
                print(f"[ERROR] Failed to generate sample for {voice_id}: {e}")
        
        return jsonify({
            'success': True,
            'message': f'Da tao {len(generated)} file mau thanh cong',
            'generated': generated,
            'failed': failed,
            'total': len(voices)
        })
        
    except Exception as e:
        print(f"[ERROR] Generate voice samples error: {e}")
        return jsonify({'success': False, 'message': f'Loi: {str(e)}'}), 500
    finally:
        if conn:
            conn.close()

@app.route('/api/convert', methods=['POST'])
def convert_text_to_speech():
    """Chuyển văn bản thành giọng nói"""
    conn = None
    conversion_id = None
    
    try:
        # Kiểm tra đăng nhập
        if not is_logged_in():
            return jsonify({'success': False, 'message': 'Vui lòng đăng nhập'}), 401
        
        # Lấy dữ liệu từ request
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'message': 'Dữ liệu không hợp lệ'}), 400
        
        text = data.get('text', '').strip()
        voice_id = data.get('voice_id', 'Binh')
        
        # V2: Check if custom voice
        is_custom_voice = voice_id.startswith('custom_')
        custom_voice_data = None
        base_voice_id = voice_id
        pitch_adjustment = 0
        speed_adjustment = 1.0
        
        if is_custom_voice:
            # Extract custom voice ID
            try:
                custom_voice_id = int(voice_id.replace('custom_', ''))
                
                # Fetch custom voice details (include voice_type, sample_audio_path, ref_transcript for zero_shot)
                conn = get_db_connection()
                if conn:
                    cursor = conn.cursor()
                    try:
                        cursor.execute("""
                            SELECT base_voice_id, pitch_adjustment, speed_adjustment, 
                                   energy_adjustment, voice_name, voice_type, sample_audio_path, ref_transcript
                            FROM custom_voices 
                            WHERE id = %s AND user_id = %s AND status = 'completed'
                        """, (custom_voice_id, session['user_id']))
                    except Exception:
                        cursor.execute("""
                            SELECT base_voice_id, pitch_adjustment, speed_adjustment, 
                                   energy_adjustment, voice_name, sample_audio_path
                            FROM custom_voices 
                            WHERE id = %s AND user_id = %s AND status = 'completed'
                        """, (custom_voice_id, session['user_id']))
                    custom_voice_data = cursor.fetchone()
                    if custom_voice_data and 'voice_type' not in custom_voice_data:
                        custom_voice_data['voice_type'] = 'rvc'
                        custom_voice_data['ref_transcript'] = None
                    conn.close()
                    
                    if custom_voice_data:
                        voice_type_cv = (custom_voice_data.get('voice_type') or 'rvc').strip().lower()
                        if voice_type_cv == 'zero_shot':
                            # Zero-shot: use ref_audio + ref_transcript at infer time
                            print(f"[CONVERT Zero-shot] Using custom voice: {custom_voice_data['voice_name']}")
                        else:
                            # RVC/V2: base voice + pitch/speed
                            base_voice_id = custom_voice_data.get('base_voice_id') or base_voice_id
                            pitch_adjustment = custom_voice_data.get('pitch_adjustment', 0)
                            speed_adjustment = custom_voice_data.get('speed_adjustment', 1.0)
                            print(f"[CONVERT V2] Using custom voice: {custom_voice_data['voice_name']}")
                            print(f"[CONVERT V2] Base voice from DB: {base_voice_id}, Pitch: {pitch_adjustment}, Speed: {speed_adjustment}")
                            
                            # Remove 'HM' suffix if present (TTS doesn't use it)
                            tts_voice_id = base_voice_id
                            if tts_voice_id and str(tts_voice_id).endswith('HM'):
                                tts_voice_id = tts_voice_id[:-2]
                            if tts_voice_id:
                                tts_voice_id = str(tts_voice_id).capitalize()
                            print(f"[CONVERT V2] TTS voice ID: {tts_voice_id}")
                            voice_id = tts_voice_id
                    else:
                        print(f"[WARNING] Custom voice {custom_voice_id} not found, using default")
                        voice_id = 'BinhHM'
            except Exception as e:
                print(f"[ERROR] Error loading custom voice: {e}")
                voice_id = 'BinhHM'
        
        if not text:
            return jsonify({'success': False, 'message': 'Vui lòng nhập văn bản'}), 400
        
        # Kiểm tra giới hạn ký tự
        text_length = len(text)
        can_convert, error_message = check_characters_limit(session['user_id'], text_length)
        if not can_convert:
            return jsonify({'success': False, 'message': error_message}), 403
        
        print(f"[CONVERT] Converting text: {text[:50]}... (length: {text_length})")
        print(f"[CONVERT] Voice ID: {voice_id}")
        
        # Generate unique filename
        filename = f"{uuid.uuid4()}.wav"
        output_path = AUDIO_OUTPUT_DIR / filename
        
        # Save conversion to database (bắt buộc để lịch sử thư viện cập nhật)
        conn = get_db_connection()
        if not conn:
            return jsonify({'success': False, 'message': 'Không thể kết nối database. Vui lòng thử lại.'}), 500
        try:
            with conn.cursor() as cursor:
                cursor.execute(
                    """INSERT INTO conversions (user_id, text_input, text_length, voice_id, status)
                       VALUES (%s, %s, %s, %s, 'processing')""",
                    (session['user_id'], text, len(text), voice_id)
                )
                conn.commit()
                conversion_id = cursor.lastrowid
                print(f"[CONVERT] Saved conversion record: ID={conversion_id}")
        except Exception as e:
            print(f"[ERROR] Error saving conversion: {e}")
            conn.close()
            return jsonify({'success': False, 'message': 'Không thể lưu bản ghi chuyển đổi. Kiểm tra database.'}), 500
        conn.close()
        conn = None  # Dùng kết nối mới khi UPDATE sau (tránh timeout sau TTS)
        
        # Lấy TTS instance và chuyển đổi
        print(f"[CONVERT] Getting TTS instance...")
        try:
            tts = get_tts_instance()
            print(f"[CONVERT] TTS instance obtained successfully")
        except Exception as tts_error:
            error_trace = traceback.format_exc()
            print(f"[ERROR] Failed to get TTS instance: {tts_error}")
            print(f"[ERROR] Traceback: {error_trace}")
            raise Exception(f"Không thể khởi tạo TTS engine: {str(tts_error)}")
        
        try:
            _cv_type = (custom_voice_data.get('voice_type') or 'rvc').strip().lower() if (is_custom_voice and custom_voice_data) else 'rvc'
            voice_data = tts.get_preset_voice(voice_id) if voice_id and _cv_type not in ('zero_shot', 'vixtts_clone') else None
            print(f"[CONVERT] Voice data obtained: {voice_id}")
        except Exception as voice_error:
            print(f"[WARNING] Could not get preset voice {voice_id}, using None: {voice_error}")
            voice_data = None
        
        print(f"[CONVERT] Converting text to speech (length: {len(text)} chars)...")
        duration_seconds = 0
        sample_rate = 24000

        # viXTTS Clone: synthesize directly with user's voice reference
        use_vixtts_clone = is_custom_voice and custom_voice_data and (custom_voice_data.get('voice_type') or 'rvc').strip().lower() == 'vixtts_clone'

        if use_vixtts_clone:
            try:
                if not VIXTTS_EMOTIONAL_AVAILABLE or VIXTTS_INSTANCE is None or VIXTTS_INSTANCE.model is None:
                    raise Exception("viXTTS model chưa được tải. Vui lòng thử lại sau vài phút.")
                ref_audio_path = resolve_audio_path(custom_voice_data.get('sample_audio_path'))
                if not ref_audio_path or not os.path.exists(ref_audio_path):
                    raise Exception(f"Không tìm thấy file audio mẫu của giọng viXTTS Clone: {ref_audio_path}")
                print(f"[CONVERT viXTTS-Clone] Using voice ref: {ref_audio_path}")
                output_path.parent.mkdir(parents=True, exist_ok=True)
                VIXTTS_INSTANCE.synthesize_with_voice(text, ref_audio_path, str(output_path))
                import librosa as _librosa
                _y, _sr = _librosa.load(str(output_path), sr=None)
                duration_seconds = len(_y) / _sr
                print(f"[CONVERT viXTTS-Clone] Audio generated, duration: {duration_seconds:.2f}s")
            except Exception as vixtts_err:
                error_trace = traceback.format_exc()
                print(f"[ERROR] viXTTS Clone failed: {vixtts_err}")
                print(f"[ERROR] Traceback: {error_trace}")
                raise Exception(f"Lỗi tạo âm thanh viXTTS Clone: {str(vixtts_err)}")
        else:
            try:
                # Zero-shot: use ref_audio + ref_text; otherwise preset voice or ref from voice_data
                use_zero_shot = is_custom_voice and custom_voice_data and (custom_voice_data.get('voice_type') or 'rvc').strip().lower() == 'zero_shot'
                if use_zero_shot:
                    ref_audio_path = resolve_audio_path(custom_voice_data.get('sample_audio_path'))
                    ref_text_zs = custom_voice_data.get('ref_transcript') or ''
                    if ref_audio_path and os.path.exists(ref_audio_path) and ref_text_zs:
                        print(f"[CONVERT Zero-shot] ref_audio={ref_audio_path}, ref_text length={len(ref_text_zs)}")
                        audio = tts.infer(text=text, ref_audio=ref_audio_path, ref_text=ref_text_zs)
                    else:
                        print(f"[WARNING] Zero-shot missing ref_audio/ref_text, using default voice")
                        audio = tts.infer(text=text, voice=voice_data if voice_data else None)
                else:
                    audio = tts.infer(text=text, voice=voice_data if voice_data else None)
                
                # Calculate duration from audio shape
                sample_rate = getattr(tts, 'sample_rate', sample_rate)
                
                if hasattr(audio, 'shape'):
                    audio_shape = audio.shape
                    audio_dtype = audio.dtype
                    if len(audio_shape) > 0:
                        duration_seconds = audio_shape[0] / sample_rate
                        print(f"[CONVERT] Audio generated: shape={audio_shape}, dtype={audio_dtype}, duration={duration_seconds:.2f}s")
                    else:
                        print(f"[CONVERT] Audio generated: shape={audio_shape}, dtype={audio_dtype}")
                else:
                    audio_length = len(audio) if hasattr(audio, '__len__') else 'unknown'
                    print(f"[CONVERT] Text inference completed: audio length={audio_length}")
                
            except Exception as infer_error:
                error_trace = traceback.format_exc()
                print(f"[ERROR] Failed to infer audio: {infer_error}")
                print(f"[ERROR] Traceback: {error_trace}")
                raise Exception(f"Lỗi tạo âm thanh: {str(infer_error)}")
            
            print(f"[CONVERT] Saving audio to: {output_path}")
            try:
                # Ensure output directory exists
                output_path.parent.mkdir(parents=True, exist_ok=True)
                
                tts.save(audio, str(output_path))
                print(f"[CONVERT] Audio saved successfully")
                
                # V2: Apply speed adjustment if custom voice (before pitch)
                if is_custom_voice and speed_adjustment != 1.0:
                    try:
                        import librosa
                        import soundfile as sf
                        print(f"[CONVERT V2] Applying speed adjustment: {speed_adjustment}x")
                        
                        # Load audio
                        audio_data, sr = librosa.load(str(output_path), sr=None)
                        
                        # Change speed
                        audio_adjusted = librosa.effects.time_stretch(audio_data, rate=speed_adjustment)
                        
                        # Save adjusted audio
                        sf.write(str(output_path), audio_adjusted, sr)
                        print(f"[CONVERT V2] Speed adjustment applied successfully")
                    except Exception as speed_error:
                        print(f"[WARNING] Could not apply speed adjustment: {speed_error}")
                
                # V2: Apply pitch adjustment if custom voice
                if is_custom_voice and pitch_adjustment != 0:
                    try:
                        print(f"[CONVERT V2] Applying pitch adjustment: {pitch_adjustment}")
                        rvc_processor = get_rvc_processor()
                        if rvc_processor.is_available():
                            adjusted_filename = f"{uuid.uuid4()}_adjusted.wav"
                            adjusted_path = AUDIO_OUTPUT_DIR / adjusted_filename
                            
                            success, msg, result_path = rvc_processor.adjust_voice(
                                str(output_path), 
                                str(adjusted_path), 
                                pitch=pitch_adjustment
                            )
                            
                            if success and result_path:
                                # Delete original, use adjusted
                                os.remove(output_path)
                                output_path = Path(adjusted_path)
                                filename = adjusted_filename
                                print(f"[CONVERT V2] Pitch adjustment applied successfully")
                            else:
                                print(f"[WARNING] Pitch adjustment failed: {msg}")
                        else:
                            print(f"[WARNING] RVC not available for pitch adjustment")
                    except Exception as pitch_error:
                        print(f"[WARNING] Could not apply pitch adjustment: {pitch_error}")
                        # Continue with original audio
                    
            except Exception as save_error:
                error_trace = traceback.format_exc()
                print(f"[ERROR] Failed to save audio: {save_error}")
                print(f"[ERROR] Traceback: {error_trace}")
                raise Exception(f"Lỗi lưu file âm thanh: {str(save_error)}")
        
        # Kiểm tra file đã được tạo chưa
        if not output_path.exists():
            raise Exception("File âm thanh không được tạo thành công")
        
        file_size = os.path.getsize(output_path)
        
        # Nếu chưa có duration từ audio shape, tính từ file size
        if duration_seconds == 0:
            # Estimate duration from file size (WAV format: 44 bytes header + 2 bytes per sample at 24kHz)
            # For mono 16-bit PCM: duration ≈ (file_size - 44) / (sample_rate * 2)
            estimated_duration = max(0, (file_size - 44) / (sample_rate * 2)) if file_size > 44 else 0
            if estimated_duration > 0:
                duration_seconds = estimated_duration
        
        print(f"[CONVERT] Audio file created: {filename}")
        print(f"[CONVERT] File size: {file_size} bytes ({file_size / (1024*1024):.2f} MB), duration: {duration_seconds:.2f}s ({duration_seconds/3600:.2f} hours)")
        
        # Get voice name
        voice_name = voice_id
        if conn:
            try:
                with conn.cursor() as cursor:
                    if is_custom_voice and custom_voice_data:
                        # Use custom voice name
                        voice_name = custom_voice_data['voice_name']
                    else:
                        cursor.execute("SELECT voice_name FROM voices WHERE voice_id = %s", (voice_id,))
                        result = cursor.fetchone()
                        if result:
                            voice_name = result['voice_name']
            except Exception as e:
                print(f"[ERROR] Error getting voice name: {e}")
        
        # Update conversion in database (dùng kết nối mới để tránh timeout sau TTS)
        if conversion_id:
            conn_update = get_db_connection()
            if conn_update:
                try:
                    with conn_update.cursor() as cursor:
                        cursor.execute(
                            """UPDATE conversions SET 
                               audio_file_path = %s, audio_file_size = %s, voice_name = %s,
                               duration_seconds = %s, status = 'completed', completed_at = NOW()
                               WHERE id = %s""",
                            (str(output_path), file_size, voice_name, duration_seconds, conversion_id)
                        )
                        conn_update.commit()
                        print(f"[CONVERT] Updated conversion record: ID={conversion_id}, duration={duration_seconds:.2f}s, size={file_size} bytes")
                        
                        # Cập nhật số ký tự đã sử dụng
                        update_characters_used(session['user_id'], text_length)
                        
                        # V2: Log custom voice usage
                        if is_custom_voice and custom_voice_data:
                            try:
                                original_voice_id = data.get('voice_id', '')
                                custom_voice_id = int(original_voice_id.replace('custom_', ''))
                                cursor.execute("""
                                    INSERT INTO voice_usage_logs 
                                    (custom_voice_id, user_id, text_input, text_length, audio_duration)
                                    VALUES (%s, %s, %s, %s, %s)
                                """, (custom_voice_id, session['user_id'], text[:500], text_length, duration_seconds))
                                cursor.execute("""
                                    UPDATE custom_voices 
                                    SET usage_count = usage_count + 1 
                                    WHERE id = %s
                                """, (custom_voice_id,))
                                conn_update.commit()
                                print(f"[CONVERT V2] Logged custom voice usage: ID={custom_voice_id}")
                            except Exception as log_error:
                                print(f"[WARNING] Could not log custom voice usage: {log_error}")
                except Exception as e:
                    print(f"[ERROR] Error updating conversion: {e}")
                finally:
                    conn_update.close()
            else:
                print(f"[WARNING] Could not get DB connection to update conversion {conversion_id}")
        
        return jsonify({
            'success': True,
            'message': 'Chuyển đổi thành công',
            'audio_url': f'/api/audio/{filename}',
            'audio_filename': filename,  # Add filename for voice conversion
            'conversion_id': conversion_id
        }), 200
        
    except Exception as e:
        error_trace = traceback.format_exc()
        print(f"[ERROR] TTS conversion error: {e}")
        print(f"[ERROR] Traceback: {error_trace}")
        
        # Update status to failed (dùng kết nối mới)
        if conversion_id:
            conn_fail = get_db_connection()
            if conn_fail:
                try:
                    with conn_fail.cursor() as cursor:
                        cursor.execute("UPDATE conversions SET status = 'failed' WHERE id = %s", (conversion_id,))
                        conn_fail.commit()
                except Exception as db_error:
                    print(f"[ERROR] Error updating failed status: {db_error}")
                finally:
                    conn_fail.close()
        
        error_message = str(e)
        return jsonify({
            'success': False, 
            'message': f'Lỗi chuyển đổi: {error_message}'
        }), 500
        
    finally:
        if conn:
            conn.close()

@app.route('/api/audio/<filename>')
def get_audio(filename):
    """Lấy file âm thanh với hỗ trợ streaming"""
    if not is_logged_in():
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401
    
    file_path = AUDIO_OUTPUT_DIR / secure_filename(filename)
    if not file_path.exists() or not file_path.is_file():
        return jsonify({'success': False, 'message': 'File not found'}), 404
    
    # send_file tự động hỗ trợ range requests cho streaming
    response = send_file(
        file_path,
        mimetype='audio/wav',
        as_attachment=False,
        download_name=filename
    )
    
    # Thêm headers để tối ưu streaming
    response.headers['Accept-Ranges'] = 'bytes'
    response.headers['Cache-Control'] = 'public, max-age=3600'
    
    return response

@app.route('/api/audio/<filename>/export')
def export_audio_file(filename):
    """Xuất audio sang MP3/OGG với bitrate tùy chọn (WAV trả file gốc)."""
    if not is_logged_in():
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401

    fmt = request.args.get('format', 'mp3').lower()
    try:
        bitrate = int(request.args.get('bitrate', 192))
    except (TypeError, ValueError):
        bitrate = 192

    safe_name = secure_filename(filename)
    file_path = AUDIO_OUTPUT_DIR / safe_name
    if not file_path.exists() or not file_path.is_file():
        return jsonify({'success': False, 'message': 'File not found'}), 404

    if fmt == 'wav':
        return send_file(file_path, mimetype='audio/wav', as_attachment=True, download_name=safe_name)

    if fmt not in SUPPORTED_FORMATS:
        return jsonify({'success': False, 'message': f'Định dạng không hỗ trợ: {fmt}'}), 400

    if not ffmpeg_available():
        return jsonify({
            'success': False,
            'message': 'ffmpeg chưa được cài đặt trên server. Chỉ hỗ trợ tải WAV.'
        }), 503

    tmp_path = None
    try:
        out_path, download_name, mimetype = export_audio(file_path, fmt, bitrate)
        tmp_path = out_path if out_path != file_path else None
        return send_file(out_path, mimetype=mimetype, as_attachment=True, download_name=download_name)
    except ValueError as e:
        return jsonify({'success': False, 'message': str(e)}), 400
    except RuntimeError as e:
        return jsonify({'success': False, 'message': f'Lỗi chuyển đổi: {e}'}), 500
    finally:
        if tmp_path and tmp_path.exists():
            try:
                tmp_path.unlink()
            except OSError:
                pass

@app.route('/api/audio/formats')
def audio_formats_info():
    """Thông tin định dạng xuất được hỗ trợ."""
    return jsonify({
        'success': True,
        'formats': ['wav', 'mp3', 'ogg'],
        'bitrates': sorted(ALLOWED_BITRATES),
        'ffmpeg': ffmpeg_available(),
    })

@app.route('/api/emotional-tts/status', methods=['GET'])
def check_emotional_tts_status():
    """
    Check if Emotional TTS is ready to use
    """
    try:
        if not VIXTTS_EMOTIONAL_AVAILABLE:
            return jsonify({
                'success': True,
                'ready': False,
                'message': 'Emotional TTS không được cài đặt (import failed)'
            }), 200
        
        # Check instance và model
        if VIXTTS_INSTANCE is None:
            return jsonify({
                'success': True,
                'ready': False,
                'message': 'Model chưa được khởi tạo'
            }), 200
        
        # Check if model is loaded
        is_ready = VIXTTS_INSTANCE.model is not None
        
        return jsonify({
            'success': True,
            'ready': is_ready,
            'message': 'Sẵn sàng' if is_ready else 'Model đang load...'
        }), 200
        
    except Exception as e:
        return jsonify({
            'success': False,
            'ready': False,
            'message': str(e)
        }), 500

@app.route('/api/convert-emotional', methods=['POST'])
def convert_text_to_speech_emotional():
    """
    Chuyển văn bản thành giọng nói VỚI EMOTION CONTROL
    Sử dụng viXTTS với tự động phát hiện emotion từ text tags
    """
    conn = None
    conversion_id = None
    
    try:
        # Kiểm tra viXTTS có available không
        if not VIXTTS_EMOTIONAL_AVAILABLE or VIXTTS_INSTANCE is None:
            return jsonify({
                'success': False,
                'message': 'Tính năng Emotional TTS chưa được cài đặt. Vui lòng liên hệ admin.'
            }), 503

        # Check model đã load chưa
        if VIXTTS_INSTANCE.model is None:
            return jsonify({
                'success': False,
                'message': 'Emotional TTS đang khởi động. Vui lòng đợi 30 giây và thử lại.'
            }), 503
        
        emotional_tts = VIXTTS_INSTANCE
        
        # Kiểm tra đăng nhập
        if not is_logged_in():
            return jsonify({'success': False, 'message': 'Vui lòng đăng nhập'}), 401
        
        # Lấy dữ liệu từ request
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'message': 'Dữ liệu không hợp lệ'}), 400
        
        text = data.get('text', '').strip()
        custom_voice_id_emotional = data.get('custom_voice_id')  # optional vixtts_clone voice
        voice_id = 'viXTTS-Emotional'  # Fixed voice ID for emotional TTS
        
        if not text:
            return jsonify({'success': False, 'message': 'Vui lòng nhập văn bản'}), 400
        
        # Look up custom voice reference audio if provided
        ref_audio_for_emotional = None
        if custom_voice_id_emotional:
            try:
                conn_cv = get_db_connection()
                if conn_cv:
                    with conn_cv.cursor() as cur_cv:
                        cur_cv.execute("""
                            SELECT sample_audio_path, voice_type, voice_name
                            FROM custom_voices
                            WHERE id = %s AND user_id = %s AND status = 'completed'
                        """, (int(custom_voice_id_emotional), session['user_id']))
                        cv_row = cur_cv.fetchone()
                    conn_cv.close()
                    if cv_row and (cv_row.get('voice_type') or '').strip().lower() == 'vixtts_clone':
                        ref_path = resolve_audio_path(cv_row.get('sample_audio_path'))
                        print(f"[CONVERT EMOTIONAL] Resolved ref_path: {ref_path}")
                        if ref_path and os.path.exists(ref_path):
                            ref_audio_for_emotional = ref_path
                            voice_id = f"viXTTS-Emotional ({cv_row['voice_name']})"
                            print(f"[CONVERT EMOTIONAL] Using custom voice: {cv_row['voice_name']}, ref={ref_path}")
                        else:
                            print(f"[ERROR] Custom voice ref audio not found at: {ref_path}")
                            return jsonify({
                                'success': False,
                                'message': f'Không tìm thấy file audio mẫu của giọng "{cv_row["voice_name"]}". '
                                           f'Vui lòng xóa giọng này và tải lại file mẫu.'
                            }), 400
                    elif cv_row:
                        print(f"[WARNING] Custom voice {custom_voice_id_emotional} has wrong type: {cv_row.get('voice_type')}")
                        return jsonify({
                            'success': False,
                            'message': 'Giọng này không phải viXTTS Clone. Chỉ giọng viXTTS Clone mới được dùng ở đây.'
                        }), 400
                    else:
                        print(f"[WARNING] Custom voice {custom_voice_id_emotional} not found in DB")
                        return jsonify({
                            'success': False,
                            'message': 'Không tìm thấy giọng clone. Vui lòng thử lại.'
                        }), 404
            except Exception as cv_err:
                print(f"[ERROR] Could not load custom voice for emotional: {cv_err}")
                return jsonify({
                    'success': False,
                    'message': f'Lỗi tải thông tin giọng clone: {str(cv_err)}'
                }), 500
        
        # Kiểm tra giới hạn ký tự
        text_length = len(text)
        can_convert, error_message = check_characters_limit(session['user_id'], text_length)
        if not can_convert:
            return jsonify({'success': False, 'message': error_message}), 403
        
        print(f"[CONVERT EMOTIONAL] Text length: {text_length} chars")
        print(f"[CONVERT EMOTIONAL] Text preview: {text[:100]}...")
        
        # Generate unique filename
        filename = f"{uuid.uuid4()}_emotional.wav"
        output_path = AUDIO_OUTPUT_DIR / filename
        
        # Save conversion to database
        conn = get_db_connection()
        if not conn:
            return jsonify({'success': False, 'message': 'Không thể kết nối database'}), 500
        
        try:
            with conn.cursor() as cursor:
                cursor.execute(
                    """INSERT INTO conversions (user_id, text_input, text_length, voice_id, status)
                       VALUES (%s, %s, %s, %s, 'processing')""",
                    (session['user_id'], text, len(text), voice_id)
                )
                conn.commit()
                conversion_id = cursor.lastrowid
                print(f"[CONVERT EMOTIONAL] Conversion ID: {conversion_id}")
        except Exception as e:
            print(f"[ERROR] Error saving conversion: {e}")
            conn.close()
            return jsonify({'success': False, 'message': 'Không thể lưu bản ghi chuyển đổi'}), 500
        
        conn.close()
        conn = None
        
        # Use the pre-loaded VIXTTS_INSTANCE (already checked at start of function)
        print(f"[CONVERT EMOTIONAL] Using pre-loaded viXTTS instance")
        
        # Generate audio with emotion control
        print(f"[CONVERT EMOTIONAL] Generating audio with emotion control...")
        try:
            if ref_audio_for_emotional:
                print(f"[CONVERT EMOTIONAL] Using custom voice ref for emotional synthesis")
                emotional_tts.synthesize_emotional_with_voice(text, ref_audio_for_emotional, str(output_path))
            else:
                emotional_tts.synthesize(text, str(output_path))
            
            if not output_path.exists():
                raise Exception("File audio không được tạo thành công")
            
            file_size = os.path.getsize(output_path)
            print(f"[CONVERT EMOTIONAL] Audio created: {file_size} bytes")
            
            # Calculate duration
            import librosa
            y, sr = librosa.load(str(output_path), sr=None)
            duration_seconds = len(y) / sr
            print(f"[CONVERT EMOTIONAL] Duration: {duration_seconds:.2f}s")
            
        except Exception as gen_error:
            error_trace = traceback.format_exc()
            print(f"[ERROR] Failed to generate audio: {gen_error}")
            print(f"[ERROR] Traceback: {error_trace}")
            raise Exception(f"Lỗi tạo âm thanh: {str(gen_error)}")
        
        # Update conversion record
        if conversion_id:
            conn_update = get_db_connection()
            if conn_update:
                try:
                    with conn_update.cursor() as cursor:
                        cursor.execute(
                            """UPDATE conversions SET 
                               audio_file_path = %s, audio_file_size = %s, voice_name = %s,
                               duration_seconds = %s, status = 'completed', completed_at = NOW()
                               WHERE id = %s""",
                            (str(output_path), file_size, 'Emotional Voice (viXTTS)', duration_seconds, conversion_id)
                        )
                        conn_update.commit()
                        
                        # Update characters used
                        update_characters_used(session['user_id'], text_length)
                        
                except Exception as e:
                    print(f"[ERROR] Error updating conversion: {e}")
                finally:
                    conn_update.close()
        
        # Return success response
        return jsonify({
            'success': True,
            'message': 'Chuyển đổi thành công với emotion control!',
            'audio_url': f'/api/audio/{filename}',
            'audio_filename': filename,
            'conversion_id': conversion_id,
            'file_size': file_size,
            'duration': round(duration_seconds, 2)
        }), 200
        
    except Exception as e:
        error_trace = traceback.format_exc()
        print(f"[ERROR] Emotional conversion error: {e}")
        print(f"[ERROR] Traceback: {error_trace}")
        
        # Mark conversion as failed
        if conversion_id:
            try:
                conn_err = get_db_connection()
                if conn_err:
                    with conn_err.cursor() as cursor:
                        cursor.execute(
                            "UPDATE conversions SET status = 'failed', completed_at = NOW() WHERE id = %s",
                            (conversion_id,)
                        )
                        conn_err.commit()
                    conn_err.close()
            except:
                pass
        
        return jsonify({
            'success': False,
            'message': f'Lỗi: {str(e)}'
        }), 500
    
    finally:
        if conn:
            conn.close()

@app.route('/api/upload/extract', methods=['POST'])
def extract_text_from_file():
    """Extract text from uploaded file (.txt, .pdf, .docx)"""
    if not is_logged_in():
        return jsonify({'success': False, 'message': 'Vui lòng đăng nhập'}), 401
    
    if 'file' not in request.files:
        return jsonify({'success': False, 'message': 'Không có file được tải lên'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'success': False, 'message': 'Không có file được chọn'}), 400
    
    # Check file extension
    filename = secure_filename(file.filename)
    file_ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
    
    if file_ext not in ['txt', 'pdf', 'docx']:
        return jsonify({'success': False, 'message': f'Định dạng file .{file_ext} không được hỗ trợ. Chỉ hỗ trợ .txt, .pdf, .docx'}), 400
    
    try:
        # Save file temporarily
        temp_filename = f"{uuid.uuid4()}.{file_ext}"
        temp_path = UPLOAD_DIR / temp_filename
        file.save(temp_path)
        
        text_content = None
        
        # Extract text based on file type
        if file_ext == 'txt':
            with open(temp_path, 'r', encoding='utf-8') as f:
                text_content = f.read()
        
        elif file_ext == 'docx':
            try:
                from docx import Document
                doc = Document(temp_path)
                paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
                text_content = '\n'.join(paragraphs)
            except ImportError:
                # Clean up temp file
                if temp_path.exists():
                    temp_path.unlink()
                return jsonify({
                    'success': False, 
                    'message': 'Thư viện python-docx chưa được cài đặt. Chạy: pip install python-docx'
                }), 500
            except Exception as e:
                if temp_path.exists():
                    temp_path.unlink()
                print(f"[ERROR] Error reading DOCX file: {e}")
                return jsonify({'success': False, 'message': f'Lỗi đọc file Word: {str(e)}'}), 500
        
        elif file_ext == 'pdf':
            try:
                import PyPDF2
                with open(temp_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    paragraphs = []
                    for page in pdf_reader.pages:
                        text = page.extract_text()
                        if text.strip():
                            paragraphs.append(text)
                    text_content = '\n'.join(paragraphs)
            except ImportError:
                if temp_path.exists():
                    temp_path.unlink()
                return jsonify({
                    'success': False, 
                    'message': 'Thư viện PyPDF2 chưa được cài đặt. Chạy: pip install PyPDF2'
                }), 500
            except Exception as e:
                if temp_path.exists():
                    temp_path.unlink()
                print(f"[ERROR] Error reading PDF file: {e}")
                return jsonify({'success': False, 'message': f'Lỗi đọc file PDF: {str(e)}'}), 500
        
        # Clean up temp file
        if temp_path.exists():
            temp_path.unlink()
        
        if not text_content or not text_content.strip():
            return jsonify({'success': False, 'message': 'Không tìm thấy văn bản trong file'}), 400
        
        return jsonify({
            'success': True,
            'text': text_content,
            'message': f'Đã đọc {len(text_content)} ký tự từ file {filename}'
        })
        
    except Exception as e:
        error_trace = traceback.format_exc()
        print(f"[ERROR] Error extracting text from file: {e}")
        print(f"[ERROR] Traceback: {error_trace}")
        
        # Clean up temp file if exists
        temp_path = UPLOAD_DIR / temp_filename if 'temp_filename' in locals() else None
        if temp_path and temp_path.exists():
            temp_path.unlink()
        
        return jsonify({'success': False, 'message': f'Lỗi xử lý file: {str(e)}'}), 500

@app.route('/api/history')
def get_history():
    """Lấy lịch sử chuyển đổi"""
    if not is_logged_in():
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401
    
    page = int(request.args.get('page', 1))
    per_page = int(request.args.get('per_page', 10))
    search = request.args.get('search', '')
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': 'Database connection failed'}), 500
    
    try:
        with conn.cursor() as cursor:
            query = "SELECT * FROM conversions WHERE user_id = %s"
            params = [session['user_id']]
            
            if search:
                query += " AND text_input LIKE %s"
                params.append(f'%{search}%')
            
            query += " ORDER BY created_at DESC LIMIT %s OFFSET %s"
            params.extend([per_page, (page - 1) * per_page])
            
            cursor.execute(query, params)
            conversions = cursor.fetchall()
            
            # Get total count
            count_query = "SELECT COUNT(*) as total FROM conversions WHERE user_id = %s"
            count_params = [session['user_id']]
            if search:
                count_query += " AND text_input LIKE %s"
                count_params.append(f'%{search}%')
            
            cursor.execute(count_query, count_params)
            total = cursor.fetchone()['total']
            
            return jsonify({
                'success': True,
                'conversions': conversions,
                'total': total,
                'page': page,
                'per_page': per_page
            })
    except Exception as e:
        print(f"[ERROR] Get history error: {e}")
        return jsonify({'success': False, 'message': f'Lỗi: {str(e)}'}), 500
    finally:
        conn.close()

@app.route('/api/statistics')
def get_statistics():
    """Lấy thống kê tổng quan mở rộng"""
    if not is_logged_in():
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': 'Database connection failed'}), 500
    
    try:
        with conn.cursor() as cursor:
            # If admin, show system-wide statistics
            if is_admin():
                # Total users
                cursor.execute("SELECT COUNT(*) as total FROM users")
                total_users = cursor.fetchone()['total']
                
                # Total conversions
                cursor.execute("SELECT COUNT(*) as total FROM conversions")
                total_conversions = cursor.fetchone()['total']
                
                # Total voices
                cursor.execute("SELECT COUNT(*) as total FROM voices")
                total_voices = cursor.fetchone()['total']
                
                # Total characters processed
                cursor.execute("SELECT SUM(text_length) as total FROM conversions WHERE text_length IS NOT NULL")
                result = cursor.fetchone()
                total_characters = result['total'] or 0
                
                # Total audio file size (MB)
                cursor.execute("SELECT SUM(audio_file_size) as total FROM conversions WHERE audio_file_size IS NOT NULL")
                result = cursor.fetchone()
                total_audio_size_bytes = result['total'] or 0
                total_audio_size_mb = round(total_audio_size_bytes / (1024 * 1024), 2)
                
                # Total duration (seconds -> hours)
                cursor.execute("SELECT SUM(duration_seconds) as total FROM conversions WHERE duration_seconds IS NOT NULL")
                result = cursor.fetchone()
                total_duration_seconds = result['total'] or 0
                total_duration_hours = round(total_duration_seconds / 3600, 2)
                
                # Average text length
                cursor.execute("SELECT AVG(text_length) as avg FROM conversions WHERE text_length IS NOT NULL")
                result = cursor.fetchone()
                avg_text_length = round(result['avg'] or 0, 0)
                
                # Success rate
                cursor.execute("SELECT COUNT(*) as total FROM conversions WHERE status = 'completed'")
                completed = cursor.fetchone()['total']
                success_rate = round((completed / total_conversions * 100) if total_conversions > 0 else 0, 1)
                
                # Active users (users with at least 1 conversion)
                cursor.execute("SELECT COUNT(DISTINCT user_id) as total FROM conversions")
                active_users = cursor.fetchone()['total']
                
                return jsonify({
                    'success': True,
                    'statistics': {
                        'total_users': total_users,
                        'active_users': active_users,
                        'total_conversions': total_conversions,
                        'total_voices': total_voices,
                        'total_characters': total_characters,
                        'total_audio_size_mb': total_audio_size_mb,
                        'total_duration_hours': total_duration_hours,
                        'avg_text_length': int(avg_text_length),
                        'success_rate': success_rate
                    }
                })
            else:
                # Regular user statistics
                cursor.execute("SELECT COUNT(*) as total FROM conversions WHERE user_id = %s", (session['user_id'],))
                total_conversions = cursor.fetchone()['total']
                
                cursor.execute("SELECT SUM(text_length) as total FROM conversions WHERE user_id = %s AND text_length IS NOT NULL", (session['user_id'],))
                result = cursor.fetchone()
                total_characters = result['total'] or 0
                
                cursor.execute("SELECT SUM(audio_file_size) as total FROM conversions WHERE user_id = %s AND audio_file_size IS NOT NULL", (session['user_id'],))
                result = cursor.fetchone()
                total_audio_size_bytes = result['total'] or 0
                total_audio_size_mb = round(total_audio_size_bytes / (1024 * 1024), 2)
                
                cursor.execute("SELECT SUM(duration_seconds) as total FROM conversions WHERE user_id = %s AND duration_seconds IS NOT NULL", (session['user_id'],))
                result = cursor.fetchone()
                total_duration_seconds = result['total'] or 0
                total_duration_hours = round(total_duration_seconds / 3600, 2)
                
                return jsonify({
                    'success': True,
                    'statistics': {
                        'total_conversions': total_conversions,
                        'total_characters': total_characters,
                        'total_audio_size_mb': total_audio_size_mb,
                        'total_duration_hours': total_duration_hours
                    }
                })
    except Exception as e:
        print(f"[ERROR] Get statistics error: {e}")
        return jsonify({'success': False, 'message': f'Lỗi: {str(e)}'}), 500
    finally:
        conn.close()

@app.route('/api/statistics/time-based')
def get_time_based_statistics():
    """Lấy thống kê theo thời gian (hôm nay/tuần/tháng)"""
    if not is_logged_in() or not is_admin():
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': 'Database connection failed'}), 500
    
    try:
        from datetime import datetime, timedelta
        
        with conn.cursor() as cursor:
            now = datetime.now()
            today_start = now.replace(hour=0, minute=0, second=0, microsecond=0)
            week_start = today_start - timedelta(days=now.weekday())
            month_start = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
            
            # Today statistics
            cursor.execute("""
                SELECT COUNT(*) as conversions,
                       SUM(text_length) as characters,
                       SUM(audio_file_size) as audio_size,
                       SUM(duration_seconds) as duration
                FROM conversions
                WHERE created_at >= %s AND created_at < %s
            """, (today_start, today_start + timedelta(days=1)))
            today = cursor.fetchone()
            
            # This week statistics
            cursor.execute("""
                SELECT COUNT(*) as conversions,
                       SUM(text_length) as characters,
                       SUM(audio_file_size) as audio_size,
                       SUM(duration_seconds) as duration
                FROM conversions
                WHERE created_at >= %s
            """, (week_start,))
            week = cursor.fetchone()
            
            # This month statistics
            cursor.execute("""
                SELECT COUNT(*) as conversions,
                       SUM(text_length) as characters,
                       SUM(audio_file_size) as audio_size,
                       SUM(duration_seconds) as duration
                FROM conversions
                WHERE created_at >= %s
            """, (month_start,))
            month = cursor.fetchone()
            
            # Last 7 days for chart
            chart_data = []
            for i in range(6, -1, -1):
                day_start = today_start - timedelta(days=i)
                day_end = day_start + timedelta(days=1)
                cursor.execute("""
                    SELECT COUNT(*) as count
                    FROM conversions
                    WHERE created_at >= %s AND created_at < %s
                """, (day_start, day_end))
                result = cursor.fetchone()
                chart_data.append({
                    'date': day_start.strftime('%Y-%m-%d'),
                    'label': day_start.strftime('%d/%m'),
                    'conversions': result['count'] or 0
                })
            
            return jsonify({
                'success': True,
                'today': {
                    'conversions': today['conversions'] or 0,
                    'characters': today['characters'] or 0,
                    'audio_size_mb': round((today['audio_size'] or 0) / (1024 * 1024), 2),
                    'duration_hours': round((today['duration'] or 0) / 3600, 2)
                },
                'week': {
                    'conversions': week['conversions'] or 0,
                    'characters': week['characters'] or 0,
                    'audio_size_mb': round((week['audio_size'] or 0) / (1024 * 1024), 2),
                    'duration_hours': round((week['duration'] or 0) / 3600, 2)
                },
                'month': {
                    'conversions': month['conversions'] or 0,
                    'characters': month['characters'] or 0,
                    'audio_size_mb': round((month['audio_size'] or 0) / (1024 * 1024), 2),
                    'duration_hours': round((month['duration'] or 0) / 3600, 2)
                },
                'chart_data': chart_data
            })
    except Exception as e:
        print(f"[ERROR] Get time-based statistics error: {e}")
        import traceback
        print(traceback.format_exc())
        return jsonify({'success': False, 'message': f'Lỗi: {str(e)}'}), 500
    finally:
        conn.close()

@app.route('/api/statistics/top-rankings')
def get_top_rankings():
    """Lấy top rankings (users, voices)"""
    if not is_logged_in() or not is_admin():
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': 'Database connection failed'}), 500
    
    try:
        with conn.cursor() as cursor:
            # Top 5 users by conversions
            cursor.execute("""
                SELECT u.id, u.username, u.full_name, u.email,
                       COUNT(c.id) as conversion_count,
                       SUM(c.text_length) as total_characters,
                       SUM(c.duration_seconds) as total_duration
                FROM users u
                LEFT JOIN conversions c ON u.id = c.user_id
                GROUP BY u.id, u.username, u.full_name, u.email
                ORDER BY conversion_count DESC
                LIMIT 5
            """)
            top_users = cursor.fetchall()
            
            # Top 5 voices by usage - chỉ lấy các giọng đã được sử dụng
            cursor.execute("""
                SELECT voice_id, voice_name,
                       COUNT(*) as usage_count,
                       SUM(text_length) as total_characters,
                       SUM(duration_seconds) as total_duration
                FROM conversions
                WHERE voice_id IS NOT NULL 
                  AND voice_name IS NOT NULL 
                  AND voice_name != 'null'
                  AND voice_name != ''
                GROUP BY voice_id, voice_name
                HAVING COUNT(*) > 0
                ORDER BY usage_count DESC
                LIMIT 5
            """)
            top_voices = cursor.fetchall()
            
            # Voice distribution (for pie chart) - chỉ các giọng đã được sử dụng
            cursor.execute("""
                SELECT voice_name, COUNT(*) as count
                FROM conversions
                WHERE voice_name IS NOT NULL 
                  AND voice_name != 'null'
                  AND voice_name != ''
                GROUP BY voice_name
                HAVING COUNT(*) > 0
                ORDER BY count DESC
            """)
            voice_distribution = cursor.fetchall()
            
            return jsonify({
                'success': True,
                'top_users': [
                    {
                        'id': u['id'],
                        'username': u['username'],
                        'full_name': u['full_name'] or '',
                        'email': u['email'],
                        'conversion_count': u['conversion_count'] or 0,
                        'total_characters': u['total_characters'] or 0,
                        'total_duration_hours': round((u['total_duration'] or 0) / 3600, 2)
                    }
                    for u in top_users
                ],
                'top_voices': [
                    {
                        'voice_id': v['voice_id'],
                        'voice_name': v['voice_name'],
                        'usage_count': v['usage_count'] or 0,
                        'total_characters': v['total_characters'] or 0,
                        'total_duration_hours': round((v['total_duration'] or 0) / 3600, 2)
                    }
                    for v in top_voices
                ],
                'voice_distribution': [
                    {
                        'voice_name': v['voice_name'],
                        'count': v['count']
                    }
                    for v in voice_distribution
                ]
            })
    except Exception as e:
        print(f"[ERROR] Get top rankings error: {e}")
        import traceback
        print(traceback.format_exc())
        return jsonify({'success': False, 'message': f'Lỗi: {str(e)}'}), 500
    finally:
        conn.close()

@app.route('/api/admin/users', methods=['GET'])
def get_all_users():
    """Lấy danh sách tất cả người dùng (chỉ admin)"""
    if not is_logged_in() or not is_admin():
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': 'Database connection failed'}), 500
    
    try:
        with conn.cursor() as cursor:
            cursor.execute("""
                SELECT id, username, email, full_name, role, is_active, 
                       created_at, updated_at,
                       (SELECT COUNT(*) FROM conversions WHERE user_id = users.id) as total_conversions
                FROM users
                ORDER BY id ASC
            """)
            users = cursor.fetchall()
            
            # Convert to list of dicts
            users_list = []
            for user in users:
                users_list.append({
                    'id': user['id'],
                    'username': user['username'],
                    'email': user['email'],
                    'full_name': user['full_name'] or '',
                    'role': user['role'],
                    'is_active': bool(user['is_active']),
                    'created_at': user['created_at'].isoformat() if user['created_at'] else None,
                    'updated_at': user['updated_at'].isoformat() if user['updated_at'] else None,
                    'total_conversions': user['total_conversions'] or 0
                })
            
            return jsonify({
                'success': True,
                'users': users_list
            })
    except Exception as e:
        print(f"[ERROR] Get all users error: {e}")
        return jsonify({'success': False, 'message': f'Lỗi: {str(e)}'}), 500
    finally:
        conn.close()

@app.route('/api/admin/users/<int:user_id>/role', methods=['PUT'])
def update_user_role(user_id):
    """Cập nhật vai trò của người dùng (chỉ admin)"""
    if not is_logged_in() or not is_admin():
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401
    
    data = request.get_json()
    new_role = data.get('role')
    
    if new_role not in ['admin', 'user']:
        return jsonify({'success': False, 'message': 'Vai trò không hợp lệ'}), 400
    
    # Không cho phép tự thay đổi role của chính mình
    if user_id == session['user_id']:
        return jsonify({'success': False, 'message': 'Không thể thay đổi vai trò của chính mình'}), 400
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': 'Database connection failed'}), 500
    
    try:
        with conn.cursor() as cursor:
            # Check if user exists
            cursor.execute("SELECT id FROM users WHERE id = %s", (user_id,))
            if not cursor.fetchone():
                return jsonify({'success': False, 'message': 'Người dùng không tồn tại'}), 404
            
            # Update role
            cursor.execute("UPDATE users SET role = %s WHERE id = %s", (new_role, user_id))
            conn.commit()
            
            return jsonify({
                'success': True,
                'message': f'Đã cập nhật vai trò thành {new_role}'
            })
    except Exception as e:
        conn.rollback()
        print(f"[ERROR] Update user role error: {e}")
        return jsonify({'success': False, 'message': f'Lỗi: {str(e)}'}), 500
    finally:
        conn.close()

@app.route('/api/admin/users/<int:user_id>/status', methods=['PUT'])
def update_user_status(user_id):
    """Cập nhật trạng thái (khóa/mở khóa) của người dùng (chỉ admin)"""
    if not is_logged_in() or not is_admin():
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401
    
    data = request.get_json()
    is_active = data.get('is_active')
    
    if is_active not in [True, False]:
        return jsonify({'success': False, 'message': 'Trạng thái không hợp lệ'}), 400
    
    # Không cho phép tự khóa tài khoản của chính mình
    if user_id == session['user_id']:
        return jsonify({'success': False, 'message': 'Không thể khóa tài khoản của chính mình'}), 400
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': 'Database connection failed'}), 500
    
    try:
        with conn.cursor() as cursor:
            # Check if user exists
            cursor.execute("SELECT id FROM users WHERE id = %s", (user_id,))
            if not cursor.fetchone():
                return jsonify({'success': False, 'message': 'Người dùng không tồn tại'}), 404
            
            # Update status
            cursor.execute("UPDATE users SET is_active = %s WHERE id = %s", (is_active, user_id))
            conn.commit()
            
            status_text = 'kích hoạt' if is_active else 'khóa'
            return jsonify({
                'success': True,
                'message': f'Đã {status_text} tài khoản'
            })
    except Exception as e:
        conn.rollback()
        print(f"[ERROR] Update user status error: {e}")
        return jsonify({'success': False, 'message': f'Lỗi: {str(e)}'}), 500
    finally:
        conn.close()

@app.route('/api/admin/users/<int:user_id>', methods=['DELETE'])
def delete_user(user_id):
    """Xóa người dùng (chỉ admin)"""
    if not is_logged_in() or not is_admin():
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401
    
    # Không cho phép tự xóa tài khoản của chính mình
    if user_id == session['user_id']:
        return jsonify({'success': False, 'message': 'Không thể xóa tài khoản của chính mình'}), 400
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': 'Database connection failed'}), 500
    
    try:
        with conn.cursor() as cursor:
            # Check if user exists
            cursor.execute("SELECT id, username FROM users WHERE id = %s", (user_id,))
            user = cursor.fetchone()
            if not user:
                return jsonify({'success': False, 'message': 'Người dùng không tồn tại'}), 404
            
            # Delete user (cascade will delete conversions)
            cursor.execute("DELETE FROM users WHERE id = %s", (user_id,))
            conn.commit()
            
            return jsonify({
                'success': True,
                'message': f'Đã xóa người dùng {user["username"]}'
            })
    except Exception as e:
        conn.rollback()
        print(f"[ERROR] Delete user error: {e}")
        return jsonify({'success': False, 'message': f'Lỗi: {str(e)}'}), 500
    finally:
        conn.close()

@app.route('/api/admin/account-deletions', methods=['GET'])
def admin_list_account_deletions():
    """Danh sách yêu cầu xóa tài khoản đang chờ duyệt (chỉ admin)"""
    if not is_logged_in() or not is_admin():
        return jsonify({'success': False, 'message': 'Không có quyền truy cập'}), 403

    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': 'Database connection failed'}), 500

    try:
        with conn.cursor() as cursor:
            cursor.execute("""
                SELECT id, username, email, full_name, delete_reason, delete_requested_at, delete_status
                FROM users
                WHERE delete_status = 'pending'
                ORDER BY delete_requested_at ASC
            """)
            rows = cursor.fetchall()

        requests = []
        for row in rows:
            requests.append({
                'id': row['id'],
                'username': row['username'],
                'email': row.get('email') or '',
                'full_name': row.get('full_name') or '',
                'delete_reason': row.get('delete_reason') or '',
                'delete_requested_at': _format_datetime_vn(row.get('delete_requested_at')),
                'delete_status': row.get('delete_status') or 'pending',
            })
        return jsonify({'success': True, 'requests': requests})
    except Exception as e:
        print(f"[ERROR] admin_list_account_deletions: {e}")
        return jsonify({'success': False, 'message': f'Lỗi: {str(e)}'}), 500
    finally:
        conn.close()

@app.route('/api/admin/account-deletions/<int:user_id>/approve', methods=['POST'])
def admin_approve_account_deletion(user_id):
    """Admin duyệt xóa tài khoản — vô hiệu hóa, không xóa vật lý"""
    if not is_logged_in() or not is_admin():
        return jsonify({'success': False, 'message': 'Không có quyền truy cập'}), 403

    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': 'Database connection failed'}), 500

    try:
        with conn.cursor() as cursor:
            cursor.execute("SELECT * FROM users WHERE id = %s", (user_id,))
            user = cursor.fetchone()
            if not user:
                return jsonify({'success': False, 'message': 'Người dùng không tồn tại'}), 404

            if _is_account_deleted(user):
                return jsonify({'success': False, 'message': 'Tài khoản đã bị xóa hoặc vô hiệu hóa'}), 400

            if (user.get('delete_status') or 'none') != 'pending':
                return jsonify({'success': False, 'message': 'Yêu cầu xóa không ở trạng thái chờ duyệt'}), 400

            if user.get('role') == 'admin':
                return jsonify({'success': False, 'message': 'Không thể xóa tài khoản admin'}), 400

            requested_at = user.get('delete_requested_at')
            now = datetime.now()
            effective_at = now + timedelta(days=ACCOUNT_DELETION_GRACE_DAYS)
            cursor.execute("""
                UPDATE users SET
                    is_active = 0,
                    status = 'deactivated',
                    delete_status = 'approved',
                    deleted_at = %s,
                    deletion_effective_at = %s,
                    restore_requested = 0,
                    restore_requested_at = NULL
                WHERE id = %s
            """, (now, effective_at, user_id))
            conn.commit()

            user['deleted_at'] = now
            user['deletion_effective_at'] = effective_at
            email_sent = _send_user_delete_approved_email(user, requested_at, now, effective_at)

        msg = f'Đã vô hiệu hóa tài khoản trong {ACCOUNT_DELETION_GRACE_DAYS} ngày (có thể khôi phục)'
        if not email_sent:
            msg += ' (gửi email xác nhận thất bại — kiểm tra cấu hình SMTP)'
        return jsonify({'success': True, 'message': msg, 'email_sent': email_sent})
    except Exception as e:
        conn.rollback()
        print(f"[ERROR] admin_approve_account_deletion: {e}")
        return jsonify({'success': False, 'message': f'Lỗi: {str(e)}'}), 500
    finally:
        conn.close()

@app.route('/api/admin/account-deletions/<int:user_id>/reject', methods=['POST'])
def admin_reject_account_deletion(user_id):
    """Admin từ chối yêu cầu xóa tài khoản"""
    if not is_logged_in() or not is_admin():
        return jsonify({'success': False, 'message': 'Không có quyền truy cập'}), 403

    data = request.get_json() or {}
    note = (data.get('note') or data.get('admin_delete_note') or '').strip()

    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': 'Database connection failed'}), 500

    try:
        with conn.cursor() as cursor:
            cursor.execute("SELECT * FROM users WHERE id = %s", (user_id,))
            user = cursor.fetchone()
            if not user:
                return jsonify({'success': False, 'message': 'Người dùng không tồn tại'}), 404

            if (user.get('delete_status') or 'none') != 'pending':
                return jsonify({'success': False, 'message': 'Yêu cầu xóa không ở trạng thái chờ duyệt'}), 400

            cursor.execute("""
                UPDATE users SET
                    delete_requested = 0,
                    delete_status = 'rejected',
                    admin_delete_note = %s
                WHERE id = %s
            """, (note or None, user_id))
            conn.commit()

            email_sent = _send_user_delete_rejected_email(user, note)

        msg = 'Đã từ chối yêu cầu xóa tài khoản'
        if not email_sent and (user.get('email') or '').strip():
            msg += ' (gửi email thông báo thất bại — kiểm tra cấu hình SMTP)'
        return jsonify({'success': True, 'message': msg, 'email_sent': email_sent})
    except Exception as e:
        conn.rollback()
        print(f"[ERROR] admin_reject_account_deletion: {e}")
        return jsonify({'success': False, 'message': f'Lỗi: {str(e)}'}), 500
    finally:
        conn.close()

@app.route('/api/admin/account-deletions/grace-period', methods=['GET'])
def admin_list_grace_period_accounts():
    """Danh sách tài khoản đang vô hiệu hóa (30 ngày, có thể khôi phục)"""
    if not is_logged_in() or not is_admin():
        return jsonify({'success': False, 'message': 'Không có quyền truy cập'}), 403

    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': 'Database connection failed'}), 500

    try:
        with conn.cursor() as cursor:
            _finalize_expired_account_deletions(cursor)
            conn.commit()
            cursor.execute("""
                SELECT id, username, email, full_name, deleted_at, deletion_effective_at,
                       restore_requested, restore_requested_at, delete_reason
                FROM users
                WHERE delete_status = 'approved'
                  AND status = 'deactivated'
                  AND (deletion_effective_at IS NULL OR deletion_effective_at > NOW())
                ORDER BY deletion_effective_at ASC
            """)
            rows = cursor.fetchall()

        accounts = []
        for row in rows:
            accounts.append({
                'id': row['id'],
                'username': row['username'],
                'email': row.get('email') or '',
                'full_name': row.get('full_name') or '',
                'deleted_at': _format_datetime_vn(row.get('deleted_at')),
                'deletion_effective_at': _format_datetime_vn(row.get('deletion_effective_at')),
                'restore_requested': bool(row.get('restore_requested')),
                'restore_requested_at': _format_datetime_vn(row.get('restore_requested_at')),
                'delete_reason': row.get('delete_reason') or '',
            })
        return jsonify({'success': True, 'accounts': accounts})
    except Exception as e:
        print(f"[ERROR] admin_list_grace_period_accounts: {e}")
        return jsonify({'success': False, 'message': f'Lỗi: {str(e)}'}), 500
    finally:
        conn.close()

@app.route('/api/admin/account-deletions/<int:user_id>/restore', methods=['POST'])
def admin_restore_account(user_id):
    """Admin khôi phục tài khoản trong thời gian grace period"""
    if not is_logged_in() or not is_admin():
        return jsonify({'success': False, 'message': 'Không có quyền truy cập'}), 403

    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': 'Database connection failed'}), 500

    try:
        with conn.cursor() as cursor:
            cursor.execute("SELECT * FROM users WHERE id = %s", (user_id,))
            user = cursor.fetchone()
            if not user:
                return jsonify({'success': False, 'message': 'Người dùng không tồn tại'}), 404

            if not _is_in_deletion_grace(user):
                if _is_account_deleted(user):
                    return jsonify({'success': False, 'message': 'Tài khoản đã bị xóa vĩnh viễn, không thể khôi phục'}), 400
                return jsonify({'success': False, 'message': 'Tài khoản không ở trạng thái chờ xóa (grace period)'}), 400

            cursor.execute("""
                UPDATE users SET
                    is_active = 1,
                    status = 'active',
                    delete_requested = 0,
                    delete_status = 'none',
                    delete_requested_at = NULL,
                    delete_reason = NULL,
                    deleted_at = NULL,
                    deletion_effective_at = NULL,
                    admin_delete_note = NULL,
                    restore_requested = 0,
                    restore_requested_at = NULL
                WHERE id = %s
            """, (user_id,))
            conn.commit()

            email_sent = _send_user_restored_email(user)

        msg = 'Đã khôi phục tài khoản thành công'
        if not email_sent and (user.get('email') or '').strip():
            msg += ' (gửi email thông báo thất bại — kiểm tra cấu hình SMTP)'
        return jsonify({'success': True, 'message': msg, 'email_sent': email_sent})
    except Exception as e:
        conn.rollback()
        print(f"[ERROR] admin_restore_account: {e}")
        return jsonify({'success': False, 'message': f'Lỗi: {str(e)}'}), 500
    finally:
        conn.close()

@app.route('/api/public/request-account-restore', methods=['POST'])
def public_request_account_restore():
    """User yêu cầu khôi phục tài khoản (không cần đăng nhập)"""
    data = request.get_json() or {}
    email = (data.get('email') or '').strip().lower()
    if not email:
        return jsonify({'success': False, 'message': 'Vui lòng nhập email'}), 400

    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': 'Database connection failed'}), 500

    try:
        with conn.cursor() as cursor:
            _finalize_expired_account_deletions(cursor)
            conn.commit()
            cursor.execute(
                "SELECT * FROM users WHERE LOWER(TRIM(email)) = %s",
                (email,),
            )
            user = cursor.fetchone()
            if not user:
                return jsonify({'success': False, 'message': 'Không tìm thấy tài khoản với email này'}), 404

            if _is_account_deleted(user):
                return jsonify({
                    'success': False,
                    'message': 'Tài khoản đã bị xóa vĩnh viễn và không thể khôi phục.'
                }), 400

            if not _is_in_deletion_grace(user):
                return jsonify({
                    'success': False,
                    'message': 'Tài khoản không ở trạng thái vô hiệu hóa tạm thời (30 ngày).'
                }), 400

            if user.get('restore_requested'):
                return jsonify({
                    'success': False,
                    'message': 'Bạn đã gửi yêu cầu khôi phục. Vui lòng chờ admin xử lý.'
                }), 400

            now = datetime.now()
            cursor.execute("""
                UPDATE users SET restore_requested = 1, restore_requested_at = %s
                WHERE id = %s
            """, (now, user['id']))
            conn.commit()

            user['restore_requested_at'] = now
            email_sent = _send_admin_restore_request_email(user)

        msg = 'Yêu cầu khôi phục đã được gửi. Admin sẽ xử lý trong thời gian sớm nhất.'
        if not email_sent:
            msg += ' (gửi email thông báo admin thất bại — kiểm tra cấu hình SMTP)'
        return jsonify({'success': True, 'message': msg, 'email_sent': email_sent})
    except Exception as e:
        conn.rollback()
        print(f"[ERROR] public_request_account_restore: {e}")
        return jsonify({'success': False, 'message': f'Lỗi: {str(e)}'}), 500
    finally:
        conn.close()

@app.route('/history')
def history():
    """Trang lịch sử"""
    if not is_logged_in():
        return redirect(url_for('login'))
    return render_template('history.html')

@app.route('/admin')
def admin():
    """Trang quản trị"""
    if not is_logged_in() or not is_admin():
        return redirect(url_for('index'))
    return render_template('admin.html')


# ── LANDING PAGE CONTENT ──────────────────────────────────────────
LANDING_CONTENT_FILE = os.path.join(os.path.dirname(__file__), 'landing_content.json')
LANDING_CONTENT_EN_FILE = os.path.join(os.path.dirname(__file__), 'landing_content_en.json')

def load_landing_content():
    """Đọc nội dung landing page từ JSON"""
    try:
        with open(LANDING_CONTENT_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception:
        return {}

def save_landing_content(data):
    """Lưu nội dung landing page vào JSON"""
    with open(LANDING_CONTENT_FILE, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    try:
        if os.path.exists(LANDING_CONTENT_EN_FILE):
            os.remove(LANDING_CONTENT_EN_FILE)
    except Exception:
        pass


def _landing_content_hash(data):
    import hashlib
    raw = json.dumps(data, ensure_ascii=False, sort_keys=True)
    return hashlib.sha256(raw.encode('utf-8')).hexdigest()


def _load_landing_en_cache():
    try:
        with open(LANDING_CONTENT_EN_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception:
        return None


def _save_landing_en_cache(vi_hash, en_lp):
    with open(LANDING_CONTENT_EN_FILE, 'w', encoding='utf-8') as f:
        json.dump({'vi_hash': vi_hash, 'lp': en_lp}, f, ensure_ascii=False, indent=2)


_LANDING_SKIP_KEYS = frozenset({
    'href', 'icon', 'mst', 'phone', 'since', 'representative', 'address', 'email',
})


def _translate_landing_dict(data):
    """Dịch đệ quy nội dung landing (cache theo từng chuỗi trong translate_service)."""
    from translate_service import translate_text_safe
    if isinstance(data, list):
        return [_translate_landing_dict(item) for item in data]
    if isinstance(data, dict):
        out = {}
        for key, val in data.items():
            if key in _LANDING_SKIP_KEYS:
                out[key] = val
            else:
                out[key] = _translate_landing_dict(val)
        return out
    if isinstance(data, str) and data.strip():
        return translate_text_safe(data, 'en')
    return data


def get_landing_for_display(lang='vi'):
    """Nội dung landing theo ngôn ngữ UI."""
    import copy
    vi_lp = load_landing_content()
    lang = (lang or 'vi').strip().lower()
    if lang != 'en':
        return vi_lp
    vi_hash = _landing_content_hash(vi_lp)
    cached = _load_landing_en_cache()
    if cached and cached.get('vi_hash') == vi_hash and cached.get('lp'):
        return cached['lp']
    en_lp = _translate_landing_dict(copy.deepcopy(vi_lp))
    _save_landing_en_cache(vi_hash, en_lp)
    return en_lp

# ── SITE SETTINGS & LEGAL CONTENT ─────────────────────────────────
SITE_SETTINGS_FILE = os.path.join(os.path.dirname(__file__), 'site_settings.json')
LEGAL_CONTENT_FILE = os.path.join(os.path.dirname(__file__), 'legal_content.json')
LEGAL_CONTENT_EN_FILE = os.path.join(os.path.dirname(__file__), 'legal_content_en.json')
SUPPORT_CONTENT_FILE = os.path.join(os.path.dirname(__file__), 'support_content.json')
SUPPORT_CONTENT_EN_FILE = os.path.join(os.path.dirname(__file__), 'support_content_en.json')
try:
    from legal_defaults import get_legal_defaults
except ImportError:
    def get_legal_defaults():
        return {}
try:
    from support_defaults import get_support_defaults
except ImportError:
    def get_support_defaults():
        return {}
try:
    from support_defaults_en import get_support_defaults_en
except ImportError:
    def get_support_defaults_en():
        return {}
SITE_LOGO_DIR = os.path.join(os.path.dirname(__file__), 'static', 'img')
ALLOWED_LOGO_EXTENSIONS = {'png', 'jpg', 'jpeg', 'webp', 'svg'}
MAX_LOGO_BYTES = 2 * 1024 * 1024

DEFAULT_SITE_SETTINGS = {
    'site_name': 'VietVoice',
    'logo_url': '',
    'support_email': 'support@vietvoice.app',
    'contact_email': 'support@vietvoice.app',
    'smtp_from_display': '',
    'company_name': '',
    'company_phone': '',
}

def load_site_settings():
    """Đọc cấu hình site từ JSON"""
    try:
        with open(SITE_SETTINGS_FILE, 'r', encoding='utf-8') as f:
            data = json.load(f)
            return {**DEFAULT_SITE_SETTINGS, **data}
    except Exception:
        return dict(DEFAULT_SITE_SETTINGS)

def save_site_settings(data):
    """Lưu cấu hình site vào JSON"""
    current = load_site_settings()
    for key in DEFAULT_SITE_SETTINGS:
        if key in data and data[key] is not None:
            current[key] = str(data[key]).strip()
    with open(SITE_SETTINGS_FILE, 'w', encoding='utf-8') as f:
        json.dump(current, f, ensure_ascii=False, indent=2)

def get_support_email():
    """Email hỗ trợ (ưu tiên site_settings, fallback env/default)"""
    email = load_site_settings().get('support_email', '').strip()
    if email:
        return email
    return SMTP_FROM or 'support@vietvoice.app'


def get_contact_email():
    """Email liên hệ form (ưu tiên site_settings, fallback support)"""
    email = load_site_settings().get('contact_email', '').strip()
    if email:
        return email
    return get_support_email()


@app.template_global('support_email_addr')
def template_global_support_email():
    """Luôn đọc email hỗ trợ mới nhất từ site_settings.json khi render template."""
    return get_support_email()


@app.template_global('contact_email_addr')
def template_global_contact_email():
    """Luôn đọc email liên hệ mới nhất từ site_settings.json khi render template."""
    return get_contact_email()

def clean_legal_html(html):
    """Loại placeholder và markup rác từ Quill trước khi lưu."""
    if not html:
        return ''
    text = str(html)
    text = re.sub(
        r'<p>\s*(Nội dung (cảnh báo|ghi chú)\.\.\.|'
        r'⚠️ Thêm nội dung cảnh báo quan trọng tại đây\.|'
        r'ℹ️ Thêm nội dung ghi chú hoặc lưu ý tại đây\.)\s*</p>',
        '',
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(r'<p>\s*dung cảnh báo\.\.\.\s*</p>', '', text, flags=re.IGNORECASE)
    text = re.sub(r'<p>\s*(<br\s*/?>)?\s*</p>', '', text, flags=re.IGNORECASE)
    text = re.sub(r'<div class="warning-box">\s*<br\s*/?></div>', '', text, flags=re.IGNORECASE)
    text = re.sub(r'<div class="warning-box">\s*</div>', '', text, flags=re.IGNORECASE)
    text = re.sub(r'<span class="ql-ui"[^>]*></span>', '', text)
    return text.strip()


def normalize_legal_page(page):
    """Loại mục rỗng trước khi lưu / hiển thị."""
    default = {'updated': '', 'sections': [], 'body_html': ''}
    if not page:
        return dict(default)
    page = {**default, **page}
    cleaned_sections = []
    for section in (page.get('sections') or []):
        sec = dict(section)
        sec['title'] = (sec.get('title') or '').strip()
        sec['content'] = clean_legal_html(sec.get('content') or '')
        if legal_section_has_content(sec):
            cleaned_sections.append(sec)
    page['sections'] = cleaned_sections
    body = clean_legal_html(page.get('body_html') or '')
    page['body_html'] = body if _strip_html_text(body) else ''
    return page


def normalize_legal_content(data):
    """Chuẩn hóa toàn bộ legal_content.json."""
    normalized = {}
    for key in LEGAL_CONTENT_PAGE_KEYS:
        normalized[key] = normalize_legal_page(data.get(key))
    return normalized


def resolve_legal_placeholders(text):
    """Thay placeholder trong nội dung pháp lý (email, URL động)."""
    if not text:
        return text
    try:
        pricing_url = url_for('pricing')
        contact_url = url_for('contact')
        support_url = url_for('support')
    except RuntimeError:
        pricing_url = '/pricing'
        contact_url = '/contact'
        support_url = '/support'
    replacements = {
        '__SUPPORT_EMAIL__': get_support_email(),
        '__CONTACT_EMAIL__': get_contact_email(),
        '__PRICING_URL__': pricing_url,
        '__CONTACT_URL__': contact_url,
        '__SUPPORT_URL__': support_url,
    }
    for key, value in replacements.items():
        text = text.replace(key, value)
    return text


def resolve_legal_page(page):
    """Áp placeholder cho một trang pháp lý trước khi hiển thị."""
    if not page:
        return page
    page = dict(page)
    page['sections'] = [
        {
            **section,
            'title': resolve_legal_placeholders(section.get('title') or ''),
            'content': resolve_legal_placeholders(section.get('content') or ''),
        }
        for section in (page.get('sections') or [])
    ]
    body = page.get('body_html') or ''
    if body:
        page['body_html'] = resolve_legal_placeholders(body)
    return page


def merge_legal_page_with_defaults(page, page_key):
    """Gộp nội dung mặc định (form cũ) khi chưa có bản tùy chỉnh."""
    if page_key in ('user_guide', 'installation_guide'):
        from guide_content_loader import load_guide_from_markdown
        md_page = load_guide_from_markdown(page_key)
        if md_page and (md_page.get('sections') or md_page.get('body_html')):
            return resolve_legal_page(md_page)
    defaults = get_legal_defaults().get(page_key) or {}
    base = page or {}
    merged = {
        'updated': base.get('updated') or defaults.get('updated') or '',
        'sections': [dict(s) for s in defaults.get('sections', [])],
        'body_html': defaults.get('body_html') or '',
    }
    return resolve_legal_page(merged)


def get_legal_for_admin():
    """Nội dung cho admin — luôn có form/sections mặc định nếu chưa lưu tùy chỉnh."""
    content = load_legal_content()
    result = {}
    for key in LEGAL_CONTENT_PAGE_KEYS:
        page = content.get(key) or {}
        if key in GUIDE_MD_PAGE_KEYS:
            result[key] = merge_legal_page_with_defaults(page, key)
        elif legal_page_has_custom_content(page):
            result[key] = resolve_legal_page(page)
        else:
            result[key] = merge_legal_page_with_defaults(page, key)
    return result


def merge_legal_page_with_defaults_en(page, page_key):
    """English defaults for public legal pages."""
    from legal_defaults_en import get_legal_defaults_en
    defaults = get_legal_defaults_en().get(page_key) or {}
    merged = {
        'updated': defaults.get('updated') or '',
        'sections': [dict(s) for s in defaults.get('sections', [])],
        'body_html': defaults.get('body_html') or '',
    }
    return resolve_legal_page(merged)


def _translate_legal_page_to_en(vi_page):
    """Translate custom Vietnamese legal content to English via LLM (cached server-side)."""
    from translate_service import translate_text_safe, translate_text_chunked_safe
    sections = []
    for section in vi_page.get('sections') or []:
        title = section.get('title') or ''
        content = section.get('content') or ''
        sections.append({
            'title': translate_text_safe(title, 'en'),
            'content': translate_text_chunked_safe(content, 'en'),
        })
    updated = vi_page.get('updated') or ''
    body = vi_page.get('body_html') or ''
    return resolve_legal_page({
        'updated': translate_text_safe(updated, 'en') if updated else '',
        'sections': sections,
        'body_html': translate_text_chunked_safe(body, 'en') if body else '',
    })


def _get_legal_en_for_template(page_key, vi_page):
    """EN embed for Jinja — only use cache; never block page render on live translation."""
    vi_hash = _legal_page_hash(vi_page)
    cached = _load_legal_en_cache().get(page_key)
    if cached and cached.get('vi_hash') == vi_hash and cached.get('page'):
        return cached['page']
    return vi_page


def _legal_translation_applied(en_page, vi_page):
    """Return True if EN page differs from VI (LLM translation likely succeeded)."""
    vi_sections = vi_page.get('sections') or []
    en_sections = en_page.get('sections') or []
    if not vi_sections or len(vi_sections) != len(en_sections):
        return False
    for vi_sec, en_sec in zip(vi_sections, en_sections):
        if (en_sec.get('title') or '').strip() != (vi_sec.get('title') or '').strip():
            return True
        vi_text = _strip_html_text(vi_sec.get('content') or '')
        en_text = _strip_html_text(en_sec.get('content') or '')
        if en_text and vi_text and en_text != vi_text:
            return True
    return False


def _strip_guide_keys_from_legal_data(data):
    """Xóa nội dung hướng dẫn khỏi JSON — file Markdown là nguồn duy nhất."""
    for key in GUIDE_MD_PAGE_KEYS:
        data[key] = {'updated': '', 'sections': [], 'body_html': ''}


def get_legal_for_display(page_key, lang='vi'):
    """Nội dung hiển thị trang công khai — luôn resolve placeholder, giữ layout legal-card."""
    lang = (lang or 'vi').strip().lower()
    stored = load_legal_content().get(page_key) or {}
    is_guide_md = page_key in GUIDE_MD_PAGE_KEYS
    if is_guide_md:
        vi_page = merge_legal_page_with_defaults(stored, page_key)
    elif legal_page_has_custom_content(stored):
        vi_page = resolve_legal_page(stored)
    else:
        vi_page = merge_legal_page_with_defaults(stored, page_key)

    if lang != 'en':
        return vi_page

    vi_hash = _legal_page_hash(vi_page)
    cached = _load_legal_en_cache().get(page_key)
    if cached and cached.get('vi_hash') == vi_hash and cached.get('page'):
        return cached['page']

    en_defaults = merge_legal_page_with_defaults_en(stored, page_key)
    if is_guide_md:
        translated = _translate_legal_page_to_en(vi_page)
        if _legal_translation_applied(translated, vi_page):
            _save_legal_en_cache(page_key, vi_hash, translated)
            return translated
        _save_legal_en_cache(page_key, vi_hash, en_defaults)
        return en_defaults

    if not legal_page_has_custom_content(stored):
        _save_legal_en_cache(page_key, vi_hash, en_defaults)
        return en_defaults

    translated = _translate_legal_page_to_en(vi_page)
    if _legal_translation_applied(translated, vi_page):
        _save_legal_en_cache(page_key, vi_hash, translated)
        return translated
    _save_legal_en_cache(page_key, vi_hash, en_defaults)
    return en_defaults


def get_legal_for_public(page_key, lang='vi'):
    """Alias — dùng get_legal_for_display để đồng bộ mặc định và nội dung đã sửa."""
    return get_legal_for_display(page_key, lang=lang)

def load_legal_content():
    """Đọc nội dung trang pháp lý từ JSON"""
    default = {key: {'updated': '', 'sections': [], 'body_html': ''} for key in LEGAL_CONTENT_PAGE_KEYS}
    try:
        with open(LEGAL_CONTENT_FILE, 'r', encoding='utf-8') as f:
            data = json.load(f)
            for key in default:
                if key not in data:
                    data[key] = default[key]
                else:
                    data[key] = {**default[key], **data[key]}
            return normalize_legal_content(data)
    except Exception:
        return default

def save_legal_content(data):
    """Lưu nội dung trang pháp lý"""
    with open(LEGAL_CONTENT_FILE, 'w', encoding='utf-8') as f:
        json.dump(normalize_legal_content(data), f, ensure_ascii=False, indent=2)
    try:
        if os.path.exists(LEGAL_CONTENT_EN_FILE):
            os.remove(LEGAL_CONTENT_EN_FILE)
    except Exception:
        pass


def _legal_page_hash(page):
    import hashlib
    raw = json.dumps(page, ensure_ascii=False, sort_keys=True)
    return hashlib.sha256(raw.encode('utf-8')).hexdigest()


def _load_legal_en_cache():
    try:
        with open(LEGAL_CONTENT_EN_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception:
        return {}


def _save_legal_en_cache(page_key, vi_hash, en_page):
    cache = _load_legal_en_cache()
    cache[page_key] = {'vi_hash': vi_hash, 'page': en_page}
    with open(LEGAL_CONTENT_EN_FILE, 'w', encoding='utf-8') as f:
        json.dump(cache, f, ensure_ascii=False, indent=2)


def support_card_has_content(card):
    if not card:
        return False
    if (card.get('title') or '').strip() or (card.get('desc') or '').strip():
        return True
    return bool((card.get('link_text') or '').strip())


def support_guide_has_content(guide):
    if not guide:
        return False
    if (guide.get('title') or '').strip():
        return True
    for step in guide.get('steps') or []:
        if _strip_html_text(step.get('text') or ''):
            return True
    return False


def support_faq_has_content(faq):
    if not faq:
        return False
    if (faq.get('question') or '').strip():
        return True
    return bool(_strip_html_text(faq.get('answer_html') or ''))


def support_has_custom_content(data):
    """Nội dung hỗ trợ có bản tùy chỉnh (không phải file rỗng)."""
    if not data:
        return False
    for card in data.get('contact_cards') or []:
        if support_card_has_content(card):
            return True
    if (data.get('guides_title') or '').strip() or (data.get('faq_title') or '').strip():
        return True
    for guide in data.get('guides') or []:
        if support_guide_has_content(guide):
            return True
    for faq in data.get('faqs') or []:
        if support_faq_has_content(faq):
            return True
    return False


def normalize_support_content(data):
    """Chuẩn hóa support_content.json — loại mục rỗng (không gộp mặc định)."""
    merged = data or {}
    cards = []
    for card in merged.get('contact_cards') or []:
        c = {
            'icon': (card.get('icon') or '').strip() or '📧',
            'title': (card.get('title') or '').strip(),
            'desc': (card.get('desc') or '').strip(),
            'link_text': (card.get('link_text') or '').strip(),
            'action': (card.get('action') or 'custom').strip(),
            'mailto_subject': (card.get('mailto_subject') or '').strip(),
        }
        if support_card_has_content(c):
            cards.append(c)
    guides = []
    for guide in merged.get('guides') or []:
        steps = []
        for step in guide.get('steps') or []:
            text = clean_legal_html(step.get('text') or '')
            if _strip_html_text(text):
                steps.append({'text': text})
        g = {
            'title': (guide.get('title') or '').strip(),
            'steps': steps,
        }
        if support_guide_has_content(g):
            guides.append(g)
    faqs = []
    for faq in merged.get('faqs') or []:
        f = {
            'question': (faq.get('question') or '').strip(),
            'answer_html': clean_legal_html(faq.get('answer_html') or ''),
        }
        if support_faq_has_content(f):
            faqs.append(f)
    return {
        'contact_cards': cards,
        'guides_title': (merged.get('guides_title') or '').strip(),
        'guides': guides,
        'faq_title': (merged.get('faq_title') or '').strip(),
        'faqs': faqs,
    }


def get_pricing_or_register_url():
    try:
        if session.get('user_id'):
            return url_for('pricing')
        return url_for('register')
    except RuntimeError:
        return '/pricing'


def resolve_support_placeholders(text):
    """Thay placeholder trong nội dung hỗ trợ (email, URL động)."""
    if not text:
        return text
    try:
        terms_url = url_for('terms')
        data_deletion_url = url_for('data_deletion')
        pricing_or_register = get_pricing_or_register_url()
    except RuntimeError:
        terms_url = '/terms'
        data_deletion_url = '/data-deletion'
        pricing_or_register = '/pricing'
    text = resolve_legal_placeholders(text)
    replacements = {
        '__TERMS_URL__': terms_url,
        '__DATA_DELETION_URL__': data_deletion_url,
        '__PRICING_OR_REGISTER_URL__': pricing_or_register,
    }
    for key, value in replacements.items():
        text = text.replace(key, value)
    return text


def resolve_support_content(content):
    """Áp placeholder cho toàn bộ nội dung hỗ trợ."""
    if not content:
        return content
    resolved = dict(content)
    resolved['contact_cards'] = [
        {
            **card,
            'title': resolve_support_placeholders(card.get('title') or ''),
            'desc': resolve_support_placeholders(card.get('desc') or ''),
            'link_text': resolve_support_placeholders(card.get('link_text') or ''),
            'mailto_subject': resolve_support_placeholders(card.get('mailto_subject') or ''),
        }
        for card in (content.get('contact_cards') or [])
    ]
    resolved['guides_title'] = resolve_support_placeholders(content.get('guides_title') or '')
    resolved['faq_title'] = resolve_support_placeholders(content.get('faq_title') or '')
    resolved['guides'] = [
        {
            'title': resolve_support_placeholders(guide.get('title') or ''),
            'steps': [
                {'text': resolve_support_placeholders(step.get('text') or '')}
                for step in (guide.get('steps') or [])
            ],
        }
        for guide in (content.get('guides') or [])
    ]
    resolved['faqs'] = [
        {
            'question': resolve_support_placeholders(faq.get('question') or ''),
            'answer_html': resolve_support_placeholders(faq.get('answer_html') or ''),
        }
        for faq in (content.get('faqs') or [])
    ]
    return resolved


def merge_support_with_defaults(stored):
    """Gộp mặc định khi chưa có bản tùy chỉnh."""
    if support_has_custom_content(stored):
        return resolve_support_content(normalize_support_content(stored))
    defaults = get_support_defaults()
    merged = normalize_support_content(defaults)
    return resolve_support_content(merged)


def merge_support_with_defaults_en(stored):
    """English defaults for support page."""
    defaults = get_support_defaults_en()
    merged = normalize_support_content(defaults)
    return resolve_support_content(merged)


def _translate_support_to_en(vi_content):
    from translate_service import translate_text_safe
    cards = []
    for card in vi_content.get('contact_cards') or []:
        cards.append({
            'icon': card.get('icon') or '📧',
            'title': translate_text_safe(card.get('title') or '', 'en'),
            'desc': translate_text_safe(card.get('desc') or '', 'en'),
            'link_text': card.get('link_text') or '',
            'action': card.get('action') or 'custom',
            'mailto_subject': translate_text_safe(card.get('mailto_subject') or '', 'en'),
        })
    guides = []
    for guide in vi_content.get('guides') or []:
        steps = []
        for step in guide.get('steps') or []:
            steps.append({'text': translate_text_safe(step.get('text') or '', 'en')})
        guides.append({
            'title': translate_text_safe(guide.get('title') or '', 'en'),
            'steps': steps,
        })
    faqs = []
    for faq in vi_content.get('faqs') or []:
        faqs.append({
            'question': translate_text_safe(faq.get('question') or '', 'en'),
            'answer_html': translate_text_safe(faq.get('answer_html') or '', 'en'),
        })
    return resolve_support_content({
        'contact_cards': cards,
        'guides_title': translate_text_safe(vi_content.get('guides_title') or '', 'en'),
        'guides': guides,
        'faq_title': translate_text_safe(vi_content.get('faq_title') or '', 'en'),
        'faqs': faqs,
    })


def _support_translation_applied(en_content, vi_content):
    if (en_content.get('guides_title') or '').strip() != (vi_content.get('guides_title') or '').strip():
        return True
    if (en_content.get('faq_title') or '').strip() != (vi_content.get('faq_title') or '').strip():
        return True
    vi_cards = vi_content.get('contact_cards') or []
    en_cards = en_content.get('contact_cards') or []
    if len(vi_cards) == len(en_cards):
        for vi_c, en_c in zip(vi_cards, en_cards):
            if (en_c.get('title') or '').strip() != (vi_c.get('title') or '').strip():
                return True
    vi_faqs = vi_content.get('faqs') or []
    en_faqs = en_content.get('faqs') or []
    if len(vi_faqs) == len(en_faqs):
        for vi_f, en_f in zip(vi_faqs, en_faqs):
            if (en_f.get('question') or '').strip() != (vi_f.get('question') or '').strip():
                return True
    return False


def _support_content_hash(content):
    import hashlib
    raw = json.dumps(content, ensure_ascii=False, sort_keys=True)
    return hashlib.sha256(raw.encode('utf-8')).hexdigest()


def _load_support_en_cache():
    try:
        with open(SUPPORT_CONTENT_EN_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception:
        return {}


def _save_support_en_cache(vi_hash, en_content):
    with open(SUPPORT_CONTENT_EN_FILE, 'w', encoding='utf-8') as f:
        json.dump({'vi_hash': vi_hash, 'content': en_content}, f, ensure_ascii=False, indent=2)


def load_support_content():
    empty = {
        'contact_cards': [],
        'guides_title': '',
        'guides': [],
        'faq_title': '',
        'faqs': [],
    }
    try:
        with open(SUPPORT_CONTENT_FILE, 'r', encoding='utf-8') as f:
            return normalize_support_content(json.load(f))
    except Exception:
        return empty


def save_support_content(data):
    with open(SUPPORT_CONTENT_FILE, 'w', encoding='utf-8') as f:
        json.dump(normalize_support_content(data), f, ensure_ascii=False, indent=2)
    try:
        if os.path.exists(SUPPORT_CONTENT_EN_FILE):
            os.remove(SUPPORT_CONTENT_EN_FILE)
    except Exception:
        pass


def get_support_for_admin():
    stored = load_support_content()
    if support_has_custom_content(stored):
        return resolve_support_content(stored)
    return merge_support_with_defaults(stored)


def get_support_for_display(lang='vi'):
    lang = (lang or 'vi').strip().lower()
    stored = load_support_content()
    if support_has_custom_content(stored):
        vi_content = resolve_support_content(stored)
    else:
        vi_content = merge_support_with_defaults(stored)

    if lang != 'en':
        return vi_content

    vi_hash = _support_content_hash(vi_content)
    cached = _load_support_en_cache()
    if cached.get('vi_hash') == vi_hash and cached.get('content'):
        return cached['content']

    en_defaults = merge_support_with_defaults_en(stored)
    if not support_has_custom_content(stored):
        _save_support_en_cache(vi_hash, en_defaults)
        return en_defaults

    translated = _translate_support_to_en(vi_content)
    if _support_translation_applied(translated, vi_content):
        _save_support_en_cache(vi_hash, translated)
        return translated
    _save_support_en_cache(vi_hash, en_defaults)
    return en_defaults


@app.template_global('support_card_href')
def template_support_card_href(card):
    action = (card.get('action') or 'custom').strip()
    support = get_support_email()
    if action == 'mailto_support':
        return f'mailto:{support}'
    if action == 'mailto_bug':
        subject = (card.get('mailto_subject') or '').strip()
        if subject:
            from urllib.parse import quote
            return f'mailto:{support}?subject={quote(subject)}'
        return f'mailto:{support}'
    if action == 'contact_page':
        try:
            return url_for('contact')
        except RuntimeError:
            return '/contact'
    href = (card.get('href') or '#').strip()
    return resolve_support_placeholders(href) or '#'


@app.template_global('support_link_text')
def template_support_link_text(card):
    text = resolve_support_placeholders(card.get('link_text') or '')
    if text == '__SUPPORT_EMAIL__':
        return get_support_email()
    return text


@app.context_processor
def inject_site_settings():
    """Inject site settings into all templates."""
    try:
        settings = load_site_settings()
    except Exception:
        settings = dict(DEFAULT_SITE_SETTINGS)
    support = get_support_email()
    contact = get_contact_email()
    return {
        'site_settings': settings,
        'support_email': support,
        'contact_email': contact,
        'legal_page_has_custom_content': legal_page_has_custom_content,
    }

@app.route('/admin/settings', methods=['GET'])
def admin_settings():
    """Trang cấu hình site (admin)"""
    if not is_logged_in() or not is_admin():
        return redirect(url_for('index'))
    return render_template('admin_settings.html')

@app.route('/admin/policies', methods=['GET'])
def admin_policies():
    """Trang cấu hình chính sách / pháp lý (lưu file JSON, không dùng database)"""
    if not is_logged_in() or not is_admin():
        return redirect(url_for('index'))
    return render_template('admin_policies.html')

@app.route('/api/admin/settings', methods=['GET'])
def api_admin_get_settings():
    """API: lấy cấu hình site"""
    if not is_logged_in() or not is_admin():
        return jsonify({'success': False, 'message': 'Unauthorized'}), 403
    settings = load_site_settings()
    settings['smtp_from_env'] = SMTP_FROM or ''
    settings['smtp_host_env'] = SMTP_HOST or ''
    return jsonify({'success': True, 'settings': settings})

@app.route('/api/admin/settings', methods=['POST'])
def api_admin_save_settings():
    """API: lưu cấu hình site"""
    if not is_logged_in() or not is_admin():
        return jsonify({'success': False, 'message': 'Unauthorized'}), 403
    try:
        data = request.get_json() or {}
        save_site_settings(data)
        return jsonify({'success': True, 'message': 'Đã lưu cấu hình site thành công!'})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/admin/settings/logo', methods=['POST'])
def api_admin_upload_logo():
    """API: upload logo site"""
    if not is_logged_in() or not is_admin():
        return jsonify({'success': False, 'message': 'Unauthorized'}), 403
    if 'logo' not in request.files:
        return jsonify({'success': False, 'message': 'Không có file logo'}), 400
    file = request.files['logo']
    if not file or not file.filename:
        return jsonify({'success': False, 'message': 'File không hợp lệ'}), 400
    ext = file.filename.rsplit('.', 1)[-1].lower() if '.' in file.filename else ''
    if ext not in ALLOWED_LOGO_EXTENSIONS:
        return jsonify({'success': False, 'message': 'Định dạng logo không hỗ trợ (png, jpg, webp, svg)'}), 400
    file.seek(0, os.SEEK_END)
    size = file.tell()
    file.seek(0)
    if size > MAX_LOGO_BYTES:
        return jsonify({'success': False, 'message': 'Logo quá lớn (tối đa 2MB)'}), 400
    os.makedirs(SITE_LOGO_DIR, exist_ok=True)
    filename = f'site_logo.{ext}'
    filepath = os.path.join(SITE_LOGO_DIR, filename)
    file.save(filepath)
    logo_url = f'img/{filename}'
    save_site_settings({'logo_url': logo_url})
    return jsonify({
        'success': True,
        'message': 'Đã cập nhật logo!',
        'logo_url': logo_url,
        'logo_src': url_for('static', filename=logo_url),
    })

@app.route('/api/admin/legal', methods=['GET'])
def api_admin_get_legal():
    """API: lấy nội dung pháp lý"""
    if not is_logged_in() or not is_admin():
        return jsonify({'success': False, 'message': 'Unauthorized'}), 403
    from guide_content_loader import load_all_guide_markdown_raw
    return jsonify({
        'success': True,
        'legal': get_legal_for_admin(),
        'guide_markdown': load_all_guide_markdown_raw(),
    })

@app.route('/api/admin/legal', methods=['POST'])
def api_admin_save_legal():
    """API: lưu nội dung pháp lý vào legal_content.json (file, không database)"""
    if not is_logged_in() or not is_admin():
        return jsonify({'success': False, 'message': 'Unauthorized'}), 403
    try:
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'message': 'Dữ liệu không hợp lệ'}), 400
        guide_md = data.pop('guide_markdown', None)
        if guide_md:
            from guide_content_loader import save_guide_markdown
            for key in GUIDE_MD_PAGE_KEYS:
                if key in guide_md and guide_md[key]:
                    save_guide_markdown(key, guide_md[key])
        _strip_guide_keys_from_legal_data(data)
        save_legal_content(data)
        from guide_content_loader import load_all_guide_markdown_raw
        return jsonify({
            'success': True,
            'message': 'Đã lưu nội dung pháp lý!',
            'legal': get_legal_for_admin(),
            'guide_markdown': load_all_guide_markdown_raw(),
        })
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/admin/support', methods=['GET'])
def api_admin_get_support():
    """API: lấy nội dung trang hỗ trợ"""
    if not is_logged_in() or not is_admin():
        return jsonify({'success': False, 'message': 'Unauthorized'}), 403
    return jsonify({'success': True, 'support': get_support_for_admin()})


@app.route('/api/admin/support', methods=['POST'])
def api_admin_save_support():
    """API: lưu nội dung hỗ trợ vào support_content.json"""
    if not is_logged_in() or not is_admin():
        return jsonify({'success': False, 'message': 'Unauthorized'}), 403
    try:
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'message': 'Dữ liệu không hợp lệ'}), 400
        save_support_content(data)
        return jsonify({
            'success': True,
            'message': 'Đã lưu nội dung hỗ trợ & FAQ!',
            'support': load_support_content(),
        })
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/admin/packages', methods=['GET'])
def api_admin_get_packages():
    """API: danh sách gói cước (admin, bao gồm inactive)"""
    if not is_logged_in() or not is_admin():
        return jsonify({'success': False, 'message': 'Unauthorized'}), 403
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': 'Database connection failed'}), 500
    try:
        with conn.cursor() as cursor:
            cursor.execute("""
                SELECT id, package_name, characters_limit, price_vnd, duration_days, is_active
                FROM subscription_packages
                ORDER BY characters_limit ASC
            """)
            rows = cursor.fetchall()
            return jsonify({
                'success': True,
                'packages': [
                    {
                        'id': r['id'],
                        'name': r['package_name'],
                        'characters': r['characters_limit'],
                        'price': r['price_vnd'],
                        'duration_days': r['duration_days'],
                        'is_active': bool(r['is_active']),
                    }
                    for r in rows
                ],
            })
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500
    finally:
        conn.close()

@app.route('/api/admin/packages', methods=['POST'])
def api_admin_create_package():
    """API: tạo gói cước mới"""
    if not is_logged_in() or not is_admin():
        return jsonify({'success': False, 'message': 'Unauthorized'}), 403
    data = request.get_json() or {}
    name = (data.get('name') or '').strip()
    characters = int(data.get('characters') or 0)
    price = int(data.get('price') or 0)
    duration_days = int(data.get('duration_days') or 30)
    is_active = 1 if data.get('is_active', True) else 0
    if not name or characters <= 0:
        return jsonify({'success': False, 'message': 'Tên gói và số ký tự là bắt buộc'}), 400
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': 'Database connection failed'}), 500
    try:
        with conn.cursor() as cursor:
            cursor.execute("""
                INSERT INTO subscription_packages (package_name, characters_limit, price_vnd, price, duration_days, is_active)
                VALUES (%s, %s, %s, %s, %s, %s)
            """, (name, characters, price, price, duration_days, is_active))
            conn.commit()
            return jsonify({'success': True, 'message': 'Đã tạo gói cước mới', 'id': cursor.lastrowid})
    except Exception as e:
        conn.rollback()
        return jsonify({'success': False, 'message': str(e)}), 500
    finally:
        conn.close()

@app.route('/api/admin/packages/<int:package_id>', methods=['PUT'])
def api_admin_update_package(package_id):
    """API: cập nhật gói cước"""
    if not is_logged_in() or not is_admin():
        return jsonify({'success': False, 'message': 'Unauthorized'}), 403
    data = request.get_json() or {}
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': 'Database connection failed'}), 500
    try:
        with conn.cursor() as cursor:
            cursor.execute("SELECT id FROM subscription_packages WHERE id = %s", (package_id,))
            if not cursor.fetchone():
                return jsonify({'success': False, 'message': 'Gói không tồn tại'}), 404
            fields = []
            values = []
            if 'name' in data:
                fields.append('package_name = %s')
                values.append(str(data['name']).strip())
            if 'characters' in data:
                fields.append('characters_limit = %s')
                values.append(int(data['characters']))
            if 'price' in data:
                price_val = int(data['price'])
                fields.append('price_vnd = %s')
                values.append(price_val)
                fields.append('price = %s')
                values.append(price_val)
            if 'duration_days' in data:
                fields.append('duration_days = %s')
                values.append(int(data['duration_days']))
            if 'is_active' in data:
                fields.append('is_active = %s')
                values.append(1 if data['is_active'] else 0)
            if not fields:
                return jsonify({'success': False, 'message': 'Không có dữ liệu cập nhật'}), 400
            values.append(package_id)
            cursor.execute(
                f"UPDATE subscription_packages SET {', '.join(fields)} WHERE id = %s",
                tuple(values),
            )
            conn.commit()
            cursor.execute("""
                SELECT id, package_name, characters_limit, price_vnd, duration_days, is_active
                FROM subscription_packages WHERE id = %s
            """, (package_id,))
            updated = cursor.fetchone()
            return jsonify({
                'success': True,
                'message': 'Đã cập nhật gói cước',
                'package': {
                    'id': updated['id'],
                    'name': updated['package_name'],
                    'characters': updated['characters_limit'],
                    'price': updated['price_vnd'],
                    'duration_days': updated['duration_days'],
                    'is_active': bool(updated['is_active']),
                },
            })
    except Exception as e:
        conn.rollback()
        return jsonify({'success': False, 'message': str(e)}), 500
    finally:
        conn.close()

@app.route('/api/admin/packages/<int:package_id>', methods=['DELETE'])
def api_admin_delete_package(package_id):
    """API: vô hiệu hóa gói cước"""
    if not is_logged_in() or not is_admin():
        return jsonify({'success': False, 'message': 'Unauthorized'}), 403
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': 'Database connection failed'}), 500
    try:
        with conn.cursor() as cursor:
            cursor.execute(
                "UPDATE subscription_packages SET is_active = 0 WHERE id = %s",
                (package_id,),
            )
            conn.commit()
            return jsonify({'success': True, 'message': 'Đã ẩn gói cước'})
    except Exception as e:
        conn.rollback()
        return jsonify({'success': False, 'message': str(e)}), 500
    finally:
        conn.close()

@app.route('/admin/landing', methods=['GET'])
def admin_landing():
    """Trang chỉnh sửa nội dung landing page"""
    if not is_logged_in() or not is_admin():
        return redirect(url_for('index'))
    content = load_landing_content()
    return render_template('admin_landing.html', content=content)

@app.route('/admin/landing/save', methods=['POST'])
def admin_landing_save():
    """Lưu nội dung landing page"""
    if not is_logged_in() or not is_admin():
        return jsonify({'success': False, 'message': 'Unauthorized'}), 403
    try:
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'message': 'Dữ liệu không hợp lệ'}), 400
        save_landing_content(data)
        return jsonify({'success': True, 'message': 'Đã lưu nội dung landing page thành công!'})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

# SePay.vn Integration Functions
def create_sepay_payment(amount, transaction_id, description, user_name):
    """Tạo thanh toán qua SePay.vn"""
    try:
        # Tạo nội dung chuyển khoản cho SePay
        content = f"{transaction_id}"
        
        # Tạo QR code bằng VietQR (SePay tương thích với VietQR)
        qr_data = create_sepay_qr_code(amount, content, description)
        
        return {
            'success': True,
            'qr_code': qr_data['qr_image'],
            'bank_info': {
                'bank_name': SEPAY_BANK_ID,
                'account_number': SEPAY_ACCOUNT_NUMBER,
                'account_name': 'TTS SYSTEM',
                'amount': amount,
                'content': content,
                'transaction_id': transaction_id
            },
            'sepay_info': {
                'account_number': SEPAY_ACCOUNT_NUMBER,
                'bank_id': SEPAY_BANK_ID,
                'api_url': SEPAY_API_URL
            }
        }
    except Exception as e:
        print(f"[ERROR] SePay payment creation failed: {e}")
        return {
            'success': False,
            'error': str(e)
        }

def create_sepay_qr_code(amount, content, description):
    """Tạo QR code cho SePay bằng VietQR API"""
    try:
        # Sử dụng VietQR API để tạo QR code cho MBBank (SePay)
        # MBBank BIN code
        mbbank_bin = "970422"
        
        # API VietQR URL
        vietqr_url = f"{SEPAY_QR_API}/{mbbank_bin}-{SEPAY_ACCOUNT_NUMBER}-compact2.jpg?amount={amount}&addInfo={content}&accountName=TTS%20SYSTEM"
        
        print(f"[INFO] Creating SePay QR with URL: {vietqr_url}")
        
        # Download QR image
        response = requests.get(vietqr_url, timeout=10)
        if response.status_code == 200:
            # Convert to base64
            img_base64 = base64.b64encode(response.content).decode()
            return {
                'qr_image': f"data:image/jpeg;base64,{img_base64}",
                'api_url': vietqr_url
            }
        else:
            raise Exception(f"VietQR API returned status {response.status_code}")
            
    except Exception as e:
        print(f"[WARNING] SePay QR generation failed, using fallback: {e}")
        # Fallback to manual QR generation
        return create_manual_qr_code(amount, content)

def create_manual_qr_code(amount, content):
    """Tạo QR code thủ công cho SePay khi API thất bại"""
    try:
        # Tạo QR code đơn giản với thông tin chuyển khoản
        qr_content = f"Account: {SEPAY_ACCOUNT_NUMBER}\\nBank: {SEPAY_BANK_ID}\\nAmount: {amount:,} VND\\nContent: {content}"
        
        qr = qrcode.QRCode(
            version=1,
            error_correction=qrcode.constants.ERROR_CORRECT_L,
            box_size=10,
            border=4,
        )
        qr.add_data(qr_content)
        qr.make(fit=True)
        
        img = qr.make_image(fill_color="black", back_color="white")
        
        # Convert to base64
        buffer = io.BytesIO()
        img.save(buffer, format='PNG')
        img_str = base64.b64encode(buffer.getvalue()).decode()
        
        return {
            'qr_image': f"data:image/png;base64,{img_str}",
            'fallback': True
        }
    except Exception as e:
        print(f"[ERROR] Manual QR generation failed: {e}")
        return {
            'qr_image': None,
            'error': str(e)
        }

def verify_sepay_transaction(transaction_id, amount):
    """Xác minh giao dịch qua SePay API - dùng endpoint /transactions/list"""
    try:
        headers = {
            'Authorization': f'Bearer {SEPAY_TOKEN}',
            'Content-Type': 'application/json'
        }
        
        print(f"[INFO] Verifying SePay transaction: {transaction_id}, amount: {amount}")
        
        # SePay correct endpoint: /transactions/list with transaction_content filter
        response = requests.get(
            f"{SEPAY_API_URL}/list",
            params={
                'account_number': SEPAY_ACCOUNT_NUMBER,
                'transaction_content': transaction_id,
                'limit': 5
            },
            headers=headers,
            timeout=8
        )
        
        print(f"[INFO] SePay API response status: {response.status_code}")
        
        if response.status_code == 200:
            data = response.json()
            print(f"[INFO] SePay API response: {data}")
            
            transactions = data.get('transactions', [])
            
            # Tìm giao dịch khớp: nội dung chứa transaction_id VÀ số tiền đúng
            txn_id_upper = transaction_id.upper()
            txn_id_normalized = ''.join(c for c in txn_id_upper if c.isalnum())
            # Phần hex thuần (bỏ TTS_ prefix)
            hex_part = txn_id_upper.replace('TTS_', '').replace('TTS', '').strip('_- ')

            for txn in transactions:
                content = str(txn.get('transaction_content', '') or '').upper()
                content_normalized = ''.join(c for c in content if c.isalnum())
                txn_amount = int(txn.get('amount_in', 0) or 0)

                content_match = (
                    txn_id_upper in content or
                    txn_id_normalized in content_normalized or
                    (len(hex_part) >= 8 and hex_part in content_normalized)
                )

                if content_match and txn_amount >= int(amount) * 0.99:
                    print(f"[INFO] SePay transaction MATCHED: {txn}")
                    return {
                        'success': True,
                        'verified': True,
                        'transaction_data': txn,
                        'source': 'sepay_api'
                    }
            
            print(f"[INFO] No matching SePay transaction found for {transaction_id}")
            return {'success': True, 'verified': False, 'source': 'sepay_api'}
        else:
            print(f"[WARNING] SePay API returned status {response.status_code}: {response.text}")
            return {'success': True, 'verified': False, 'error': f'SePay API status {response.status_code}'}
            
    except Exception as e:
        print(f"[ERROR] SePay verification failed: {e}")
        return {'success': False, 'verified': False, 'error': str(e)}

def auto_approve_by_time(transaction_id, amount):
    """Auto-approve payment sau một khoảng thời gian (fallback method)"""
    try:
        # Kiểm tra thời gian tạo payment
        conn = get_db_connection()
        if not conn:
            return {'success': False, 'error': 'Database connection failed'}
        
        with conn.cursor() as cursor:
            cursor.execute("""
                SELECT created_at, TIMESTAMPDIFF(MINUTE, created_at, NOW()) as minutes_ago
                FROM payments 
                WHERE transaction_id = %s AND amount_vnd = %s
            """, (transaction_id, amount))
            payment_time = cursor.fetchone()
            
            if payment_time:
                minutes_ago = payment_time['minutes_ago']
                
                # Auto-approve nếu đã quá 5 phút (có thể điều chỉnh)
                AUTO_APPROVE_MINUTES = 5
                
                if minutes_ago >= AUTO_APPROVE_MINUTES:
                    print(f"[INFO] Auto-approving payment after {minutes_ago} minutes")
                    return {
                        'success': True,
                        'verified': True,
                        'auto_approved': True,
                        'reason': f'Auto-approved after {minutes_ago} minutes'
                    }
                else:
                    return {
                        'success': True,
                        'verified': False,
                        'pending_minutes': AUTO_APPROVE_MINUTES - minutes_ago
                    }
            else:
                return {'success': False, 'error': 'Payment not found'}
                
    except Exception as e:
        print(f"[ERROR] Auto-approve by time failed: {e}")
        return {'success': False, 'error': str(e)}
    finally:
        if 'conn' in locals():
            conn.close()

@app.route('/pricing')
def pricing():
    """Trang thanh toán"""
    if not is_logged_in():
        return redirect(url_for('login'))
    return render_template('pricing.html', pricing_packages=fetch_active_packages())

@app.route('/payment/confirm')
def payment_confirm():
    """Trang xác nhận thanh toán"""
    if not is_logged_in():
        return redirect(url_for('login'))

    payment_id = request.args.get('id', type=int)
    if not payment_id:
        return redirect(url_for('pricing'))

    conn = get_db_connection()
    if not conn:
        return redirect(url_for('pricing'))

    try:
        with conn.cursor() as cursor:
            cursor.execute("""
                SELECT p.id AS payment_id, p.transaction_id, p.amount_vnd,
                       p.payment_status, p.created_at,
                       pk.package_name, pk.characters_limit, pk.duration_days
                FROM payments p
                JOIN subscription_packages pk ON p.package_id = pk.id
                WHERE p.id = %s AND p.user_id = %s
            """, (payment_id, session.get('user_id')))
            row = cursor.fetchone()

        if not row:
            return redirect(url_for('pricing'))

        # Generate (or regenerate) QR code for this payment
        from config import SEPAY_BANK_ID, SEPAY_ACCOUNT_NUMBER, BANK_ACCOUNT_NAME
        qr_result = create_sepay_qr_code(
            row['amount_vnd'],
            row['transaction_id'],
            f"Thanh toan {row['package_name']}"
        )

        payment = {
            'payment_id':       row['payment_id'],
            'transaction_id':   row['transaction_id'],
            'amount_vnd':       row['amount_vnd'],
            'payment_status':   row['payment_status'],
            'package_name':     row['package_name'],
            'characters_limit': row['characters_limit'],
            'duration_days':    row['duration_days'],
            'qr_code':          qr_result.get('qr_image', ''),
            'bank_name':        SEPAY_BANK_ID,
            'account_number':   SEPAY_ACCOUNT_NUMBER,
            'account_name':     BANK_ACCOUNT_NAME,
        }
        return render_template('payment_confirmation.html', payment=payment)

    except Exception as e:
        print(f"[ERROR] payment_confirm: {e}")
        return redirect(url_for('pricing'))
    finally:
        conn.close()

@app.route('/profile')
@login_required
def profile():
    """Trang hồ sơ cá nhân"""
    conn = get_db_connection()
    user_data = {}
    if conn:
        try:
            with conn.cursor() as cursor:
                cursor.execute(
                    """SELECT id, username, email, full_name, avatar_url, role, created_at,
                              delete_requested, delete_status, delete_requested_at, admin_delete_note
                       FROM users WHERE id = %s""",
                    (session['user_id'],)
                )
                row = cursor.fetchone()
                if row:
                    user_data = {
                        'id':         row['id'],
                        'username':   row['username'],
                        'email':      row['email'] or '',
                        'full_name':  row['full_name'] or '',
                        'avatar_url': row.get('avatar_url') or '',
                        'role':       row['role'],
                        'created_at': row['created_at'].strftime('%d/%m/%Y') if row.get('created_at') else '',
                        'delete_requested': bool(row.get('delete_requested')),
                        'delete_status': row.get('delete_status') or 'none',
                        'delete_requested_at': _format_datetime_vn(row.get('delete_requested_at')),
                        'admin_delete_note': row.get('admin_delete_note') or '',
                    }
        except Exception as e:
            print(f'[ERROR] profile: {e}')
        finally:
            conn.close()
    if user_data:
        session['avatar_url'] = user_data.get('avatar_url') or ''
    system_voices, custom_voices, emotional_voices = _fetch_profile_voice_options(session['user_id'])
    return render_template(
        'profile.html',
        user=user_data,
        system_voices=system_voices,
        custom_voices=custom_voices,
        emotional_voices=emotional_voices,
    )


def _fetch_profile_voice_options(user_id):
    """Danh sách giọng cho form mặc định TTS trên profile."""
    system_voices = []
    custom_voices = []
    emotional_voices = []
    conn = get_db_connection()
    if not conn:
        return system_voices, custom_voices, emotional_voices
    try:
        with conn.cursor() as cursor:
            cursor.execute(
                "SELECT voice_id, voice_name FROM voices WHERE is_active = 1 ORDER BY voice_id"
            )
            system_voices = cursor.fetchall() or []
            try:
                cursor.execute(
                    """
                    SELECT id, voice_name, voice_type, quality_score
                    FROM custom_voices
                    WHERE user_id = %s AND status = 'completed'
                    ORDER BY created_at DESC
                    """,
                    (user_id,),
                )
            except Exception:
                cursor.execute(
                    """
                    SELECT id, voice_name, quality_score
                    FROM custom_voices
                    WHERE user_id = %s AND status = 'completed'
                    ORDER BY created_at DESC
                    """,
                    (user_id,),
                )
            custom_voices = cursor.fetchall() or []
    except Exception as e:
        print(f"[ERROR] _fetch_profile_voice_options: {e}")
    finally:
        conn.close()
    for v in custom_voices:
        vt = (v.get('voice_type') or 'rvc').strip().lower()
        if vt == 'vixtts_clone':
            emotional_voices.append(v)
    return system_voices, custom_voices, emotional_voices


@app.route('/api/user/account-deletion-status', methods=['GET'])
@login_required
def user_account_deletion_status():
    """Trạng thái yêu cầu xóa tài khoản của user hiện tại"""
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': 'Database connection failed'}), 500
    try:
        with conn.cursor() as cursor:
            cursor.execute(
                """SELECT delete_requested, delete_status, delete_requested_at, admin_delete_note, status
                   FROM users WHERE id = %s""",
                (session['user_id'],),
            )
            row = cursor.fetchone()
            if not row:
                return jsonify({'success': False, 'message': 'Người dùng không tồn tại'}), 404
        return jsonify({
            'success': True,
            'delete_requested': bool(row.get('delete_requested')),
            'delete_status': row.get('delete_status') or 'none',
            'delete_requested_at': _format_datetime_vn(row.get('delete_requested_at')),
            'admin_delete_note': row.get('admin_delete_note') or '',
            'account_deleted': _is_account_deleted(row),
        })
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500
    finally:
        conn.close()


@app.route('/api/user/request-account-deletion', methods=['POST'])
@login_required
def user_request_account_deletion():
    """User gửi yêu cầu xóa tài khoản"""
    data = request.get_json() or {}
    reason = (data.get('delete_reason') or data.get('reason') or '').strip()
    if len(reason) > 2000:
        return jsonify({'success': False, 'message': 'Lý do quá dài (tối đa 2000 ký tự)'}), 400

    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': 'Database connection failed'}), 500

    try:
        with conn.cursor() as cursor:
            cursor.execute("SELECT * FROM users WHERE id = %s", (session['user_id'],))
            user = cursor.fetchone()
            if not user:
                return jsonify({'success': False, 'message': 'Người dùng không tồn tại'}), 404

            if _is_account_deleted(user):
                return jsonify({
                    'success': False,
                    'message': 'Tài khoản của bạn đã bị xóa hoặc vô hiệu hóa.'
                }), 400

            delete_status = (user.get('delete_status') or 'none').lower()
            if delete_status == 'pending':
                return jsonify({
                    'success': False,
                    'message': 'Bạn đã gửi yêu cầu xóa tài khoản. Vui lòng chờ admin xử lý.'
                }), 400

            now = datetime.now()
            cursor.execute("""
                UPDATE users SET
                    delete_requested = 1,
                    delete_status = 'pending',
                    delete_requested_at = %s,
                    delete_reason = %s,
                    admin_delete_note = NULL
                WHERE id = %s
            """, (now, reason or None, session['user_id']))
            conn.commit()

            user['delete_reason'] = reason
            user['delete_requested_at'] = now
            email_sent = _send_admin_delete_request_email(user)

        msg = 'Yêu cầu xóa tài khoản đã được gửi. Admin sẽ xử lý trong thời gian sớm nhất.'
        if not email_sent:
            msg += ' (gửi email thông báo admin thất bại — kiểm tra cấu hình SMTP)'
        return jsonify({'success': True, 'message': msg, 'email_sent': email_sent})
    except Exception as e:
        conn.rollback()
        print(f"[ERROR] user_request_account_deletion: {e}")
        return jsonify({'success': False, 'message': f'Lỗi: {str(e)}'}), 500
    finally:
        conn.close()


@app.route('/api/user/update-profile', methods=['POST'])
@login_required
def update_profile():
    """Cập nhật thông tin cá nhân (full_name)"""
    data = request.get_json() or {}
    full_name = (data.get('full_name') or '').strip()
    if not full_name:
        return jsonify({'success': False, 'message': 'Tên không được để trống'}), 400

    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': 'Database connection failed'}), 500
    try:
        with conn.cursor() as cursor:
            cursor.execute(
                "UPDATE users SET full_name = %s WHERE id = %s",
                (full_name, session['user_id'])
            )
            conn.commit()
        session['full_name'] = full_name
        return jsonify({'success': True, 'message': 'Cập nhật thành công'})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500
    finally:
        conn.close()


@app.route('/uploads/<path:subpath>')
def serve_upload(subpath):
    """Phục vụ file tải lên (avatar, v.v.)"""
    base = UPLOAD_DIR.resolve()
    file_path = (UPLOAD_DIR / subpath).resolve()
    if not str(file_path).startswith(str(base)) or not file_path.is_file():
        abort(404)
    return send_file(file_path)


@app.route('/api/user/upload-avatar', methods=['POST'])
@login_required
def upload_avatar():
    """Upload / thay ảnh đại diện"""
    if 'avatar' not in request.files:
        return jsonify({'success': False, 'message': 'Không có file ảnh'}), 400

    file = request.files['avatar']
    if not file or not file.filename:
        return jsonify({'success': False, 'message': 'Không có file ảnh'}), 400

    ext = file.filename.rsplit('.', 1)[-1].lower() if '.' in file.filename else ''
    if ext not in ALLOWED_AVATAR_EXTENSIONS:
        return jsonify({'success': False, 'message': 'Chỉ chấp nhận JPG, PNG hoặc WEBP'}), 400

    file.seek(0, os.SEEK_END)
    size = file.tell()
    file.seek(0)
    if size > MAX_AVATAR_BYTES:
        return jsonify({'success': False, 'message': 'Ảnh tối đa 2MB'}), 400

    user_id = session['user_id']
    filename = f'user_{user_id}.{ext}'
    save_path = AVATAR_UPLOAD_DIR / filename
    avatar_url = f'/uploads/avatars/{filename}'

    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': 'Database connection failed'}), 500

    try:
        old_url = ''
        with conn.cursor() as cursor:
            cursor.execute("SELECT avatar_url FROM users WHERE id = %s", (user_id,))
            row = cursor.fetchone()
            if row:
                old_url = (row.get('avatar_url') or '').split('?')[0]

        for old_file in AVATAR_UPLOAD_DIR.glob(f'user_{user_id}.*'):
            try:
                old_file.unlink()
            except OSError:
                pass

        file.save(str(save_path))

        with conn.cursor() as cursor:
            cursor.execute(
                "UPDATE users SET avatar_url = %s WHERE id = %s",
                (avatar_url, user_id)
            )
            conn.commit()

        session['avatar_url'] = avatar_url
        return jsonify({
            'success': True,
            'message': 'Cập nhật ảnh đại diện thành công',
            'avatar_url': f'{avatar_url}?t={int(time.time())}'
        })
    except Exception as e:
        print(f'[ERROR] upload_avatar: {e}')
        return jsonify({'success': False, 'message': str(e)}), 500
    finally:
        conn.close()


@app.route('/api/user/change-password', methods=['POST'])
@login_required
def change_password():
    """Đổi mật khẩu"""
    data = request.get_json() or {}
    current_pw  = data.get('current_password', '')
    new_pw      = data.get('new_password', '')
    confirm_pw  = data.get('confirm_password', '')

    if not current_pw or not new_pw:
        return jsonify({'success': False, 'message': 'Vui lòng điền đầy đủ thông tin'}), 400
    if new_pw != confirm_pw:
        return jsonify({'success': False, 'message': 'Mật khẩu xác nhận không khớp'}), 400
    if len(new_pw) < 6:
        return jsonify({'success': False, 'message': 'Mật khẩu phải có ít nhất 6 ký tự'}), 400

    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': 'Database connection failed'}), 500
    try:
        with conn.cursor() as cursor:
            cursor.execute("SELECT password FROM users WHERE id = %s", (session['user_id'],))
            row = cursor.fetchone()
            if not row:
                return jsonify({'success': False, 'message': 'Người dùng không tồn tại'}), 404
            if not check_password_hash(row['password'], current_pw):
                return jsonify({'success': False, 'message': 'Mật khẩu hiện tại không đúng'}), 400
            new_hash = generate_password_hash(new_pw)
            cursor.execute(
                "UPDATE users SET password = %s WHERE id = %s",
                (new_hash, session['user_id'])
            )
            conn.commit()
        return jsonify({'success': True, 'message': 'Đổi mật khẩu thành công'})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500
    finally:
        conn.close()


@app.route('/api/user/settings', methods=['GET'])
@login_required
def get_user_settings_api():
    """Lấy cài đặt TTS & thông báo email của user."""
    settings = get_user_settings(session['user_id'])
    return jsonify({'success': True, 'settings': settings})


@app.route('/api/user/settings', methods=['POST'])
@login_required
def update_user_settings_api():
    """Cập nhật cài đặt TTS & thông báo email."""
    data = request.get_json() or {}
    ok, result = save_user_settings(session['user_id'], data)
    if not ok:
        return jsonify({'success': False, 'message': result}), 500
    return jsonify({'success': True, 'message': 'Cập nhật cài đặt thành công', 'settings': result})


@app.route('/api/user/usage-chart', methods=['GET'])
@login_required
def get_user_usage_chart():
    """Biểu đồ ký tự đã dùng theo ngày (7 ngày gần nhất)."""
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': 'Database connection failed'}), 500
    try:
        now = datetime.now()
        today_start = now.replace(hour=0, minute=0, second=0, microsecond=0)
        chart_data = []
        for i in range(6, -1, -1):
            day_start = today_start - timedelta(days=i)
            day_end = day_start + timedelta(days=1)
            with conn.cursor() as cursor:
                cursor.execute(
                    """
                    SELECT COALESCE(SUM(text_length), 0) AS chars,
                           COUNT(*) AS conversions
                    FROM conversions
                    WHERE user_id = %s AND created_at >= %s AND created_at < %s
                    """,
                    (session['user_id'], day_start, day_end),
                )
                row = cursor.fetchone()
            chart_data.append({
                'date': day_start.strftime('%Y-%m-%d'),
                'label': day_start.strftime('%d/%m'),
                'characters': int(row.get('chars') or 0),
                'conversions': int(row.get('conversions') or 0),
            })
        return jsonify({'success': True, 'chart': chart_data})
    except Exception as e:
        print(f"[ERROR] get_user_usage_chart: {e}")
        return jsonify({'success': False, 'message': str(e)}), 500
    finally:
        conn.close()


def _pdf_register_viet_fonts():
    """Đăng ký font Unicode cho PDF (Arial/Verdana trên Windows)."""
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    font_reg, font_bold = 'Helvetica', 'Helvetica-Bold'
    win_fonts = [
        ('C:/Windows/Fonts/arial.ttf', 'C:/Windows/Fonts/arialbd.ttf'),
        ('C:/Windows/Fonts/verdana.ttf', 'C:/Windows/Fonts/verdanab.ttf'),
    ]
    for reg_path, bold_path in win_fonts:
        if os.path.exists(reg_path):
            try:
                pdfmetrics.registerFont(TTFont('VV_Reg', reg_path))
                font_reg = 'VV_Reg'
                if os.path.exists(bold_path):
                    pdfmetrics.registerFont(TTFont('VV_Bold', bold_path))
                    font_bold = 'VV_Bold'
                break
            except Exception:
                pass
    return font_reg, font_bold


def _prepare_personal_export_sections(user, subscription, payments, conversions, settings):
    """Chuẩn bị dữ liệu sections cho PDF/Word export."""
    exported_at = datetime.now().strftime('%d/%m/%Y %H:%M')
    account = [
        ('Họ và tên', user.get('full_name') or '—'),
        ('Tên đăng nhập', user.get('username') or '—'),
        ('Email', user.get('email') or '—'),
        ('Vai trò', user.get('role') or '—'),
        ('Ngày tạo tài khoản', _format_datetime_vn(user.get('created_at'))),
        ('Liên kết Google', 'Có' if user.get('google_id') else 'Không'),
    ]
    if subscription:
        limit = int(subscription.get('characters_limit') or 0)
        used = int(subscription.get('characters_used') or 0)
        start_d = subscription.get('start_date')
        end_d = subscription.get('end_date')
        period = (
            f'{start_d.strftime("%d/%m/%Y") if start_d else "—"}'
            f' → {end_d.strftime("%d/%m/%Y") if end_d else "—"}'
        )
        subscription_rows = [
            ('Gói hiện tại', subscription.get('package_name') or 'Free / Mặc định'),
            ('Ký tự giới hạn', _fmt_num(limit)),
            ('Đã sử dụng', _fmt_num(used)),
            ('Còn lại', _fmt_num(max(0, limit - used))),
            ('Thời hạn gói', period),
        ]
    else:
        subscription_rows = [('Gói hiện tại', 'Không có gói active')]

    lang = settings.get('default_language') or 'vi'
    tts = [
        ('Giọng TTS cơ bản', settings.get('default_voice_id') or 'Không đặt'),
        (
            'Giọng Emotional TTS',
            settings.get('default_emotional_voice_id') or 'Mặc định (base_voice.wav)',
        ),
        ('Cao độ (pitch)', str(settings.get('default_pitch', 0))),
        ('Tốc độ', f'{settings.get("default_speed", 1)}x'),
        ('Định dạng xuất', (settings.get('default_export_format') or 'wav').upper()),
        ('Bitrate', f'{settings.get("default_export_bitrate", 192)} kbps'),
        ('Ngôn ngữ', 'Tiếng Việt' if lang == 'vi' else 'English'),
    ]
    notify = [
        ('Ký tự sắp hết (<10%)', 'Bật' if settings.get('notify_chars_low') else 'Tắt'),
        ('Thanh toán', 'Bật' if settings.get('notify_payment') else 'Tắt'),
        ('Gói sắp hết hạn', 'Bật' if settings.get('notify_plan_expiry') else 'Tắt'),
        ('Tin tức sản phẩm', 'Bật' if settings.get('notify_marketing') else 'Tắt'),
    ]
    pay_header = ['Gói dịch vụ', 'Mã giao dịch', 'Số tiền', 'Trạng thái', 'Ngày']
    pay_rows = []
    for p in payments[:50]:
        pay_rows.append([
            str(p.get('package_name') or '—'),
            str(p.get('transaction_id') or '—'),
            _fmt_vnd(p.get('amount_vnd')),
            _payment_status_vn(p.get('payment_status')),
            _format_datetime_vn(p.get('created_at')),
        ])
    conv_header = ['Tên / Giọng', 'Ký tự', 'Thời lượng', 'Trạng thái', 'Ngày']
    conv_rows = []
    total_chars = 0
    for cv in conversions[:50]:
        total_chars += int(cv.get('text_length') or 0)
        dur = cv.get('duration_seconds')
        conv_rows.append([
            str(cv.get('display_name') or cv.get('voice_id') or '—'),
            _fmt_num(cv.get('text_length')),
            f'{float(dur):.1f}s' if dur else '—',
            str(cv.get('status') or '—'),
            _format_datetime_vn(cv.get('created_at')),
        ])
    return {
        'exported_at': exported_at,
        'account': account,
        'subscription': subscription_rows,
        'tts': tts,
        'notify': notify,
        'payments_header': pay_header,
        'payments': pay_rows,
        'payments_total': len(payments),
        'conversions_header': conv_header,
        'conversions': conv_rows,
        'conversions_total': len(conversions),
        'total_chars': sum(int(c.get('text_length') or 0) for c in conversions),
    }


def _build_personal_data_pdf(user, subscription, payments, conversions, settings):
    """Báo cáo dữ liệu cá nhân PDF — layout chuyên nghiệp (Platypus)."""
    from io import BytesIO
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle,
    )

    buffer = BytesIO()
    font_reg, font_bold = _pdf_register_viet_fonts()
    data = _prepare_personal_export_sections(user, subscription, payments, conversions, settings)
    page_w, page_h = A4
    margin = 14 * mm

    C_NAVY = colors.HexColor('#051424')
    C_NAVY_MID = colors.HexColor('#122131')
    C_PRIMARY = colors.HexColor('#7078ff')
    C_PRIMARY_LIGHT = colors.HexColor('#eef2ff')
    C_TEXT = colors.HexColor('#1e293b')
    C_MUTED = colors.HexColor('#64748b')
    C_BORDER = colors.HexColor('#e2e8f0')
    C_ROW_ALT = colors.HexColor('#f8fafc')
    C_WHITE = colors.white
    C_ACCENT = colors.HexColor('#2fd9f4')

    styles = {
        'subtitle': ParagraphStyle(
            'vv_sub', fontName=font_reg, fontSize=10, leading=14,
            textColor=C_MUTED, alignment=TA_CENTER,
        ),
        'section': ParagraphStyle(
            'vv_sec', fontName=font_bold, fontSize=12, leading=16,
            textColor=C_NAVY, spaceBefore=14, spaceAfter=8,
        ),
        'cell_label': ParagraphStyle(
            'vv_cl', fontName=font_bold, fontSize=9, leading=12, textColor=C_TEXT,
        ),
        'cell_value': ParagraphStyle(
            'vv_cv', fontName=font_reg, fontSize=9, leading=12, textColor=C_TEXT,
        ),
        'note': ParagraphStyle(
            'vv_note', fontName=font_reg, fontSize=8, leading=11,
            textColor=C_MUTED, spaceBefore=4,
        ),
        'footer': ParagraphStyle(
            'vv_foot', fontName=font_reg, fontSize=8, leading=11,
            textColor=C_MUTED, alignment=TA_CENTER,
        ),
    }

    def on_page(canvas, _doc):
        canvas.saveState()
        canvas.setStrokeColor(C_BORDER)
        canvas.setLineWidth(0.5)
        canvas.line(margin, 22, page_w - margin, 22)
        canvas.setFont(font_reg, 8)
        canvas.setFillColor(C_MUTED)
        canvas.drawString(margin, 10, 'VietVoice AI — Báo cáo dữ liệu cá nhân')
        canvas.drawRightString(page_w - margin, 10, f'Trang {_doc.page}')
        canvas.restoreState()

    def on_first_page(canvas, _doc):
        canvas.saveState()
        canvas.setFillColor(C_NAVY)
        canvas.rect(0, page_h - 32 * mm, page_w, 32 * mm, fill=1, stroke=0)
        canvas.setFillColor(C_PRIMARY)
        canvas.setFont(font_bold, 20)
        canvas.drawString(margin, page_h - 18 * mm, 'VietVoice AI')
        canvas.setFillColor(C_WHITE)
        canvas.setFont(font_bold, 13)
        canvas.drawRightString(page_w - margin, page_h - 16 * mm, 'BÁO CÁO DỮ LIỆU CÁ NHÂN')
        canvas.setFont(font_reg, 9)
        canvas.setFillColor(colors.HexColor('#94a3b8'))
        canvas.drawRightString(
            page_w - margin, page_h - 22 * mm,
            f'Xuất ngày: {data["exported_at"]}',
        )
        on_page(canvas, _doc)
        canvas.restoreState()

    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        leftMargin=margin, rightMargin=margin,
        topMargin=38 * mm, bottomMargin=28 * mm,
        title='VietVoice — Dữ liệu cá nhân',
    )

    def kv_table(rows):
        body = [
            [Paragraph(label, styles['cell_label']), Paragraph(str(val), styles['cell_value'])]
            for label, val in rows
        ]
        tbl = Table(body, colWidths=[42 * mm, 118 * mm])
        tbl.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), C_PRIMARY_LIGHT),
            ('BACKGROUND', (1, 0), (1, -1), C_WHITE),
            ('BOX', (0, 0), (-1, -1), 0.6, C_BORDER),
            ('INNERGRID', (0, 0), (-1, -1), 0.4, C_BORDER),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('LEFTPADDING', (0, 0), (-1, -1), 8),
            ('RIGHTPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 7),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 7),
        ]))
        return tbl

    def data_table(headers, rows, col_widths):
        header_cells = [
            Paragraph(f'<font color="#ffffff">{h}</font>', ParagraphStyle(
                'hdr', fontName=font_bold, fontSize=8, leading=10, textColor=C_WHITE,
            ))
            for h in headers
        ]
        body = []
        for row in rows:
            body.append([
                Paragraph(str(c), ParagraphStyle(
                    'td', fontName=font_reg, fontSize=8, leading=10, textColor=C_TEXT,
                ))
                for c in row
            ])
        tbl = Table([header_cells] + body, colWidths=col_widths, repeatRows=1)
        style_cmds = [
            ('BACKGROUND', (0, 0), (-1, 0), C_NAVY),
            ('TEXTCOLOR', (0, 0), (-1, 0), C_WHITE),
            ('FONTNAME', (0, 0), (-1, 0), font_bold),
            ('FONTSIZE', (0, 0), (-1, 0), 8),
            ('BOX', (0, 0), (-1, -1), 0.6, C_BORDER),
            ('INNERGRID', (0, 0), (-1, -1), 0.4, C_BORDER),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('LEFTPADDING', (0, 0), (-1, -1), 6),
            ('RIGHTPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ]
        for i in range(1, len(body) + 1):
            bg = C_WHITE if i % 2 == 1 else C_ROW_ALT
            style_cmds.append(('BACKGROUND', (0, i), (-1, i), bg))
        tbl.setStyle(TableStyle(style_cmds))
        return tbl

    story = [
        Paragraph(
            'Tài liệu tóm tắt thông tin tài khoản, gói dịch vụ, giao dịch và lịch sử sử dụng '
            '(không chứa file âm thanh).',
            styles['subtitle'],
        ),
        Spacer(1, 6 * mm),
        Paragraph('1. Thông tin tài khoản', styles['section']),
        kv_table(data['account']),
        Paragraph('2. Gói dịch vụ & sử dụng', styles['section']),
        kv_table(data['subscription']),
        Paragraph('3. Cài đặt mặc định TTS', styles['section']),
        kv_table(data['tts']),
        Paragraph('4. Thông báo email', styles['section']),
        kv_table(data['notify']),
        Paragraph(
            f'5. Lịch sử giao dịch ({data["payments_total"]} bản ghi)',
            styles['section'],
        ),
    ]
    if data['payments']:
        story.append(data_table(
            data['payments_header'], data['payments'],
            [32 * mm, 38 * mm, 28 * mm, 24 * mm, 32 * mm],
        ))
        if data['payments_total'] > len(data['payments']):
            story.append(Paragraph(
                f'Ghi chú: hiển thị {len(data["payments"])}/{data["payments_total"]} giao dịch gần nhất.',
                styles['note'],
            ))
    else:
        story.append(Paragraph('Chưa có giao dịch.', styles['note']))

    story.append(Paragraph(
        f'6. Lịch sử chuyển đổi ({data["conversions_total"]} bản ghi)',
        styles['section'],
    ))
    if data['conversions']:
        story.append(kv_table([('Tổng ký tự (tất cả)', _fmt_num(data['total_chars']))]))
        story.append(Spacer(1, 3 * mm))
        story.append(data_table(
            data['conversions_header'], data['conversions'],
            [38 * mm, 22 * mm, 22 * mm, 24 * mm, 32 * mm],
        ))
        if data['conversions_total'] > len(data['conversions']):
            story.append(Paragraph(
                f'Ghi chú: hiển thị {len(data["conversions"])}/{data["conversions_total"]} bản ghi gần nhất.',
                styles['note'],
            ))
    else:
        story.append(Paragraph('Chưa có lịch sử chuyển đổi.', styles['note']))

    story.append(Spacer(1, 8 * mm))
    story.append(Paragraph(
        'Để xóa tài khoản, xem Chính sách xóa dữ liệu trên website VietVoice.',
        styles['footer'],
    ))

    doc.build(story, onFirstPage=on_first_page, onLaterPages=on_page)
    buffer.seek(0)
    return buffer


def _docx_set_cell_shading(cell, fill_hex):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill_hex)
    shd.set(qn('w:val'), 'clear')
    tc_pr.append(shd)


def _docx_write_cell(cell, text, bold=False, size=10, color_rgb=(0x1e, 0x29, 0x3b), center=False):
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    cell.text = ''
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(*color_rgb)


def _docx_kv_table(doc, rows):
    from docx.shared import Pt, RGBColor, Cm
    table = doc.add_table(rows=len(rows), cols=2)
    table.autofit = False
    for row_idx, (label, value) in enumerate(rows):
        c0, c1 = table.rows[row_idx].cells
        c0.width = Cm(4.5)
        c1.width = Cm(12)
        _docx_set_cell_shading(c0, 'EEF2FF')
        _docx_write_cell(c0, label, bold=True, size=10, color_rgb=(0x37, 0x41, 0x51))
        _docx_write_cell(c1, value, size=10)
    return table


def _docx_data_table(doc, headers, rows):
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.autofit = False
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        _docx_set_cell_shading(hdr[i], '051424')
        _docx_write_cell(hdr[i], h, bold=True, size=9, color_rgb=(0xff, 0xff, 0xff), center=True)
    for ri, row_data in enumerate(rows):
        cells = table.rows[ri + 1].cells
        bg = 'FFFFFF' if ri % 2 == 0 else 'F8FAFC'
        for ci, val in enumerate(row_data):
            _docx_set_cell_shading(cells[ci], bg)
            _docx_write_cell(cells[ci], val, size=9)
    return table


def _build_personal_data_docx(user, subscription, payments, conversions, settings):
    """Báo cáo dữ liệu cá nhân Word — layout chuyên nghiệp."""
    from io import BytesIO
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()
    data = _prepare_personal_export_sections(user, subscription, payments, conversions, settings)

    for section in doc.sections:
        section.top_margin = Cm(1.2)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    normal = doc.styles['Normal']
    normal.font.name = 'Arial'
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    normal.font.size = Pt(10)

    # Banner header
    banner = doc.add_table(rows=1, cols=1)
    banner_cell = banner.rows[0].cells[0]
    _docx_set_cell_shading(banner_cell, '051424')
    p_brand = banner_cell.paragraphs[0]
    p_brand.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_brand = p_brand.add_run('VietVoice AI')
    r_brand.bold = True
    r_brand.font.size = Pt(22)
    r_brand.font.color.rgb = RGBColor(0xa5, 0xb4, 0xfc)
    p_title = banner_cell.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_title = p_title.add_run('BÁO CÁO DỮ LIỆU CÁ NHÂN')
    r_title.bold = True
    r_title.font.size = Pt(14)
    r_title.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
    p_date = banner_cell.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_date = p_date.add_run(f'Xuất ngày: {data["exported_at"]}')
    r_date.font.size = Pt(9)
    r_date.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

    doc.add_paragraph()
    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_intro = intro.add_run(
        'Tài liệu tóm tắt thông tin tài khoản, gói dịch vụ, giao dịch và lịch sử sử dụng '
        '(không chứa file âm thanh).'
    )
    r_intro.font.size = Pt(10)
    r_intro.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

    def section_title(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x05, 0x14, 0x24)
        # Accent line under title
        bar = doc.add_table(rows=1, cols=1)
        bar_cell = bar.rows[0].cells[0]
        _docx_set_cell_shading(bar_cell, '7078FF')
        bar_cell.height = Cm(0.08)

    section_title('1. Thông tin tài khoản')
    _docx_kv_table(doc, data['account'])

    section_title('2. Gói dịch vụ & sử dụng')
    _docx_kv_table(doc, data['subscription'])

    section_title('3. Cài đặt mặc định TTS')
    _docx_kv_table(doc, data['tts'])

    section_title('4. Thông báo email')
    _docx_kv_table(doc, data['notify'])

    section_title(f'5. Lịch sử giao dịch ({data["payments_total"]} bản ghi)')
    if data['payments']:
        _docx_data_table(doc, data['payments_header'], data['payments'])
        if data['payments_total'] > len(data['payments']):
            note = doc.add_paragraph()
            r = note.add_run(
                f'Ghi chú: hiển thị {len(data["payments"])}/{data["payments_total"]} giao dịch gần nhất.'
            )
            r.font.size = Pt(8)
            r.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
    else:
        doc.add_paragraph('Chưa có giao dịch.')

    section_title(f'6. Lịch sử chuyển đổi ({data["conversions_total"]} bản ghi)')
    if data['conversions']:
        _docx_kv_table(doc, [('Tổng ký tự (tất cả)', _fmt_num(data['total_chars']))])
        doc.add_paragraph()
        _docx_data_table(doc, data['conversions_header'], data['conversions'])
        if data['conversions_total'] > len(data['conversions']):
            note = doc.add_paragraph()
            r = note.add_run(
                f'Ghi chú: hiển thị {len(data["conversions"])}/{data["conversions_total"]} bản ghi gần nhất.'
            )
            r.font.size = Pt(8)
            r.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
    else:
        doc.add_paragraph('Chưa có lịch sử chuyển đổi.')

    doc.add_paragraph()
    foot_bar = doc.add_table(rows=1, cols=1)
    foot_cell = foot_bar.rows[0].cells[0]
    _docx_set_cell_shading(foot_cell, 'F1F5F9')
    fp = foot_cell.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run(
        'VietVoice AI — Báo cáo dữ liệu cá nhân — vietvoice.ai\n'
        'Để xóa tài khoản, xem Chính sách xóa dữ liệu trên website.'
    )
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out


@app.route('/api/user/export-data', methods=['GET'])
@login_required
def export_user_data():
    """Xuất báo cáo dữ liệu cá nhân (PDF hoặc Word)."""
    user_id = session['user_id']
    fmt = (request.args.get('format') or 'pdf').strip().lower()
    if fmt not in ('pdf', 'docx'):
        fmt = 'pdf'

    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': 'Database connection failed'}), 500

    try:
        with conn.cursor() as cursor:
            cursor.execute(
                """
                SELECT id, username, email, full_name, role, created_at, google_id
                FROM users WHERE id = %s
                """,
                (user_id,),
            )
            user = cursor.fetchone()
            if not user:
                return jsonify({'success': False, 'message': 'Người dùng không tồn tại'}), 404

            cursor.execute(
                """
                SELECT us.characters_limit, us.characters_used, us.start_date, us.end_date,
                       sp.package_name
                FROM user_subscriptions us
                LEFT JOIN subscription_packages sp ON us.package_id = sp.id
                WHERE us.user_id = %s AND us.is_active = 1 AND us.end_date >= CURDATE()
                ORDER BY us.end_date DESC
                LIMIT 1
                """,
                (user_id,),
            )
            subscription = cursor.fetchone()

            cursor.execute(
                """
                SELECT p.id, p.transaction_id, p.amount_vnd, p.payment_method,
                       p.payment_status, p.created_at, p.completed_at,
                       sp.package_name, sp.characters_limit
                FROM payments p
                LEFT JOIN subscription_packages sp ON p.package_id = sp.id
                WHERE p.user_id = %s
                ORDER BY p.created_at DESC
                LIMIT 500
                """,
                (user_id,),
            )
            payments = cursor.fetchall()

            cursor.execute(
                """
                SELECT id, voice_id, text_length, duration_seconds, status,
                       created_at, completed_at, display_name
                FROM conversions
                WHERE user_id = %s
                ORDER BY created_at DESC
                LIMIT 1000
                """,
                (user_id,),
            )
            conversions = cursor.fetchall()

        settings = get_user_settings(user_id)
        date_stamp = datetime.now().strftime('%Y%m%d')

        if fmt == 'docx':
            try:
                doc_buffer = _build_personal_data_docx(
                    user, subscription, payments, conversions, settings,
                )
            except ImportError:
                return jsonify({
                    'success': False,
                    'message': 'Thư viện python-docx chưa được cài. Chạy: pip install python-docx',
                }), 500
            doc_bytes = doc_buffer.getvalue()
            filename = f'vietvoice_du_lieu_ca_nhan_{user_id}_{date_stamp}.docx'
            return Response(
                doc_bytes,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                headers={
                    'Content-Disposition': f'attachment; filename="{filename}"',
                    'Content-Length': str(len(doc_bytes)),
                    'Cache-Control': 'no-store',
                },
            )

        try:
            pdf_buffer = _build_personal_data_pdf(
                user, subscription, payments, conversions, settings,
            )
        except ImportError:
            return jsonify({
                'success': False,
                'message': 'Thư viện reportlab chưa được cài. Chạy: pip install reportlab',
            }), 500

        pdf_bytes = pdf_buffer.getvalue()
        filename = f'vietvoice_du_lieu_ca_nhan_{user_id}_{date_stamp}.pdf'
        return Response(
            pdf_bytes,
            mimetype='application/pdf',
            headers={
                'Content-Disposition': f'attachment; filename="{filename}"',
                'Content-Length': str(len(pdf_bytes)),
                'Cache-Control': 'no-store',
            },
        )
    except Exception as e:
        print(f"[ERROR] export_user_data: {e}")
        return jsonify({'success': False, 'message': str(e)}), 500
    finally:
        conn.close()


@app.route('/contact')
def contact():
    """Trang liên hệ"""
    return render_template(
        'contact.html',
        support_email=get_support_email(),
        contact_email=get_contact_email(),
    )

@app.route('/privacy')
def privacy():
    """Chính sách quyền riêng tư"""
    return render_template(
        'privacy.html',
        legal=get_legal_for_display('privacy'),
        legal_en=get_legal_for_display('privacy', 'en'),
    )

@app.route('/terms')
def terms():
    """Điều khoản sử dụng"""
    return render_template(
        'terms.html',
        legal=get_legal_for_display('terms'),
        legal_en=get_legal_for_display('terms', 'en'),
    )

@app.route('/data-deletion')
def data_deletion():
    """Chính sách xóa dữ liệu"""
    return render_template(
        'data_deletion.html',
        legal=get_legal_for_display('data_deletion'),
        legal_en=get_legal_for_display('data_deletion', 'en'),
    )

@app.route('/payment-terms')
def payment_terms():
    """Điều khoản thanh toán"""
    return render_template(
        'payment_terms.html',
        legal=get_legal_for_display('payment'),
        legal_en=get_legal_for_display('payment', 'en'),
    )

@app.route('/support')
def support():
    """Trang hỗ trợ và FAQ"""
    return render_template(
        'support.html',
        support=get_support_for_display('vi'),
        support_en=get_support_for_display('en'),
        support_email=get_support_email(),
        contact_email=get_contact_email(),
    )

@app.route('/user_guide')
def user_guide_alias():
    return redirect('/user-guide')

@app.route('/user-guide')
def user_guide():
    """Hướng dẫn sử dụng"""
    legal = get_legal_for_display('user_guide')
    legal_en = _get_legal_en_for_template('user_guide', legal)
    return render_template(
        'user_guide.html',
        legal=legal,
        legal_en=legal_en,
    )

@app.route('/installation_guide')
def installation_guide_alias():
    return redirect('/installation-guide')

@app.route('/installation-guide')
def installation_guide():
    """Hướng dẫn cài đặt hệ thống"""
    legal = get_legal_for_display('installation_guide')
    legal_en = _get_legal_en_for_template('installation_guide', legal)
    return render_template(
        'installation_guide.html',
        legal=legal,
        legal_en=legal_en,
    )

@app.route('/api/contact', methods=['POST'])
def submit_contact():
    """Submit contact form"""
    try:
        data = request.get_json()
        name = data.get('name', '').strip()
        email = data.get('email', '').strip()
        subject = data.get('subject', '').strip()
        message = data.get('message', '').strip()
        
        # Validate
        if not all([name, email, subject, message]):
            return jsonify({
                'success': False,
                'message': 'Vui lòng điền đầy đủ thông tin'
            }), 400
        
        # Validate email format
        import re
        email_pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
        if not re.match(email_pattern, email):
            return jsonify({
                'success': False,
                'message': 'Email không hợp lệ'
            }), 400
        
        # Log contact (có thể lưu vào database nếu muốn)
        print(f"[CONTACT] From: {name} ({email})")
        print(f"[CONTACT] Subject: {subject}")
        print(f"[CONTACT] Message: {message}")
        
        return jsonify({
            'success': True,
            'message': 'Cảm ơn bạn đã liên hệ! Chúng tôi sẽ phản hồi trong vòng 24 giờ.'
        })
        
    except Exception as e:
        print(f"Error submitting contact: {e}")
        return jsonify({
            'success': False,
            'message': 'Đã xảy ra lỗi. Vui lòng thử lại sau.'
        }), 500

@app.route('/api/voice-conversion', methods=['POST'])
def voice_conversion():
    """
    Voice Conversion API
    Điều chỉnh giọng nói sử dụng RVC
    """
    try:
        # Check if user is logged in
        if not is_logged_in():
            return jsonify({
                'success': False,
                'message': 'Vui lòng đăng nhập để sử dụng tính năng này'
            }), 401
        
        # Check if voice conversion is available (RVC or librosa fallback)
        if RVC_AVAILABLE:
            rvc_processor = get_rvc_processor()
            if not rvc_processor.is_available():
                return jsonify({
                    'success': False,
                    'message': 'Tính năng điều chỉnh giọng tạm thời không khả dụng'
                }), 503
        else:
            return jsonify({
                'success': False,
                'message': 'Tính năng điều chỉnh giọng tạm thời không khả dụng'
            }), 503
        
        # Get request data
        data = request.get_json()
        audio_filename = data.get('audio_filename', '').strip()
        pitch = int(data.get('pitch', 0))
        index_rate = float(data.get('index_rate', 0.75))
        protect = float(data.get('protect', 0.33))
        
        # Validate
        if not audio_filename:
            return jsonify({
                'success': False,
                'message': 'Không tìm thấy file audio'
            }), 400
        
        # Validate pitch range
        if pitch < -12 or pitch > 12:
            return jsonify({
                'success': False,
                'message': 'Pitch phải trong khoảng -12 đến +12'
            }), 400
        
        # Validate index_rate
        if index_rate < 0 or index_rate > 1:
            return jsonify({
                'success': False,
                'message': 'Index rate phải trong khoảng 0 đến 1'
            }), 400
        
        # Validate protect
        if protect < 0 or protect > 0.5:
            return jsonify({
                'success': False,
                'message': 'Protect phải trong khoảng 0 đến 0.5'
            }), 400
        
        # Get input audio path
        input_path = os.path.join(AUDIO_OUTPUT_DIR, audio_filename)
        
        if not os.path.exists(input_path):
            return jsonify({
                'success': False,
                'message': 'File audio không tồn tại'
            }), 404
        
        # Get RVC processor
        rvc_processor = get_rvc_processor()
        
        # Process audio
        print(f"[VOICE CONVERSION] Processing: {audio_filename}")
        print(f"[VOICE CONVERSION] Pitch: {pitch}, Index Rate: {index_rate}, Protect: {protect}")
        
        success, output_path, message = rvc_processor.adjust_voice(
            input_audio_path=input_path,
            model_path=None,  # Use simple pitch shift for now
            f0_up_key=pitch,
            index_rate=index_rate,
            protect=protect
        )
        
        if not success:
            return jsonify({
                'success': False,
                'message': message
            }), 500
        
        # Get output filename
        output_filename = os.path.basename(output_path)
        
        # Update original conversion record (don't create new one to avoid double counting)
        user_id = session.get('user_id')
        try:
            connection = get_db_connection()
            if connection:
                with connection.cursor() as cursor:
                    # Find original conversion to update
                    # Note: audio_file_path in DB is full path, audio_filename is just filename
                    cursor.execute("""
                        SELECT id, text_input, voice_id, voice_name 
                        FROM conversions 
                        WHERE audio_file_path LIKE %s AND user_id = %s
                        ORDER BY created_at DESC LIMIT 1
                    """, (f'%{audio_filename}', user_id))
                    
                    original = cursor.fetchone()
                    
                    if original:
                        # UPDATE original record with adjusted audio file
                        # This keeps statistics correct (1 TTS = 1 conversion)
                        original_id = original['id']
                        original_text = original['text_input']
                        
                        # Get adjusted audio file info
                        adjusted_path = os.path.join(AUDIO_OUTPUT_DIR, output_filename)
                        file_size = os.path.getsize(adjusted_path) if os.path.exists(adjusted_path) else 0
                        
                        # Update with adjusted file, append adjustment info to text
                        cursor.execute("""
                            UPDATE conversions 
                            SET audio_file_path = %s,
                                audio_file_size = %s,
                                text_input = %s,
                                completed_at = NOW()
                            WHERE id = %s
                        """, (
                            adjusted_path,
                            file_size,
                            f"[Đã điều chỉnh: Pitch {pitch:+d}, Index {index_rate:.2f}] {original_text}",
                            original_id
                        ))
                        
                        connection.commit()
                        print(f"[VOICE CONVERSION] Updated conversion ID {original_id} with adjusted file: {output_filename}")
                    else:
                        # If original not found (shouldn't happen), log warning
                        print(f"[WARNING] Original conversion not found for {audio_filename}, skipping database update")
                    
        except Exception as db_error:
            print(f"[WARNING] Could not update database: {db_error}")
            import traceback
            traceback.print_exc()
        finally:
            if connection:
                connection.close()
        
        print(f"[VOICE CONVERSION] Success: {output_filename}")
        
        return jsonify({
            'success': True,
            'message': message,
            'audio_filename': output_filename,
            'audio_url': url_for('get_audio', filename=output_filename)
        })
        
    except ValueError as e:
        print(f"[ERROR] Invalid parameters: {e}")
        return jsonify({
            'success': False,
            'message': f'Tham số không hợp lệ: {str(e)}'
        }), 400
    except Exception as e:
        print(f"[ERROR] Voice conversion error: {e}")
        traceback.print_exc()
        return jsonify({
            'success': False,
            'message': 'Đã xảy ra lỗi khi xử lý. Vui lòng thử lại sau.'
        }), 500

@app.route('/api/voice-conversion/check')
def check_voice_conversion():
    """Check if voice conversion feature is available"""
    try:
        if not RVC_AVAILABLE:
            return jsonify({
                'available': False,
                'message': 'Voice conversion module not loaded'
            })
        
        rvc_processor = get_rvc_processor()
        is_available = rvc_processor.is_available()
        
        return jsonify({
            'available': is_available,
            'message': 'Voice conversion is ready (using librosa fallback)' if is_available else 'Voice conversion not available'
        })
    except Exception as e:
        return jsonify({
            'available': False,
            'message': f'Error checking availability: {str(e)}'
        })

@app.route('/audio-library')
def audio_library():
    """Trang thư viện audio"""
    if not is_logged_in():
        return redirect(url_for('login'))
    return render_template('audio_library.html', app_base_url=APP_BASE_URL.rstrip('/') if APP_BASE_URL else '')

@app.route('/api/audio-library')
def get_audio_library():
    """Lấy danh sách audio với pagination và filter"""
    if not is_logged_in():
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401
    
    page = int(request.args.get('page', 1))
    per_page = int(request.args.get('per_page', 12))
    search = request.args.get('search', '')
    voice_filter = request.args.get('voice', '')
    date_from = request.args.get('date_from', '')
    date_to = request.args.get('date_to', '')
    sort_by = request.args.get('sort_by', 'newest')
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': 'Database connection failed'}), 500
    
    try:
        with conn.cursor() as cursor:
            # Build query
            query = """
                SELECT * FROM conversions 
                WHERE user_id = %s AND status = 'completed' AND audio_file_path IS NOT NULL
            """
            params = [session['user_id']]
            
            # Search filter (tìm theo nội dung hoặc tên đặt)
            if search:
                query += " AND (text_input LIKE %s OR display_name LIKE %s)"
                params.extend([f'%{search}%', f'%{search}%'])
            
            # Voice filter
            if voice_filter:
                query += " AND voice_id = %s"
                params.append(voice_filter)
            
            # Date filters
            if date_from:
                query += " AND DATE(created_at) >= %s"
                params.append(date_from)
            if date_to:
                query += " AND DATE(created_at) <= %s"
                params.append(date_to)
            
            # Sorting
            if sort_by == 'newest':
                query += " ORDER BY created_at DESC"
            elif sort_by == 'oldest':
                query += " ORDER BY created_at ASC"
            elif sort_by == 'duration':
                query += " ORDER BY duration_seconds DESC"
            elif sort_by == 'size':
                query += " ORDER BY audio_file_size DESC"
            
            # Get total count
            count_query = query.replace("SELECT *", "SELECT COUNT(*) as total")
            cursor.execute(count_query, params)
            total = cursor.fetchone()['total']
            
            # Pagination
            query += " LIMIT %s OFFSET %s"
            params.extend([per_page, (page - 1) * per_page])
            
            cursor.execute(query, params)
            audios = cursor.fetchall()
            
            # Convert datetime to string + share URL công khai
            for audio in audios:
                if audio['created_at']:
                    audio['created_at'] = audio['created_at'].isoformat()
                if audio['completed_at']:
                    audio['completed_at'] = audio['completed_at'].isoformat()
                if audio.get('is_public') and audio.get('share_token'):
                    audio['share_url'] = _share_audio_url(audio['share_token'])
                else:
                    audio['share_url'] = None
            
            return jsonify({
                'success': True,
                'audios': audios,
                'total': total,
                'page': page,
                'per_page': per_page,
                'total_pages': (total + per_page - 1) // per_page
            })
    except Exception as e:
        print(f"[ERROR] Get audio library error: {e}")
        return jsonify({'success': False, 'message': f'Loi: {str(e)}'}), 500
    finally:
        conn.close()

@app.route('/api/audio-library/<int:audio_id>', methods=['DELETE'])
def delete_audio(audio_id):
    """Xoa audio file va record trong DB"""
    if not is_logged_in():
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': 'Database connection failed'}), 500
    
    try:
        with conn.cursor() as cursor:
            # Get audio info and verify ownership
            cursor.execute("""
                SELECT audio_file_path FROM conversions 
                WHERE id = %s AND user_id = %s
            """, (audio_id, session['user_id']))
            audio = cursor.fetchone()
            
            if not audio:
                return jsonify({'success': False, 'message': 'Audio khong tim thay hoac ban khong co quyen'}), 404
            
            # Delete physical file
            if audio['audio_file_path']:
                file_path = Path(audio['audio_file_path'])
                if file_path.exists():
                    file_path.unlink()
                    print(f"[DELETE] Deleted audio file: {file_path}")
            
            # Delete DB record
            cursor.execute("DELETE FROM conversions WHERE id = %s", (audio_id,))
            conn.commit()
            
            return jsonify({
                'success': True,
                'message': 'Da xoa audio thanh cong'
            })
    except Exception as e:
        conn.rollback()
        print(f"[ERROR] Delete audio error: {e}")
        return jsonify({'success': False, 'message': f'Loi: {str(e)}'}), 500
    finally:
        conn.close()

@app.route('/api/audio-library/<int:audio_id>/rename', methods=['PATCH'])
def rename_audio(audio_id):
    """Đặt tên hiển thị cho audio"""
    if not is_logged_in():
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401

    data = request.get_json() or {}
    display_name = data.get('display_name', '').strip()[:200]

    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': 'Database error'}), 500
    try:
        with conn.cursor() as cursor:
            cursor.execute("""
                UPDATE conversions SET display_name = %s
                WHERE id = %s AND user_id = %s
            """, (display_name or None, audio_id, session['user_id']))
            if cursor.rowcount == 0:
                return jsonify({'success': False, 'message': 'Không tìm thấy audio'}), 404
            conn.commit()
            return jsonify({'success': True, 'display_name': display_name})
    except Exception as e:
        conn.rollback()
        return jsonify({'success': False, 'message': str(e)}), 500
    finally:
        conn.close()


@app.route('/api/audio-library/<int:audio_id>/share', methods=['POST'])
def toggle_share_audio(audio_id):
    """Bật/tắt chia sẻ công khai cho audio"""
    if not is_logged_in():
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401

    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': 'Database error'}), 500
    try:
        with conn.cursor() as cursor:
            cursor.execute("""
                SELECT is_public, share_token FROM conversions
                WHERE id = %s AND user_id = %s
            """, (audio_id, session['user_id']))
            audio = cursor.fetchone()
            if not audio:
                return jsonify({'success': False, 'message': 'Không tìm thấy audio'}), 404

            if audio['is_public']:
                # Tắt chia sẻ
                cursor.execute("UPDATE conversions SET is_public = 0 WHERE id = %s", (audio_id,))
                conn.commit()
                return jsonify({'success': True, 'is_public': False, 'share_token': None, 'share_url': None})
            else:
                # Bật chia sẻ — tạo token nếu chưa có
                token = audio['share_token'] or secrets.token_urlsafe(32)
                cursor.execute(
                    "UPDATE conversions SET is_public = 1, share_token = %s WHERE id = %s",
                    (token, audio_id)
                )
                conn.commit()
                return jsonify({
                    'success': True,
                    'is_public': True,
                    'share_token': token,
                    'share_url': _share_audio_url(token),
                })
    except Exception as e:
        conn.rollback()
        return jsonify({'success': False, 'message': str(e)}), 500
    finally:
        conn.close()


@app.route('/audio/share/<token>')
def share_audio_page(token):
    """Trang xem audio công khai — không cần đăng nhập"""
    conn = get_db_connection()
    if not conn:
        return "Lỗi kết nối server", 500
    try:
        with conn.cursor() as cursor:
            cursor.execute("""
                SELECT c.id, c.text_input, c.display_name, c.voice_name, c.voice_id,
                       c.duration_seconds, c.audio_file_size, c.created_at, c.share_token,
                       u.full_name, u.username
                FROM conversions c
                JOIN users u ON c.user_id = u.id
                WHERE c.share_token = %s AND c.is_public = 1
            """, (token,))
            audio = cursor.fetchone()
            if not audio:
                return render_template('share_audio.html', audio=None, token=token), 404
            if audio['created_at']:
                audio['created_at'] = audio['created_at'].isoformat()
            return render_template('share_audio.html', audio=audio, token=token)
    except Exception as e:
        print(f"[ERROR] share_audio_page: {e}")
        return "Lỗi server", 500
    finally:
        conn.close()


@app.route('/api/audio/share/<token>')
def get_shared_audio_file(token):
    """Phục vụ file âm thanh công khai — không cần đăng nhập"""
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False}), 500
    try:
        with conn.cursor() as cursor:
            cursor.execute("""
                SELECT audio_file_path FROM conversions
                WHERE share_token = %s AND is_public = 1
            """, (token,))
            row = cursor.fetchone()
            if not row or not row['audio_file_path']:
                return jsonify({'success': False, 'message': 'File not found'}), 404
            file_path = Path(row['audio_file_path'])
            if not file_path.exists():
                return jsonify({'success': False, 'message': 'File not found on disk'}), 404
            resp = send_file(file_path, mimetype='audio/wav', as_attachment=False)
            resp.headers['Accept-Ranges'] = 'bytes'
            resp.headers['Cache-Control'] = 'public, max-age=86400'
            return resp
    except Exception as e:
        print(f"[ERROR] get_shared_audio_file: {e}")
        return jsonify({'success': False}), 500
    finally:
        conn.close()


@app.route('/api/subscription/status')
def get_subscription_status():
    """Lấy trạng thái subscription của user"""
    if not is_logged_in():
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401
    
    limit_info = get_user_characters_limit(session['user_id'])
    if not limit_info:
        return jsonify({'success': False, 'message': 'Không thể lấy thông tin subscription'}), 500
    
    return jsonify({
        'success': True,
        'subscription': limit_info
    })

def fetch_active_packages():
    """Đọc gói cước active từ DB (dùng cho API và render trang pricing)."""
    conn = get_db_connection()
    if not conn:
        return []
    try:
        with conn.cursor() as cursor:
            cursor.execute("""
                SELECT id, package_name, characters_limit, price_vnd, duration_days
                FROM subscription_packages
                WHERE is_active = 1
                ORDER BY characters_limit ASC
            """)
            rows = cursor.fetchall()
            return [
                {
                    'id': p['id'],
                    'name': p['package_name'],
                    'characters': int(p['characters_limit']),
                    'price': int(p['price_vnd']),
                    'duration_days': int(p['duration_days']),
                }
                for p in rows
            ]
    except Exception as e:
        print(f"[ERROR] fetch_active_packages: {e}")
        return []
    finally:
        conn.close()

@app.route('/api/packages')
def get_packages():
    """Lấy danh sách các gói thanh toán"""
    packages = fetch_active_packages()
    resp = jsonify({'success': True, 'packages': packages})
    resp.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate'
    resp.headers['Pragma'] = 'no-cache'
    return resp

@app.route('/api/payment/create', methods=['POST'])
def create_payment():
    """Tạo payment request với SePay.vn"""
    if not is_logged_in():
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401
    
    data = request.get_json()
    package_id = data.get('package_id')
    payment_method = 'bank_qr'  # Sử dụng bank_qr cho cả SePay và bank transfer
    
    if not package_id:
        return jsonify({'success': False, 'message': 'Thiếu thông tin thanh toán'}), 400
    
    # Debug: Log user info
    current_user_id = session.get('user_id')
    current_username = session.get('username', 'Unknown')
    print(f"[DEBUG] Create payment - User ID: {current_user_id}, Username: {current_username}, Package ID: {package_id}")
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': 'Database connection failed'}), 500
    
    try:
        with conn.cursor() as cursor:
            # Kiểm tra user tồn tại trong DB (tránh lỗi FK khi DB mới tạo lại / session cũ)
            cursor.execute("SELECT id, username FROM users WHERE id = %s", (current_user_id,))
            user_check = cursor.fetchone()
            if not user_check:
                print(f"[ERROR] User {current_user_id} not found in database")
                return jsonify({
                    'success': False,
                    'message': 'Phiên đăng nhập không còn hợp lệ (tài khoản không tồn tại trong hệ thống). Vui lòng đăng xuất và đăng nhập lại.'
                }), 401
            
            print(f"[DEBUG] User verified: {user_check['username']} (ID: {user_check['id']})")
            
            # Lấy thông tin package
            cursor.execute("""
                SELECT id, package_name, characters_limit, price_vnd, duration_days
                FROM subscription_packages
                WHERE id = %s AND is_active = 1
            """, (package_id,))
            package = cursor.fetchone()
            
            if not package:
                return jsonify({'success': False, 'message': 'Gói thanh toán không tồn tại'}), 404
            
            print(f"[DEBUG] Package: {package['package_name']} - {package['characters_limit']:,} chars - {package['price_vnd']:,}đ")
            
            # Tạo payment record  
            transaction_id = f"TTS{uuid.uuid4().hex[:16].upper()}"
            cursor.execute("""
                INSERT INTO payments (user_id, package_id, amount_vnd, payment_method, payment_status, transaction_id)
                VALUES (%s, %s, %s, %s, 'pending', %s)
            """, (current_user_id, package_id, package['price_vnd'], payment_method, transaction_id))
            payment_id = cursor.lastrowid
            conn.commit()
            
            print(f"[DEBUG] Payment created: ID {payment_id}, Transaction: {transaction_id}")
            
            # Tạo thanh toán SePay
            sepay_result = create_sepay_payment(
                package['price_vnd'], 
                transaction_id, 
                f"Thanh toán {package['package_name']}",
                current_username
            )
            
            if sepay_result['success']:
                return jsonify({
                    'success': True,
                    'payment_id': payment_id,
                    'transaction_id': transaction_id,
                    'qr_code': sepay_result['qr_code'],
                    'bank_info': sepay_result['bank_info'],
                    'sepay_info': sepay_result.get('sepay_info'),
                    'payment_type': 'sepay',
                    'package_info': {
                        'name': package['package_name'],
                        'characters': package['characters_limit'],
                        'price': package['price_vnd'],
                        'duration': package['duration_days']
                    },
                    'user_info': {
                        'username': current_username,
                        'user_id': current_user_id
                    }
                })
            else:
                # Fallback to bank transfer if SePay fails
                print(f"[WARNING] SePay failed, falling back to bank transfer: {sepay_result.get('error')}")
                
                # Create bank QR as fallback
                qr_data = create_bank_transfer_qr(
                    package['price_vnd'], 
                    transaction_id, 
                    package['package_name'],
                    current_username
                )
                return jsonify({
                    'success': True,
                    'payment_id': payment_id,
                    'transaction_id': transaction_id,
                    'qr_code': qr_data['qr_image'],
                    'bank_info': qr_data['bank_info'],
                    'payment_type': 'bank_qr',
                    'fallback': True,
                    'message': 'SePay không khả dụng, sử dụng chuyển khoản ngân hàng',
                    'package_info': {
                        'name': package['package_name'],
                        'characters': package['characters_limit'],
                        'price': package['price_vnd'],
                        'duration': package['duration_days']
                    },
                    'user_info': {
                        'username': current_username,
                        'user_id': current_user_id
                    }
                })
                
    except Exception as e:
        conn.rollback()
        print(f"[ERROR] Create payment error: {e}")
        return jsonify({'success': False, 'message': f'Lỗi: {str(e)}'}), 500
    finally:
        conn.close()

def create_bank_transfer_qr(amount, transaction_id, package_name, user_name):
    """Tạo QR code cho chuyển khoản ngân hàng sử dụng API VietQR chính thức"""
    import unicodedata
    import requests
    
    def remove_accents(text):
        """Chuyển đổi tiếng Việt có dấu sang không dấu"""
        if not text:
            return ''
        # Chuyển sang NFD (Normalization Form Decomposed)
        nfd = unicodedata.normalize('NFD', text)
        # Loại bỏ các ký tự dấu
        return ''.join(c for c in nfd if unicodedata.category(c) != 'Mn')
    
    def to_uppercase_no_accent(text, max_length=None):
        """Chuyển sang chữ hoa, không dấu, giới hạn độ dài"""
        result = remove_accents(str(text).upper().strip())
        if max_length and len(result) > max_length:
            result = result[:max_length]
        return result
    
    # Tạo nội dung chuyển khoản
    content = f"{transaction_id}"
    content = remove_accents(content)
    if len(content) > 25:
        content = content[:25]
    
    # Làm sạch số tài khoản
    account_number = str(BANK_ACCOUNT_NUMBER).strip().replace(' ', '').replace('-', '')
    
    # BIN code của TPBank
    tpbank_bin = "970423"
    
    # Thử sử dụng API VietQR (https://api.vietqr.io)
    try:
        vietqr_url = "https://img.vietqr.io/image/{}-{}-compact2.jpg?amount={}&addInfo={}&accountName={}"
        qr_url = vietqr_url.format(
            tpbank_bin,
            account_number,
            amount,
            content,
            to_uppercase_no_accent(BANK_ACCOUNT_NAME, 50)
        )
        
        # Download QR image
        response = requests.get(qr_url, timeout=10)
        if response.status_code == 200:
            # Convert to base64
            img_base64 = base64.b64encode(response.content).decode()
            return {
                'qr_image': f"data:image/jpeg;base64,{img_base64}",
                'bank_info': {
                    'bank_name': BANK_NAME,
                    'account_number': account_number,
                    'account_name': BANK_ACCOUNT_NAME,
                    'branch': BANK_BRANCH,
                    'amount': amount,
                    'content': content,
                    'transaction_id': transaction_id
                },
                'qr_format': 'VietQR API',
                'api_url': qr_url
            }
    except Exception as e:
        print(f"[WARNING] VietQR API failed, falling back to manual generation: {e}")
    
    # Fallback: Tạo QR thủ công nếu API thất bại
    
    def add_field(id_code, value):
        """Thêm field vào EMV QR payload"""
        if value is None or value == '':
            return ''
        value_str = str(value)
        length = len(value_str)
        return f"{id_code:02d}{length:02d}{value_str}"
    
    # Build EMV QR Code payload đầy đủ theo chuẩn VietQR
    payload = ""
    
    # 00: Payload Format Indicator - luôn là "01"
    payload += add_field(0, "01")
    
    # 01: Point of Initiation Method - "12" = dynamic QR
    payload += add_field(1, "12")
    
    # 38: Merchant Account Information theo chuẩn VietQR
    merchant_info = ""
    merchant_info += add_field(0, "A000000727")  # GUID for VietQR
    # 01: Acquirer - BIN code của ngân hàng
    merchant_info += add_field(1, tpbank_bin)
    # 02: Merchant ID - Số tài khoản
    merchant_info += add_field(2, account_number)
    payload += add_field(38, merchant_info)
    
    # 52: Merchant Category Code - "0000" = không phân loại
    payload += add_field(52, "0000")
    
    # 53: Transaction Currency - "704" = VND
    payload += add_field(53, "704")
    
    # 54: Transaction Amount
    payload += add_field(54, str(amount))
    
    # 58: Country Code - "VN"
    payload += add_field(58, "VN")
    
    # 59: Merchant Name - Viết hoa, không dấu, tối đa 50 ký tự
    merchant_name = to_uppercase_no_accent(BANK_ACCOUNT_NAME, 50)
    payload += add_field(59, merchant_name)
    
    # 60: Merchant City - Viết hoa, không dấu
    merchant_city = to_uppercase_no_accent(BANK_BRANCH, 15)
    payload += add_field(60, merchant_city)
    
    # 62: Additional Data Field Template
    additional_data = ""
    # 05: Purpose of Transaction (nội dung chuyển khoản)
    additional_data += add_field(5, content)
    # 08: Reference Label (transaction_id)
    additional_data += add_field(8, transaction_id)
    payload += add_field(62, additional_data)
    
    # Tính CRC16-CCITT checksum
    def crc16_ccitt(data):
        """Tính CRC16-CCITT cho EMV QR Code"""
        crc = 0xFFFF
        polynomial = 0x1021
        for byte in data.encode('utf-8'):
            crc ^= (byte << 8)
            for _ in range(8):
                if crc & 0x8000:
                    crc = (crc << 1) ^ polynomial
                else:
                    crc <<= 1
                crc &= 0xFFFF
        return format(crc, '04X')
    
    # 63: CRC16 checksum
    checksum = crc16_ccitt(payload)
    qr_data = payload + "6304" + checksum
    
    # Tạo QR code với chất lượng cao theo chuẩn VietQR
    qr = qrcode.QRCode(
        version=None,  # Tự động chọn version phù hợp
        error_correction=qrcode.constants.ERROR_CORRECT_M,  # Mức sửa lỗi trung bình
        box_size=15,  # Tăng kích thước để dễ quét hơn
        border=4,  # Border rõ ràng (quiet zone theo chuẩn)
    )
    qr.add_data(qr_data)
    qr.make(fit=True)
    
    # Tạo image với độ phân giải cao
    img = qr.make_image(fill_color="black", back_color="white")
    
    # Resize để đảm bảo chất lượng tốt (tối thiểu 500x500)
    from PIL import Image
    if img.size[0] < 500:
        img = img.resize((500, 500), Image.Resampling.LANCZOS)
    
    # Convert to base64
    buffer = io.BytesIO()
    img.save(buffer, format='PNG', optimize=False)  # Không optimize để giữ chất lượng
    img_str = base64.b64encode(buffer.getvalue()).decode()
    
    return {
        'qr_image': f"data:image/png;base64,{img_str}",
        'bank_info': {
            'bank_name': BANK_NAME,
            'account_number': account_number,  # Sử dụng số tài khoản đã làm sạch
            'account_name': BANK_ACCOUNT_NAME,
            'branch': BANK_BRANCH,
            'amount': amount,
            'content': content,
            'transaction_id': transaction_id
        },
        'qr_format': 'VietQR/EMV',
        'merchant_name': merchant_name,
        'merchant_city': merchant_city
    }

@app.route('/payment/bank/verify', methods=['POST'])
def verify_bank_transfer():
    """Xác nhận đã chuyển khoản (manual verification cho bank_qr)"""
    if not is_logged_in():
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401
    
    data = request.get_json()
    payment_id = data.get('payment_id')
    transaction_proof = data.get('transaction_proof', '')
    
    # Debug: Log user info
    current_user_id = session.get('user_id')
    current_username = session.get('username', 'Unknown')
    print(f"[DEBUG] Verify payment - User ID: {current_user_id}, Username: {current_username}, Payment ID: {payment_id}")
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': 'Database connection failed'}), 500
    
    try:
        with conn.cursor() as cursor:
            cursor.execute("""
                SELECT p.*, sp.characters_limit, sp.duration_days, sp.package_name
                FROM payments p
                LEFT JOIN subscription_packages sp ON p.package_id = sp.id
                WHERE p.id = %s AND p.user_id = %s AND p.payment_method = 'bank_qr'
            """, (payment_id, current_user_id))
            payment = cursor.fetchone()
            
            if not payment:
                print(f"[ERROR] Payment not found - ID: {payment_id}, User ID: {current_user_id}")
                return jsonify({'success': False, 'message': 'Payment not found'}), 404
            
            if payment['payment_status'] == 'completed':
                return jsonify({'success': False, 'message': 'Payment đã được xác nhận'}), 400
            
            print(f"[DEBUG] Payment found: {payment}")
            
            # Thử verify qua SePay trước
            verification = verify_sepay_transaction(
                payment['transaction_id'], 
                payment['amount_vnd']
            )
            
            print(f"[DEBUG] SePay verification: {verification}")
            
            if verification['success'] and verification['verified']:
                # Thanh toán thành công qua SePay
                cursor.execute("""
                    UPDATE payments
                    SET payment_status = 'completed',
                        bank_transaction_id = %s,
                        description = %s,
                        completed_at = NOW()
                    WHERE id = %s
                """, (
                    payment['transaction_id'], 
                    'Thanh toán SePay thành công (auto verified)', 
                    payment_id
                ))
                
                # Cập nhật subscription
                print(f"[DEBUG] Updating subscription for user {current_user_id}")
                success = update_user_subscription(
                    current_user_id,
                    payment['characters_limit'],
                    payment['duration_days'],
                    payment.get('package_id'),
                )
                
                print(f"[DEBUG] Subscription update result: {success}")
                
                conn.commit()
                
                if success:
                    try:
                        _notify_payment_success_from_row(payment)
                    except Exception as ne:
                        print(f"[WARN] payment notify: {ne}")
                    return jsonify({
                        'success': True,
                        'message': f'🎉 THANH TOÁN THÀNH CÔNG!\n\n✅ Đã mua gói {payment.get("package_name", "Basic Plan")} thành công!\n💰 Số tiền: {payment["amount_vnd"]:,}đ\n📝 Ký tự được thêm: +{payment["characters_limit"]:,}\n⏰ Thời hạn: +{payment["duration_days"]} ngày\n\n🚀 Bạn có thể sử dụng dịch vụ ngay bây giờ!',
                        'auto_verified': True,
                        'purchase_info': {
                            'package_name': payment.get('package_name', 'Basic Plan'),
                            'amount': payment['amount_vnd'],
                            'characters_added': payment['characters_limit'],
                            'duration_days': payment['duration_days']
                        }
                    })
                else:
                    return jsonify({
                        'success': True, 
                        'message': '✅ Thanh toán thành công nhưng có lỗi cập nhật gói dịch vụ. Vui lòng liên hệ admin để được hỗ trợ.',
                        'auto_verified': True
                    })
            else:
                # Không tìm thấy giao dịch SePay, chuyển sang manual verification
                cursor.execute("""
                    UPDATE payments
                    SET payment_status = 'pending',
                        bank_transaction_id = %s,
                        description = %s
                    WHERE id = %s
                """, (transaction_proof, f'Chờ admin duyệt - Tham chiếu: {transaction_proof}', payment_id))
                conn.commit()
                
                return jsonify({
                    'success': True,
                    'message': f'✅ Đã ghi nhận yêu cầu thanh toán!\n\n📋 Gói: {payment.get("package_name", "Basic Plan")}\n💰 Số tiền: {payment["amount_vnd"]:,}đ\n🏦 Mã tham chiếu: {transaction_proof}\n\n⏳ Admin sẽ duyệt thanh toán trong vòng 24 giờ.\nBạn sẽ nhận được thông báo khi gói dịch vụ được kích hoạt.',
                    'manual_verification': True,
                    'pending_info': {
                        'package_name': payment.get('package_name', 'Basic Plan'),
                        'amount': payment['amount_vnd'],
                        'reference': transaction_proof
                    }
                })
            
            return jsonify({
                'success': True,
                'message': 'Đã gửi yêu cầu xác nhận. Admin sẽ duyệt trong vòng 24 giờ.'
            })
    except Exception as e:
        conn.rollback()
        print(f"[ERROR] Verify bank transfer error: {e}")
        return jsonify({'success': False, 'message': f'Lỗi: {str(e)}'}), 500
    finally:
        conn.close()

@app.route('/api/payment/sepay/verify', methods=['POST'])
def verify_sepay_payment():
    """Xác minh thanh toán SePay"""
    if not is_logged_in():
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401
    
    data = request.get_json()
    payment_id = data.get('payment_id')
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': 'Database connection failed'}), 500
    
    try:
        with conn.cursor() as cursor:
            cursor.execute("""
                SELECT p.*, sp.characters_limit, sp.duration_days
                FROM payments p
                LEFT JOIN subscription_packages sp ON p.package_id = sp.id
                WHERE p.id = %s AND p.user_id = %s AND p.payment_method = 'sepay'
            """, (payment_id, session['user_id']))
            payment = cursor.fetchone()
            
            if not payment:
                return jsonify({'success': False, 'message': 'Payment not found'}), 404
            
            if payment['payment_status'] == 'completed':
                return jsonify({'success': False, 'message': 'Payment đã được xác nhận'}), 400
            
            # Xác minh qua SePay API
            verification = verify_sepay_transaction(
                payment['transaction_id'], 
                payment['amount_vnd']
            )
            
            if verification['success'] and verification['verified']:
                # Thanh toán thành công, cập nhật database
                cursor.execute("""
                    UPDATE payments
                    SET payment_status = 'completed',
                        bank_transaction_id = %s,
                        description = %s,
                        completed_at = NOW()
                    WHERE id = %s
                """, (
                    payment['transaction_id'], 
                    'Thanh toán SePay thành công', 
                    payment_id
                ))
                
                # Cập nhật subscription cho user
                update_user_subscription(
                    session['user_id'],
                    payment['characters_limit'],
                    payment['duration_days'],
                    payment.get('package_id'),
                )
                
                conn.commit()
                
                try:
                    _notify_payment_success_from_row(payment)
                except Exception as ne:
                    print(f"[WARN] payment notify: {ne}")
                
                return jsonify({
                    'success': True,
                    'message': 'Thanh toán thành công! Đã cập nhật gói dịch vụ.',
                    'verified': True
                })
            else:
                return jsonify({
                    'success': False,
                    'message': 'Chưa tìm thấy giao dịch. Vui lòng thử lại sau.',
                    'verified': False
                })
                
    except Exception as e:
        conn.rollback()
        print(f"[ERROR] Verify SePay payment error: {e}")
        return jsonify({'success': False, 'message': f'Lỗi: {str(e)}'}), 500
    finally:
        conn.close()

@app.route('/api/payment/debug/sub')
def debug_subscription():
    """DEBUG ONLY - xem subscription hiện tại"""
    conn = get_db_connection()
    if not conn:
        return jsonify({'error': 'DB failed'})
    try:
        with conn.cursor() as cursor:
            cursor.execute("""
                SELECT us.*, u.username FROM user_subscriptions us
                JOIN users u ON us.user_id = u.id
                ORDER BY us.created_at DESC LIMIT 5
            """)
            rows = cursor.fetchall()
            # Serialize dates
            result = []
            for r in rows:
                d = dict(r)
                for k, v in d.items():
                    if hasattr(v, 'isoformat'):
                        d[k] = v.isoformat()
                result.append(d)
            return jsonify({'subscriptions': result})
    except Exception as e:
        return jsonify({'error': str(e)})
    finally:
        conn.close()

@app.route('/api/payment/debug/pending')
def debug_pending_payments():
    """DEBUG ONLY - xem pending payments trong DB"""
    conn = get_db_connection()
    if not conn:
        return jsonify({'error': 'DB connection failed'})
    try:
        with conn.cursor() as cursor:
            cursor.execute("""
                SELECT p.id, p.transaction_id, p.amount_vnd, p.payment_status, p.created_at,
                       u.username
                FROM payments p
                JOIN users u ON p.user_id = u.id
                WHERE p.payment_status IN ('pending', 'completed')
                ORDER BY p.created_at DESC LIMIT 10
            """)
            rows = cursor.fetchall()
            return jsonify({'payments': [dict(r) for r in rows]})
    except Exception as e:
        return jsonify({'error': str(e)})
    finally:
        conn.close()

@app.route('/api/payment/sepay/webhook', methods=['POST'])
@app.route('/webhook/sepay', methods=['POST'])
def sepay_webhook():
    """
    Xử lý webhook từ SePay - tự động duyệt thanh toán và cộng ký tự.
    SePay webhook format:
    {
      "id": 123456,
      "gateway": "MBBank",
      "transactionDate": "2024-01-15 10:30:00",
      "accountNumber": "0866005541",
      "content": "TXN12345ABC",
      "transferType": "in",
      "amount": 50000,
      "accumulated": 100000,
      "referenceCode": "...",
      "code": null
    }
    """
    try:
        data = request.get_json(force=True) or {}
        print(f"[WEBHOOK] ===== SePay webhook received =====")
        print(f"[WEBHOOK] Full data: {data}")
        
        # Chỉ xử lý giao dịch tiền VÀO tài khoản
        transfer_type = str(data.get('transferType', '') or data.get('transfer_type', '') or '').strip()
        print(f"[WEBHOOK] transferType='{transfer_type}'")
        if transfer_type and transfer_type.lower() not in ('in', 'credit', 'receive'):
            print(f"[WEBHOOK] Ignored - transfer type is outgoing: {transfer_type}")
            return jsonify({'success': True, 'message': 'Ignored (not incoming)'})
        
        # SePay gửi số tiền ở field 'transferAmount' (không phải 'amount')
        webhook_amount = int(float(
            data.get('transferAmount') or
            data.get('transfer_amount') or
            data.get('amount') or
            0
        ))
        # SePay có thể gửi content ở nhiều field khác nhau
        raw_content = (
            data.get('content') or
            data.get('transaction_content') or
            data.get('description') or
            data.get('memo') or
            data.get('reference') or ''
        )
        webhook_content = str(raw_content).upper().strip()
        
        print(f"[WEBHOOK] transferAmount={data.get('transferAmount')}, amount={data.get('amount')}, "
              f"resolved_amount={webhook_amount}")
        print(f"[WEBHOOK] raw_content='{raw_content}', content_upper='{webhook_content}'")
        
        if not webhook_content or webhook_amount <= 0:
            print(f"[WEBHOOK] Missing content or amount — cannot match")
            return jsonify({'success': True, 'message': 'No content/amount'})
        
        conn = get_db_connection()
        if not conn:
            return jsonify({'success': False, 'message': 'DB error'}), 500
        
        try:
            with conn.cursor() as cursor:
                # Lấy TẤT CẢ pending payments (bỏ filter amount ở SQL để debug)
                cursor.execute("""
                    SELECT p.*, sp.characters_limit, sp.duration_days, sp.package_name,
                           u.username
                    FROM payments p
                    JOIN subscription_packages sp ON p.package_id = sp.id
                    JOIN users u ON p.user_id = u.id
                    WHERE p.payment_status = 'pending'
                      AND UPPER(p.transaction_id) != ''
                    ORDER BY p.created_at DESC
                    LIMIT 50
                """)
                
                pending_payments = cursor.fetchall()
                print(f"[WEBHOOK] Found {len(pending_payments)} pending payments")
                matched_payment = None
                
                # Chuẩn hóa: chỉ giữ alphanum, chuyển uppercase
                def normalize(s):
                    return ''.join(c for c in str(s).upper() if c.isalnum())
                
                webhook_norm = normalize(webhook_content)
                print(f"[WEBHOOK] webhook_norm='{webhook_norm}'")

                for pmt in pending_payments:
                    txn_id      = str(pmt['transaction_id']).upper().strip()
                    txn_norm    = normalize(txn_id)
                    # Bỏ prefix TTS (có hoặc không có gạch dưới)
                    hex_part    = txn_norm
                    for pfx in ('TTS_', 'TTS'):
                        if hex_part.startswith(pfx):
                            hex_part = hex_part[len(pfx):]
                            break
                    
                    amount_ok = abs(webhook_amount - int(pmt['amount_vnd'])) <= 2000

                    # So khớp linh hoạt: full match HOẶC hex part match
                    content_match = (
                        txn_norm    in webhook_norm or
                        webhook_norm in txn_norm   or
                        (len(hex_part) >= 8 and hex_part in webhook_norm)
                    )

                    print(f"[WEBHOOK] Check pmt#{pmt['id']}: txn={txn_id}, txn_norm={txn_norm}, hex={hex_part}, "
                          f"content_match={content_match}, amount_ok={amount_ok} "
                          f"(pmt_amount={pmt['amount_vnd']}, webhook_amount={webhook_amount})")

                    if content_match and amount_ok:
                        matched_payment = pmt
                        print(f"[WEBHOOK] ✅ MATCHED payment #{pmt['id']} for user {pmt['username']}")
                        break
                
                if not matched_payment:
                    print(f"[WEBHOOK] ❌ No matching pending payment. content='{webhook_content}', amount={webhook_amount}")
                    return jsonify({'success': True, 'message': 'No matching payment found'})
                
                print(f"[WEBHOOK] Matched payment ID={matched_payment['id']} for user={matched_payment['username']}")
                
                # Cập nhật payment thành completed
                cursor.execute("""
                    UPDATE payments
                    SET payment_status = 'completed',
                        bank_transaction_id = %s,
                        description = %s,
                        completed_at = NOW()
                    WHERE id = %s AND payment_status = 'pending'
                """, (
                    str(data.get('id', '') or data.get('referenceCode', '')),
                    f"SePay webhook auto-approved: {webhook_content}",
                    matched_payment['id']
                ))
                
                if cursor.rowcount == 0:
                    # Đã được duyệt bởi request khác (race condition)
                    print(f"[WEBHOOK] Payment {matched_payment['id']} already processed")
                    conn.commit()
                    return jsonify({'success': True, 'message': 'Already processed'})
                
                conn.commit()
                
                # Cộng ký tự cho user
                sub_success = update_user_subscription(
                    matched_payment['user_id'],
                    matched_payment['characters_limit'],
                    matched_payment['duration_days'],
                    matched_payment.get('package_id'),
                )
                
                if sub_success:
                    print(f"[WEBHOOK] ✅ Auto-approved payment {matched_payment['id']} for user {matched_payment['username']}: +{matched_payment['characters_limit']:,} chars")
                    try:
                        _notify_payment_success_from_row(matched_payment)
                    except Exception as ne:
                        print(f"[WARN] payment notify: {ne}")
                else:
                    print(f"[WEBHOOK] ⚠️ Payment approved but subscription update failed for user {matched_payment['username']}")
                
                return jsonify({
                    'success': True,
                    'message': 'Payment auto-approved',
                    'payment_id': matched_payment['id'],
                    'user': matched_payment['username'],
                    'characters_added': matched_payment['characters_limit']
                })
                
        except Exception as e:
            import traceback
            conn.rollback()
            print(f"[WEBHOOK ERROR] {e}")
            print(traceback.format_exc())
            return jsonify({'success': False, 'message': f'DB error: {str(e)}'}), 500
        finally:
            conn.close()
            
    except Exception as e:
        import traceback
        print(f"[WEBHOOK FATAL] {e}")
        print(traceback.format_exc())
        return jsonify({'success': False, 'message': str(e)}), 500

def update_user_subscription(user_id, characters_limit, duration_days, package_id=None):
    """Cập nhật subscription cho user (cộng ký tự, gia hạn, gán package_id)."""
    print(
        f"[DEBUG] update_user_subscription - User ID: {user_id}, Characters: {characters_limit}, "
        f"Days: {duration_days}, Package ID: {package_id}"
    )
    
    conn = get_db_connection()
    if not conn:
        print("[ERROR] Database connection failed in update_user_subscription")
        return False
        
    try:
        with conn.cursor() as cursor:
            # Kiểm tra user tồn tại
            cursor.execute("SELECT username FROM users WHERE id = %s", (user_id,))
            user = cursor.fetchone()
            if not user:
                print(f"[ERROR] User {user_id} not found")
                return False
            
            print(f"[DEBUG] Updating subscription for user: {user['username']}")
            
            # Kiểm tra subscription hiện tại
            cursor.execute("""
                SELECT * FROM user_subscriptions 
                WHERE user_id = %s AND is_active = 1
                ORDER BY created_at DESC LIMIT 1
            """, (user_id,))
            current_sub = cursor.fetchone()
            
            if current_sub:
                print(f"[DEBUG] Found existing subscription: {current_sub}")
                
                # Gia hạn subscription hiện tại
                new_end_date = datetime.now() + timedelta(days=duration_days)
                new_characters_used = max(0, current_sub['characters_used'] - characters_limit)  # Giảm characters_used
                new_characters_limit = current_sub['characters_limit'] + characters_limit  # Tăng limit
                
                print(f"[DEBUG] Updating - New limit: {new_characters_limit}, New used: {new_characters_used}, New end date: {new_end_date}")
                
                cursor.execute("""
                    UPDATE user_subscriptions
                    SET characters_used = %s,
                        characters_limit = %s,
                        end_date = %s,
                        package_id = CASE WHEN %s IS NOT NULL THEN %s ELSE package_id END
                    WHERE id = %s
                """, (
                    new_characters_used,
                    new_characters_limit,
                    new_end_date,
                    package_id,
                    package_id,
                    current_sub['id'],
                ))
                
                print(f"[DEBUG] Updated existing subscription for user {user['username']}")
            else:
                print(f"[DEBUG] No existing subscription, creating new one")
                
                # Tạo subscription mới
                start_date = datetime.now().date()
                end_date = start_date + timedelta(days=duration_days)
                cursor.execute("""
                    INSERT INTO user_subscriptions
                    (user_id, package_id, characters_limit, characters_used, start_date, end_date, is_active)
                    VALUES (%s, %s, %s, 0, %s, %s, 1)
                """, (user_id, package_id, characters_limit, start_date, end_date))
                
                print(f"[DEBUG] Created new subscription for user {user['username']}")
            
            conn.commit()
            print(f"[SUCCESS] Subscription updated successfully for user {user['username']}")
            return True
    except Exception as e:
        print(f"[ERROR] Update subscription error: {e}")
        print(f"[ERROR] Error details: {traceback.format_exc()}")
        conn.rollback()
        return False
    finally:
        conn.close()

@app.route('/api/admin/payment/approve', methods=['POST'])
def admin_approve_payment():
    """Admin duyệt thanh toán thủ công — bắt buộc verify SePay.
    Nếu SePay không tìm thấy giao dịch → đánh dấu FAILED, không cộng ký tự.
    """
    if not is_logged_in() or not is_admin():
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401
    
    data       = request.get_json()
    payment_id = data.get('payment_id')
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': 'Database connection failed'}), 500
    
    try:
        with conn.cursor() as cursor:
            cursor.execute("""
                SELECT p.*, sp.characters_limit, sp.duration_days, sp.package_name,
                       u.username
                FROM payments p
                LEFT JOIN subscription_packages sp ON p.package_id = sp.id
                LEFT JOIN users u ON p.user_id = u.id
                WHERE p.id = %s AND p.payment_status = 'pending'
            """, (payment_id,))
            payment = cursor.fetchone()
            
            if not payment:
                return jsonify({'success': False, 'message': 'Không tìm thấy thanh toán hoặc đã được xử lý'}), 404
            
            # ── Verify SePay bắt buộc ──
            verified    = False
            verify_note = ''
            try:
                verification = verify_sepay_transaction(
                    payment['transaction_id'],
                    payment['amount_vnd']
                )
                verified    = verification.get('verified', False)
                verify_note = verification.get('note', '')
                print(f"[ADMIN-APPROVE] SePay verify: verified={verified}, note={verify_note}")
            except Exception as ve:
                print(f"[ADMIN-APPROVE] SePay verify error: {ve}")
                verified    = False
                verify_note = str(ve)
            
            if not verified:
                # ── Không verify được → đánh dấu FAILED, kết thúc thanh toán ──
                cursor.execute("""
                    UPDATE payments
                    SET payment_status = 'failed',
                        description    = CONCAT(IFNULL(description,''), ' - Admin verify failed: không tìm thấy giao dịch SePay'),
                        completed_at   = NOW()
                    WHERE id = %s AND payment_status = 'pending'
                """, (payment_id,))
                conn.commit()
                print(f"[ADMIN-APPROVE] Payment {payment_id} marked FAILED — no SePay transaction found")
                return jsonify({
                    'success': False,
                    'failed': True,
                    'message': (
                        f'❌ Thanh toán #{payment_id} không thể xác minh qua SePay.\n\n'
                        f'Lý do: {verify_note or "Không có giao dịch nào khớp với mã giao dịch và số tiền"}\n\n'
                        f'⚠️ Người dùng CHƯA chuyển khoản hoặc số tiền không khớp.\n\n'
                        f'Thanh toán đã bị đánh dấu THẤT BẠI. Người dùng cần tạo đơn thanh toán mới.'
                    )
                })
            
            # ── SePay xác nhận → tiến hành duyệt ──
            cursor.execute("""
                UPDATE payments
                SET payment_status = 'completed',
                    description    = CONCAT(IFNULL(description,''), ' - SePay verified + Admin approved'),
                    completed_at   = NOW()
                WHERE id = %s AND payment_status = 'pending'
            """, (payment_id,))
            
            if cursor.rowcount == 0:
                conn.rollback()
                return jsonify({'success': False, 'message': 'Thanh toán đã được xử lý bởi request khác'}), 409
            
            conn.commit()
            
            # Cộng ký tự cho user
            success = update_user_subscription(
                payment['user_id'],
                payment['characters_limit'],
                payment['duration_days'],
                payment.get('package_id'),
            )
            
            if success:
                try:
                    _notify_payment_success_from_row(payment)
                except Exception as ne:
                    print(f"[WARN] payment notify: {ne}")
                return jsonify({
                    'success': True,
                    'verified': True,
                    'message': (
                        f'🎉 Đã duyệt thanh toán thành công! ✅ (Đã xác minh SePay)\n\n'
                        f'👤 User: {payment["username"]}\n'
                        f'📋 Gói: {payment.get("package_name","Custom")}\n'
                        f'💰 Số tiền: {payment["amount_vnd"]:,}đ\n'
                        f'📝 Ký tự thêm: +{payment["characters_limit"]:,}\n'
                        f'⏰ Thời hạn thêm: +{payment["duration_days"]} ngày\n\n'
                        f'✅ Gói dịch vụ đã được kích hoạt cho user.'
                    ),
                    'approval_info': {
                        'user': payment["username"],
                        'package_name': payment.get("package_name", "Custom"),
                        'amount': payment["amount_vnd"],
                        'characters_added': payment["characters_limit"],
                        'duration_days': payment["duration_days"]
                    }
                })
            else:
                return jsonify({
                    'success': True,
                    'message': f'⚠️ Đã duyệt cho user {payment["username"]} nhưng có lỗi cập nhật gói. Kiểm tra lại DB.',
                    'warning': True
                })
                
    except Exception as e:
        conn.rollback()
        print(f"[ERROR] Admin approve payment error: {e}")
        return jsonify({'success': False, 'message': f'Lỗi: {str(e)}'}), 500
    finally:
        conn.close()

@app.route('/api/user/subscription/status', methods=['GET'])
def get_user_subscription_status():
    """Lấy trạng thái subscription hiện tại của user"""
    if not is_logged_in():
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': 'Database connection failed'}), 500
    
    try:
        with conn.cursor() as cursor:
            cursor.execute("""
                SELECT us.*, sp.package_name
                FROM user_subscriptions us
                LEFT JOIN subscription_packages sp ON us.package_id = sp.id
                WHERE us.user_id = %s AND us.is_active = 1
                ORDER BY us.created_at DESC LIMIT 1
            """, (session['user_id'],))
            subscription = cursor.fetchone()
            
            if subscription:
                package_name = subscription.get('package_name')
                if not package_name or not subscription.get('package_id'):
                    cursor.execute("""
                        SELECT sp.package_name, sp.id AS package_id
                        FROM payments p
                        JOIN subscription_packages sp ON p.package_id = sp.id
                        WHERE p.user_id = %s AND p.payment_status = 'completed'
                        ORDER BY p.completed_at DESC, p.id DESC
                        LIMIT 1
                    """, (session['user_id'],))
                    pay_pkg = cursor.fetchone()
                    if pay_pkg:
                        package_name = pay_pkg.get('package_name') or package_name
                        if not subscription.get('package_id') and pay_pkg.get('package_id'):
                            cursor.execute(
                                "UPDATE user_subscriptions SET package_id = %s WHERE id = %s",
                                (pay_pkg['package_id'], subscription['id']),
                            )
                            conn.commit()

                characters_remaining = max(0, subscription['characters_limit'] - subscription['characters_used'])
                days_remaining = (subscription['end_date'] - datetime.now().date()).days
                display_package = package_name or 'Gói hiện tại'
                
                # Tạo message status dựa trên subscription
                if characters_remaining > 1000000:  # > 1M chars
                    status_message = f'🚀 Bạn đang có gói {display_package}! Còn {characters_remaining:,} ký tự và {days_remaining} ngày.'
                elif characters_remaining > 100000:  # > 100K chars  
                    status_message = f'✅ Gói {display_package} còn {characters_remaining:,} ký tự và {days_remaining} ngày.'
                elif days_remaining <= 7:  # Sắp hết hạn
                    status_message = f'⚠️ Gói {display_package} sắp hết hạn ({days_remaining} ngày). Hãy gia hạn sớm!'
                    try:
                        _send_user_notification_email(
                            session['user_id'],
                            'plan_expiry',
                            '[VietVoice] Gói dịch vụ sắp hết hạn',
                            f'Gói của bạn còn {days_remaining} ngày. Hãy gia hạn tại trang Bảng giá.',
                        )
                    except Exception as ne:
                        print(f"[WARN] plan expiry notify: {ne}")
                elif characters_remaining <= 10000:  # Sắp hết ký tự
                    status_message = f'⚠️ Gói {display_package} còn ít ký tự ({characters_remaining:,}). Hãy nâng cấp!'
                else:
                    status_message = f'📋 Gói {display_package}: {characters_remaining:,} ký tự, {days_remaining} ngày.'
                
                return jsonify({
                    'success': True,
                    'subscription': {
                        'package_name': display_package,
                        'characters_limit': subscription['characters_limit'],
                        'characters_used': subscription['characters_used'],
                        'characters_remaining': characters_remaining,
                        'end_date': subscription['end_date'].strftime('%Y-%m-%d'),
                        'days_remaining': days_remaining,
                        'is_active': subscription['is_active'],
                        'status_message': status_message
                    }
                })
            else:
                return jsonify({
                    'success': True,
                    'subscription': {
                        'package_name': 'Chưa có gói',
                        'characters_limit': 0,
                        'characters_used': 0,
                        'characters_remaining': 0,
                        'end_date': None,
                        'days_remaining': 0,
                        'is_active': False,
                        'status_message': '📭 Bạn chưa có gói dịch vụ nào. Hãy mua gói để sử dụng!'
                    }
                })
                
    except Exception as e:
        print(f"[ERROR] Get subscription status error: {e}")
        return jsonify({'success': False, 'message': f'Lỗi: {str(e)}'}), 500
    finally:
        conn.close()

@app.route('/api/admin/payments', methods=['GET'])
def admin_get_payments():
    """Admin xem danh sách payments (có phân trang)"""
    if not is_logged_in() or not is_admin():
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401

    page = max(request.args.get('page', 1, type=int), 1)
    per_page = request.args.get('per_page', 15, type=int)
    per_page = min(max(per_page, 5), 100)

    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': 'Database connection failed'}), 500

    try:
        with conn.cursor() as cursor:
            cursor.execute("SELECT COUNT(*) AS total FROM payments")
            total_row = cursor.fetchone()
            total = total_row['total'] if total_row else 0
            total_pages = max(1, (total + per_page - 1) // per_page) if total > 0 else 1

            if page > total_pages and total > 0:
                page = total_pages

            offset = (page - 1) * per_page
            cursor.execute("""
                SELECT p.*, u.username, sp.package_name
                FROM payments p
                LEFT JOIN users u ON p.user_id = u.id
                LEFT JOIN subscription_packages sp ON p.package_id = sp.id
                ORDER BY p.created_at DESC
                LIMIT %s OFFSET %s
            """, (per_page, offset))
            payments = cursor.fetchall()

            for payment in payments:
                if payment['created_at']:
                    payment['created_at'] = payment['created_at'].strftime('%Y-%m-%d %H:%M:%S')
                if payment['updated_at']:
                    payment['updated_at'] = payment['updated_at'].strftime('%Y-%m-%d %H:%M:%S')
                if payment['completed_at']:
                    payment['completed_at'] = payment['completed_at'].strftime('%Y-%m-%d %H:%M:%S')

            return jsonify({
                'success': True,
                'payments': payments,
                'page': page,
                'per_page': per_page,
                'total': total,
                'total_pages': total_pages,
            })

    except Exception as e:
        print(f"[ERROR] Admin get payments error: {e}")
        return jsonify({'success': False, 'message': f'Lỗi: {str(e)}'}), 500
    finally:
        conn.close()

@app.route('/api/admin/auto-approve', methods=['POST'])
def admin_bulk_auto_approve():
    """Admin kích hoạt auto-approve cho tất cả pending payments"""
    if not is_logged_in() or not is_admin():
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401
    
    data = request.get_json()
    force_approve = data.get('force', False)  # Force approve even without verification
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': 'Database connection failed'}), 500
    
    try:
        with conn.cursor() as cursor:
            # Get all pending payments
            cursor.execute("""
                SELECT p.*, sp.characters_limit, sp.duration_days, u.username
                FROM payments p
                LEFT JOIN subscription_packages sp ON p.package_id = sp.id
                LEFT JOIN users u ON p.user_id = u.id
                WHERE p.payment_status = 'pending'
                ORDER BY p.created_at ASC
            """)
            pending_payments = cursor.fetchall()
            
            approved_count = 0
            failed_count = 0
            results = []
            
            for payment in pending_payments:
                try:
                    payment_id = payment['id']
                    
                    # Try SePay verification first
                    verification = verify_sepay_transaction(
                        payment['transaction_id'], 
                        payment['amount_vnd']
                    )
                    
                    should_approve = force_approve or (verification.get('success') and verification.get('verified'))
                    
                    if should_approve:
                        # Approve payment
                        cursor.execute("""
                            UPDATE payments
                            SET payment_status = 'completed',
                                description = CONCAT(IFNULL(description, ''), ' - Bulk auto-approved'),
                                completed_at = NOW()
                            WHERE id = %s
                        """, (payment_id,))
                        
                        # Update user subscription
                        success = update_user_subscription(
                            payment['user_id'],
                            payment['characters_limit'],
                            payment['duration_days'],
                            payment.get('package_id'),
                        )
                        
                        if success:
                            approved_count += 1
                            results.append({
                                'payment_id': payment_id,
                                'user': payment['username'],
                                'amount': payment['amount_vnd'],
                                'status': 'approved',
                                'method': 'sepay_verified' if verification.get('verified') else 'force_approved'
                            })
                        else:
                            failed_count += 1
                            results.append({
                                'payment_id': payment_id,
                                'user': payment['username'],
                                'amount': payment['amount_vnd'],
                                'status': 'failed',
                                'error': 'Subscription update failed'
                            })
                    else:
                        results.append({
                            'payment_id': payment_id,
                            'user': payment['username'],
                            'amount': payment['amount_vnd'],
                            'status': 'skipped',
                            'reason': 'Not verified'
                        })
                        
                except Exception as e:
                    failed_count += 1
                    results.append({
                        'payment_id': payment.get('id', 'unknown'),
                        'user': payment.get('username', 'unknown'),  
                        'status': 'error',
                        'error': str(e)
                    })
            
            conn.commit()
            
            success_message = f'🎉 BULK AUTO-APPROVE HOÀN THÀNH!\n\n📊 Kết quả:\n✅ Đã duyệt: {approved_count} payments\n❌ Thất bại: {failed_count} payments\n📋 Tổng xử lý: {len(pending_payments)} payments'
            
            if approved_count > 0:
                success_message += f'\n\n💫 {approved_count} user đã được kích hoạt gói dịch vụ!'
            
            return jsonify({
                'success': True,
                'message': success_message,
                'approved_count': approved_count,
                'failed_count': failed_count,
                'total_processed': len(pending_payments),
                'results': results,
                'summary': f'Processed {len(pending_payments)} payments: {approved_count} approved, {failed_count} failed'
            })
            
    except Exception as e:
        conn.rollback()
        print(f"[ERROR] Bulk auto-approve error: {e}")
        return jsonify({'success': False, 'message': f'Lỗi: {str(e)}'}), 500
    finally:
        conn.close()

@app.route('/api/payment/status/<int:payment_id>', methods=['GET'])
def get_payment_status(payment_id):
    """Kiểm tra trạng thái thanh toán cụ thể"""
    if not is_logged_in():
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': 'Database connection failed'}), 500
    
    try:
        with conn.cursor() as cursor:
            cursor.execute("""
                SELECT p.*, sp.package_name, sp.characters_limit, sp.duration_days
                FROM payments p
                LEFT JOIN subscription_packages sp ON p.package_id = sp.id
                WHERE p.id = %s AND p.user_id = %s
            """, (payment_id, session['user_id']))
            payment = cursor.fetchone()
            
            if not payment:
                return jsonify({'success': False, 'message': 'Payment not found'}), 404
            
            # Tạo message dựa trên trạng thái
            if payment['payment_status'] == 'completed':
                message = f'🎉 THANH TOÁN THÀNH CÔNG!\n\n✅ Gói {payment["package_name"]} đã được kích hoạt!\n💰 Số tiền: {payment["amount_vnd"]:,}đ\n📝 Ký tự: +{payment["characters_limit"]:,}\n⏰ Thời hạn: +{payment["duration_days"]} ngày\n\n🚀 Bạn có thể sử dụng dịch vụ ngay!'
                status_icon = '✅'
            elif payment['payment_status'] == 'pending':
                message = f'⏳ Đang chờ xác nhận thanh toán\n\n📋 Gói: {payment["package_name"]}\n💰 Số tiền: {payment["amount_vnd"]:,}đ\n🏦 Mã giao dịch: {payment["transaction_id"]}\n\n⏰ Admin sẽ duyệt trong vòng 24 giờ.'
                status_icon = '⏳'
            elif payment['payment_status'] == 'failed':
                message = f'❌ Thanh toán thất bại\n\n📋 Gói: {payment["package_name"]}\n💰 Số tiền: {payment["amount_vnd"]:,}đ\n\n💡 Vui lòng thử lại hoặc liên hệ support.'
                status_icon = '❌'
            else:
                message = f'📋 Trạng thái: {payment["payment_status"]}\nGói: {payment["package_name"]}\nSố tiền: {payment["amount_vnd"]:,}đ'
                status_icon = '📋'
            
            return jsonify({
                'success': True,
                'payment': {
                    'id': payment['id'],
                    'status': payment['payment_status'],
                    'status_icon': status_icon,
                    'message': message,
                    'package_name': payment['package_name'],
                    'amount': payment['amount_vnd'],
                    'transaction_id': payment['transaction_id'],
                    'created_at': payment['created_at'].strftime('%Y-%m-%d %H:%M:%S') if payment['created_at'] else None,
                    'completed_at': payment['completed_at'].strftime('%Y-%m-%d %H:%M:%S') if payment['completed_at'] else None
                }
            })
            
    except Exception as e:
        print(f"[ERROR] Get payment status error: {e}")
        return jsonify({'success': False, 'message': f'Lỗi: {str(e)}'}), 500
    finally:
        conn.close()

# ─────────────────────────────────────────────────────────────────────────────
# INVOICE / HÓA ĐƠN
# ─────────────────────────────────────────────────────────────────────────────

def _build_invoice_pdf(payment, user_info):
    """Generate PDF invoice bytes using reportlab."""
    from io import BytesIO
    from reportlab.pdfgen import canvas as rl_canvas
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    buffer = BytesIO()
    W, H = A4  # 595 x 842 pt

    # ── Font registration (try Windows Arial for Vietnamese support) ──
    font_reg  = 'Helvetica'
    font_bold = 'Helvetica-Bold'
    _win_fonts = [
        ('C:/Windows/Fonts/arial.ttf',   'C:/Windows/Fonts/arialbd.ttf'),
        ('C:/Windows/Fonts/verdana.ttf', 'C:/Windows/Fonts/verdanab.ttf'),
    ]
    for reg_path, bold_path in _win_fonts:
        if os.path.exists(reg_path):
            try:
                pdfmetrics.registerFont(TTFont('VV_Reg',  reg_path))
                font_reg  = 'VV_Reg'
                if os.path.exists(bold_path):
                    pdfmetrics.registerFont(TTFont('VV_Bold', bold_path))
                    font_bold = 'VV_Bold'
                break
            except Exception:
                pass

    c = rl_canvas.Canvas(buffer, pagesize=A4)

    # ── Palette ──
    COL_BG        = (0.020, 0.082, 0.122)   # #051424
    COL_PRIMARY   = (0.439, 0.471, 1.000)   # purplish-blue
    COL_TEXT      = (0.831, 0.894, 0.980)   # #d4e4fa
    COL_MUTED     = (0.592, 0.557, 0.627)
    COL_ACCENT    = (0.188, 0.851, 0.957)   # tertiary teal
    COL_WHITE     = (1, 1, 1)
    COL_SURFACE   = (0.071, 0.133, 0.192)

    def rgb(col): c.setFillColorRGB(*col)
    def stroke(col): c.setStrokeColorRGB(*col)

    # ── Background ──
    rgb(COL_BG)
    c.rect(0, 0, W, H, fill=1, stroke=0)

    # ── Header band ──
    rgb(COL_SURFACE)
    c.rect(0, H - 90, W, 90, fill=1, stroke=0)

    # Logo / brand
    rgb(COL_PRIMARY)
    c.setFont(font_bold, 22)
    c.drawString(40, H - 55, 'VietVoice AI')
    rgb(COL_MUTED)
    c.setFont(font_reg, 10)
    c.drawString(40, H - 72, 'Ho tro: support@vietvoice.ai')

    # Invoice label (right side)
    rgb(COL_WHITE)
    c.setFont(font_bold, 28)
    c.drawRightString(W - 40, H - 52, 'HOA DON / INVOICE')
    rgb(COL_MUTED)
    c.setFont(font_reg, 10)
    inv_no = f'VV-{payment["id"]:05d}'
    c.drawRightString(W - 40, H - 68, f'So: {inv_no}')

    # ── Invoice meta ──
    y = H - 115
    rgb(COL_TEXT)
    c.setFont(font_bold, 11)
    c.drawString(40, y, 'THONG TIN HOA DON')
    rgb(COL_PRIMARY)
    c.setLineWidth(1.5)
    stroke(COL_PRIMARY)
    c.line(40, y - 5, W - 40, y - 5)

    y -= 28
    meta_rows = [
        ('Khach hang:', user_info.get('username', '—')),
        ('Ma giao dich:', payment.get('transaction_id', '—')),
        ('Ngay tao:', str(payment.get('created_at', '—'))[:19]),
        ('Ngay thanh toan:', str(payment.get('completed_at', '—'))[:19] if payment.get('completed_at') else '—'),
        ('Phuong thuc:', payment.get('payment_method', 'bank_qr').replace('_', ' ').upper()),
        ('Trang thai:', 'DA THANH TOAN' if payment.get('payment_status') == 'completed' else payment.get('payment_status', '').upper()),
    ]
    for label, val in meta_rows:
        rgb(COL_MUTED)
        c.setFont(font_reg, 10)
        c.drawString(40, y, label)
        rgb(COL_TEXT)
        c.setFont(font_bold if label in ('Trang thai:', 'Khach hang:') else font_reg, 10)
        c.drawString(200, y, val)
        y -= 20

    # ── Items table ──
    y -= 18
    rgb(COL_TEXT)
    c.setFont(font_bold, 11)
    c.drawString(40, y, 'CHI TIET DICH VU')
    stroke(COL_PRIMARY)
    c.line(40, y - 5, W - 40, y - 5)

    # Table header
    y -= 28
    rgb(COL_SURFACE)
    c.rect(38, y - 6, W - 76, 24, fill=1, stroke=0)
    rgb(COL_PRIMARY)
    c.setFont(font_bold, 10)
    c.drawString(46, y + 4, 'Dich vu / Goi')
    c.drawString(280, y + 4, 'Ky tu')
    c.drawString(380, y + 4, 'Thoi han')
    c.drawRightString(W - 46, y + 4, 'So tien (VND)')

    # Table row
    y -= 30
    rgb(COL_TEXT)
    c.setFont(font_reg, 10)
    pkg_name = payment.get('package_name', 'Goi dich vu')
    chars    = f'{payment.get("characters_limit", 0):,}'.replace(',', '.')
    days     = f'{payment.get("duration_days", 0)} ngay'
    amount_s = f'{payment.get("amount_vnd", 0):,}'.replace(',', '.') + ' d'
    c.drawString(46, y, pkg_name)
    c.drawString(280, y, chars)
    c.drawString(380, y, days)
    c.drawRightString(W - 46, y, amount_s)

    # Total row
    y -= 20
    stroke(COL_MUTED)
    c.setLineWidth(0.5)
    c.line(38, y, W - 38, y)
    y -= 22
    rgb(COL_ACCENT)
    c.setFont(font_bold, 13)
    c.drawString(46, y, 'TONG CONG:')
    c.drawRightString(W - 46, y, amount_s)

    # ── Footer ──
    stroke(COL_SURFACE)
    c.setLineWidth(1)
    c.line(40, 60, W - 40, 60)
    rgb(COL_MUTED)
    c.setFont(font_reg, 9)
    c.drawCentredString(W / 2, 44, 'Cam on ban da su dung VietVoice AI  |  vietvoice.ai  |  support@vietvoice.ai')
    c.drawCentredString(W / 2, 30, f'Tai lieu nay duoc tao tu dong - Invoice #{inv_no}')

    c.save()
    buffer.seek(0)
    return buffer


@app.route('/invoice/<int:payment_id>')
@login_required
def download_invoice(payment_id):
    """Tải hóa đơn PDF cho một payment đã hoàn thành"""
    conn = get_db_connection()
    if not conn:
        return jsonify({'error': 'Database connection failed'}), 500

    try:
        with conn.cursor() as cursor:
            cursor.execute("""
                SELECT p.id, p.transaction_id, p.amount_vnd, p.payment_method,
                       p.payment_status, p.created_at, p.completed_at,
                       sp.package_name, sp.characters_limit, sp.duration_days,
                       u.username, u.email, u.full_name
                FROM payments p
                JOIN subscription_packages sp ON p.package_id = sp.id
                JOIN users u ON p.user_id = u.id
                WHERE p.id = %s AND p.user_id = %s
            """, (payment_id, session['user_id']))
            payment = cursor.fetchone()

        if not payment:
            return jsonify({'error': 'Không tìm thấy hóa đơn'}), 404

        if payment['payment_status'] != 'completed':
            return jsonify({'error': 'Hóa đơn chỉ khả dụng sau khi thanh toán thành công'}), 403

        user_info = {
            'username':  payment.get('full_name') or payment.get('username', ''),
            'email':     payment.get('email', ''),
        }

        pdf_buffer = _build_invoice_pdf(payment, user_info)
        filename = f'VietVoice_Invoice_VV-{payment_id:05d}.pdf'

        return send_file(
            pdf_buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=filename
        )

    except Exception as e:
        print(f'[ERROR] Invoice generation: {e}')
        return jsonify({'error': str(e)}), 500
    finally:
        conn.close()


@app.route('/api/user/payments')
@login_required
def get_user_payments():
    """Lịch sử thanh toán của user hiện tại (có phân trang)"""
    page = max(request.args.get('page', 1, type=int), 1)
    per_page = request.args.get('per_page', 8, type=int)
    per_page = min(max(per_page, 5), 50)

    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'payments': []}), 500

    try:
        with conn.cursor() as cursor:
            cursor.execute(
                "SELECT COUNT(*) AS total FROM payments WHERE user_id = %s",
                (session['user_id'],),
            )
            total_row = cursor.fetchone()
            total = total_row['total'] if total_row else 0
            total_pages = max(1, (total + per_page - 1) // per_page) if total > 0 else 1

            if page > total_pages and total > 0:
                page = total_pages

            offset = (page - 1) * per_page
            cursor.execute("""
                SELECT p.id, p.transaction_id, p.amount_vnd, p.payment_method,
                       p.payment_status, p.created_at, p.completed_at,
                       sp.package_name, sp.characters_limit, sp.duration_days
                FROM payments p
                LEFT JOIN subscription_packages sp ON p.package_id = sp.id
                WHERE p.user_id = %s
                ORDER BY p.created_at DESC
                LIMIT %s OFFSET %s
            """, (session['user_id'], per_page, offset))
            rows = cursor.fetchall()

        payments = []
        for r in rows:
            payments.append({
                'id':              r['id'],
                'transaction_id':  r['transaction_id'],
                'amount_vnd':      r['amount_vnd'],
                'payment_method':  r['payment_method'],
                'payment_status':  r['payment_status'],
                'package_name':    r.get('package_name', '—'),
                'characters_limit': r.get('characters_limit', 0),
                'duration_days':   r.get('duration_days', 0),
                'created_at':      r['created_at'].strftime('%d/%m/%Y %H:%M') if r['created_at'] else '—',
                'completed_at':    r['completed_at'].strftime('%d/%m/%Y %H:%M') if r['completed_at'] else None,
            })

        return jsonify({
            'success': True,
            'payments': payments,
            'page': page,
            'per_page': per_page,
            'total': total,
            'total_pages': total_pages,
        })

    except Exception as e:
        print(f'[ERROR] get_user_payments: {e}')
        return jsonify({'success': False, 'payments': []}), 500
    finally:
        conn.close()


def _fetch_my_voices_page(user_id, page=1, per_page=9):
    """Lấy danh sách giọng có phân trang."""
    if page < 1:
        page = 1

    conn = get_db_connection()
    if not conn:
        return None

    cursor = conn.cursor()
    cursor.execute(
        "SELECT COUNT(*) AS total FROM custom_voices WHERE user_id = %s",
        (user_id,),
    )
    total_row = cursor.fetchone()
    total = total_row['total'] if total_row else 0
    total_pages = max(1, (total + per_page - 1) // per_page) if total > 0 else 1

    if page > total_pages and total > 0:
        page = total_pages

    offset = (page - 1) * per_page
    cursor.execute("""
        SELECT * FROM custom_voices
        WHERE user_id = %s
        ORDER BY created_at DESC
        LIMIT %s OFFSET %s
    """, (user_id, per_page, offset))
    voices = cursor.fetchall()
    conn.close()

    return {
        'voices': voices,
        'page': page,
        'per_page': per_page,
        'total': total,
        'total_pages': total_pages,
    }


@app.route('/my-voices')
@login_required
def my_voices():
    """Page: My custom voices"""
    user_id = session.get('user_id')
    per_page = 9
    page = request.args.get('page', 1, type=int)

    try:
        result = _fetch_my_voices_page(user_id, page, per_page)
        if result is None:
            return render_template(
                'my_voices.html',
                voices=[],
                error="Database connection failed",
                page=1,
                per_page=per_page,
                total=0,
                total_pages=1,
            )

        return render_template('my_voices.html', error=None, **result)
    except Exception as e:
        print(f"[ERROR] My voices page error: {e}")
        import traceback
        traceback.print_exc()
        return render_template(
            'my_voices.html',
            voices=[],
            error=str(e),
            page=1,
            per_page=per_page,
            total=0,
            total_pages=1,
        )
@app.route('/add-voice')
@login_required
def add_voice_page():
    """Page: Add custom voice"""
    return render_template('add_voice.html')

@app.route('/api/custom-voice/upload', methods=['POST'])
@login_required
def upload_custom_voice():
    """Upload audio for custom voice"""
    try:
        if not CUSTOM_VOICE_AVAILABLE:
            return jsonify({'success': False, 'error': 'Custom voice feature not available'}), 503
        
        user_id = session.get('user_id')
        
        # Kiểm tra user tồn tại trong DB (tránh lỗi FK khi DB mới tạo lại / session cũ)
        conn_check = get_db_connection()
        if conn_check:
            try:
                with conn_check.cursor() as cur:
                    cur.execute("SELECT id FROM users WHERE id = %s", (user_id,))
                    if not cur.fetchone():
                        return jsonify({
                            'success': False,
                            'error': 'Phiên đăng nhập không còn hợp lệ. Vui lòng đăng xuất và đăng nhập lại.'
                        }), 401
            finally:
                conn_check.close()
        
        # Get file
        if 'audio_file' not in request.files:
            return jsonify({'success': False, 'error': 'No file uploaded'}), 400
        
        audio_file = request.files['audio_file']
        voice_name = request.form.get('voice_name', 'Untitled Voice')
        description = request.form.get('description', '')
        
        # Voice type: 'rvc' (training), 'zero_shot' (clone from audio + transcript), or 'vixtts_clone' (viXTTS clone)
        voice_type = (request.form.get('voice_type') or 'rvc').strip().lower()
        if voice_type not in ('rvc', 'zero_shot', 'vixtts_clone'):
            voice_type = 'rvc'
        ref_transcript = (request.form.get('ref_transcript') or '').strip()
        if voice_type == 'zero_shot' and not ref_transcript:
            return jsonify({'success': False, 'error': 'Zero-shot cần nhập transcript (nội dung nói) của file mẫu'}), 400
        if voice_type == 'vixtts_clone' and (not VIXTTS_EMOTIONAL_AVAILABLE or VIXTTS_INSTANCE is None):
            return jsonify({'success': False, 'error': 'viXTTS model chưa sẵn sàng. Vui lòng thử lại sau vài phút.'}), 503
        
        # V2: Get base voice and adjustments (with defaults) - for RVC mode
        base_voice_id = request.form.get('base_voice_id', 'ly')
        pitch_adjustment = int(request.form.get('pitch_adjustment', 0))
        speed_adjustment = float(request.form.get('speed_adjustment', 1.0))
        energy_adjustment = float(request.form.get('energy_adjustment', 1.0))
        
        # Capitalize first letter (e.g., 'ly' -> 'Ly', 'binh' -> 'Binh')
        if base_voice_id:
            base_voice_id = base_voice_id.capitalize()
        
        # Validate file
        if not audio_file.filename:
            return jsonify({'success': False, 'error': 'No file selected'}), 400
        
        # Save file — use absolute path so the file is always at WEB_DIR/uploads/...
        filename = f"{user_id}_{int(time.time())}_{secure_filename(audio_file.filename)}"
        upload_dir = str(WEB_DIR / "uploads" / "custom_voices" / f"user_{user_id}")
        os.makedirs(upload_dir, exist_ok=True)
        audio_path = os.path.join(upload_dir, filename)
        audio_file.save(audio_path)
        
        # Validate audio
        audio_processor = get_audio_processor()
        if voice_type == 'vixtts_clone':
            # viXTTS Clone only needs 6–120 seconds of reference audio
            import librosa as _lb
            try:
                duration = _lb.get_duration(path=audio_path)
            except Exception:
                duration = 0
            if duration < 6:
                os.remove(audio_path)
                return jsonify({'success': False, 'error': 'Audio quá ngắn cho viXTTS Clone. Cần ít nhất 6 giây.'}), 400
            if duration > 120:
                os.remove(audio_path)
                return jsonify({'success': False, 'error': 'Audio quá dài cho viXTTS Clone. Tối đa 120 giây (2 phút).'}), 400
            is_valid, message = True, 'Audio hợp lệ'
        else:
            is_valid, message, duration = audio_processor.validate_audio(audio_path)
        
        if not is_valid:
            os.remove(audio_path)  # Remove invalid file
            return jsonify({'success': False, 'error': message}), 400
        
        # Check quality
        quality_score, quality_msg = audio_processor.check_audio_quality(audio_path)
        
        # Get file size
        file_size = os.path.getsize(audio_path)
        
        # Create voice record
        conn = get_db_connection()
        if not conn:
            os.remove(audio_path)
            return jsonify({'success': False, 'error': 'Database connection failed'}), 500
        
        cursor = conn.cursor()
        # Zero-shot / vixtts_clone: status='completed' immediately; RVC: status='pending' then training
        initial_status = 'completed' if voice_type in ('zero_shot', 'vixtts_clone') else 'pending'
        try:
            cursor.execute("""
                INSERT INTO custom_voices 
                (user_id, voice_name, description, sample_audio_path, sample_duration, 
                 sample_file_size, quality_score, status, base_voice_id, pitch_adjustment, 
                 speed_adjustment, energy_adjustment, voice_type, ref_transcript)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            """, (user_id, voice_name, description, audio_path, int(duration), file_size, 
                  quality_score, initial_status, base_voice_id, pitch_adjustment, speed_adjustment, 
                  energy_adjustment, voice_type, ref_transcript if voice_type == 'zero_shot' else None))
        except Exception as db_err:
            # Fallback if voice_type/ref_transcript columns don't exist yet
            if 'voice_type' in str(db_err) or 'ref_transcript' in str(db_err):
                if voice_type == 'zero_shot':
                    conn.close()
                    os.remove(audio_path)
                    return jsonify({'success': False, 'error': 'Cần chạy migration Zero-shot (file custom_voices_zero_shot.sql) trong database để dùng chế độ Zero-shot.'}), 400
                cursor.execute("""
                    INSERT INTO custom_voices 
                    (user_id, voice_name, description, sample_audio_path, sample_duration, 
                     sample_file_size, quality_score, status, base_voice_id, pitch_adjustment, 
                     speed_adjustment, energy_adjustment)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, 'pending', %s, %s, %s, %s)
                """, (user_id, voice_name, description, audio_path, int(duration), file_size, 
                      quality_score, base_voice_id, pitch_adjustment, speed_adjustment, energy_adjustment))
                initial_status = 'pending'
            else:
                raise
        conn.commit()
        custom_voice_id = cursor.lastrowid
        conn.close()
        
        # Start training only for RVC mode
        if voice_type == 'rvc':
            training_service = get_training_service()
            result = training_service.start_training(custom_voice_id, user_id, audio_path)
            return jsonify({
                'success': True,
                'custom_voice_id': custom_voice_id,
                'voice_type': 'rvc',
                'training_mode': result.get('mode'),
                'message': result.get('message'),
                'quality_score': quality_score,
                'quality_message': quality_msg,
                'duration': duration
            })
        elif voice_type == 'vixtts_clone':
            return jsonify({
                'success': True,
                'custom_voice_id': custom_voice_id,
                'voice_type': 'vixtts_clone',
                'message': 'Giọng viXTTS Clone đã sẵn sàng. Bạn có thể dùng ngay.',
                'quality_score': quality_score,
                'quality_message': quality_msg,
                'duration': duration
            })
        else:
            return jsonify({
                'success': True,
                'custom_voice_id': custom_voice_id,
                'voice_type': 'zero_shot',
                'message': 'Giọng Zero-shot đã sẵn sàng. Bạn có thể dùng ngay.',
                'quality_score': quality_score,
                'quality_message': quality_msg,
                'duration': duration
            })
        
    except Exception as e:
        print(f"[ERROR] Upload custom voice failed: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/custom-voice/<int:voice_id>/progress')
@login_required
def get_training_progress(voice_id):
    """Get training progress (for realtime updates)"""
    user_id = session.get('user_id')
    
    try:
        conn = get_db_connection()
        if not conn:
            return jsonify({'error': 'Database connection failed'}), 500
        
        cursor = conn.cursor()
        cursor.execute("""
            SELECT status, training_progress, error_message, quality_score
            FROM custom_voices 
            WHERE id = %s AND user_id = %s
        """, (voice_id, user_id))
        voice = cursor.fetchone()
        conn.close()
        
        if not voice:
            return jsonify({'error': 'Voice not found'}), 404
        
        return jsonify({
            'status': voice['status'],
            'progress': voice['training_progress'],
            'error': voice['error_message'],
            'quality_score': voice['quality_score']
        })
    except Exception as e:
        print(f"[ERROR] Get training progress failed: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/custom-voice/<int:voice_id>/test', methods=['POST'])
@login_required
def test_custom_voice(voice_id):
    """Test custom voice with sample text"""
    user_id = session.get('user_id')
    
    try:
        data = request.get_json()
        text = data.get('text', 'Xin chào, đây là giọng custom của tôi.')
        
        # Verify ownership
        conn = get_db_connection()
        if not conn:
            return jsonify({'error': 'Database connection failed'}), 500
        
        cursor = conn.cursor()
        try:
            cursor.execute("""
                SELECT model_file_path, status, voice_name, sample_audio_path,
                       base_voice_id, pitch_adjustment, speed_adjustment, energy_adjustment,
                       voice_type, ref_transcript
                FROM custom_voices 
                WHERE id = %s AND user_id = %s
            """, (voice_id, user_id))
        except Exception:
            cursor.execute("""
                SELECT model_file_path, status, voice_name, sample_audio_path,
                       base_voice_id, pitch_adjustment, speed_adjustment, energy_adjustment
                FROM custom_voices 
                WHERE id = %s AND user_id = %s
            """, (voice_id, user_id))
        voice = cursor.fetchone()
        if voice and 'voice_type' not in voice:
            voice['voice_type'] = 'rvc'
            voice['ref_transcript'] = None
        conn.close()
        
        if not voice:
            return jsonify({'error': 'Voice not found'}), 404
        
        if voice['status'] != 'completed':
            return jsonify({'error': 'Voice is not ready yet'}), 400
        
        # Test = Play sample audio OR generate with base/zero-shot
        test_text = (request.json.get('test_text') or request.json.get('text') or '').strip()
        
        # Giới hạn độ dài để không vượt context window (2048 tokens) của TTS
        TEST_TEXT_MAX_CHARS = 300
        if len(test_text) > TEST_TEXT_MAX_CHARS:
            test_text = test_text[:TEST_TEXT_MAX_CHARS]
            print(f"[TEST VOICE] Truncated test text to {TEST_TEXT_MAX_CHARS} chars")
        
        if not test_text:
            # No test text provided, just return sample audio URL
            sample_audio_path = voice.get('sample_audio_path', '')
            if sample_audio_path and os.path.exists(sample_audio_path):
                audio_url = '/' + sample_audio_path.replace('\\', '/')
                return jsonify({
                    'success': True,
                    'message': f'Đây là sample audio gốc của giọng: {voice["voice_name"]}',
                    'audio_url': audio_url,
                    'is_sample': True
                })
            else:
                return jsonify({'error': 'Sample audio not found'}), 404
        
        # Generate test audio: Zero-shot (ref_audio+ref_text) or RVC (base voice + adjustments)
        try:
            tts = get_tts_instance()
            if not tts:
                return jsonify({'error': 'TTS model not loaded'}), 500
            
            voice_type_cv = (voice.get('voice_type') or 'rvc').strip().lower()
            audio_filename = f"{uuid.uuid4()}_test.wav"
            audio_path = os.path.join(AUDIO_OUTPUT_DIR, audio_filename)
            
            if voice_type_cv == 'vixtts_clone':
                # viXTTS Clone: synthesize with user's voice reference directly to file
                pitch_adj, speed_adj = 0, 1.0
                ref_audio_path = resolve_audio_path(voice.get('sample_audio_path'))
                if not ref_audio_path or not os.path.exists(ref_audio_path):
                    return jsonify({'error': f'Không tìm thấy file audio mẫu của giọng viXTTS Clone: {ref_audio_path}'}), 400
                if not VIXTTS_EMOTIONAL_AVAILABLE or VIXTTS_INSTANCE is None or VIXTTS_INSTANCE.model is None:
                    return jsonify({'error': 'viXTTS model chưa sẵn sàng. Vui lòng thử lại sau vài phút.'}), 503
                print(f"[TEST VOICE viXTTS-Clone] ref_audio={ref_audio_path}")
                VIXTTS_INSTANCE.synthesize_with_voice(test_text, ref_audio_path, str(audio_path))
            else:
                try:
                    if voice_type_cv == 'zero_shot':
                        ref_audio_path = voice.get('sample_audio_path')
                        ref_text_zs = (voice.get('ref_transcript') or '').strip()
                        pitch_adj, speed_adj = 0, 1.0  # no adjustment for zero_shot
                        if ref_audio_path and os.path.exists(ref_audio_path) and ref_text_zs:
                            # Giới hạn ref_transcript
                            REF_TEXT_MAX_CHARS = 250
                            if len(ref_text_zs) > REF_TEXT_MAX_CHARS:
                                ref_text_zs = ref_text_zs[:REF_TEXT_MAX_CHARS]
                                print(f"[TEST VOICE Zero-shot] Truncated ref_transcript to {REF_TEXT_MAX_CHARS} chars")
                            # Mã hóa ref rồi cắt ref_codes để không vượt context 2048 (audio mẫu dài = rất nhiều token)
                            REF_CODES_MAX_FRAMES = 40
                            ref_codes_full = tts.encode_reference(ref_audio_path)
                            import numpy as np
                            if hasattr(ref_codes_full, 'cpu'):
                                ref_codes_full = ref_codes_full.cpu().numpy()
                            ref_codes_full = np.asarray(ref_codes_full).flatten()
                            ref_codes_short = ref_codes_full[:REF_CODES_MAX_FRAMES].tolist()
                            print(f"[TEST VOICE Zero-shot] ref_audio={ref_audio_path}, ref_codes frames: {len(ref_codes_full)} -> {len(ref_codes_short)}")
                            audio = tts.infer(text=test_text, ref_codes=ref_codes_short, ref_text=ref_text_zs, max_chars=150)
                        else:
                            return jsonify({'error': 'Zero-shot thiếu ref_audio hoặc ref_transcript'}), 400
                    else:
                        base_voice_id = voice.get('base_voice_id', 'ly')
                        pitch_adj = voice.get('pitch_adjustment', 0)
                        speed_adj = voice.get('speed_adjustment', 1.0)
                        tts_voice_id = base_voice_id
                        if tts_voice_id and str(tts_voice_id).endswith('HM'):
                            tts_voice_id = tts_voice_id[:-2]
                        if tts_voice_id:
                            tts_voice_id = str(tts_voice_id).capitalize()
                        voice_data = tts.get_preset_voice(tts_voice_id) if tts_voice_id else None
                        audio = tts.infer(text=test_text, voice=voice_data, max_chars=256)
                except ValueError as ve:
                    if 'context window' in str(ve) or 'exceed' in str(ve).lower():
                        return jsonify({
                            'error': 'Giới hạn model (2048 token) bị vượt. Với Zero-shot: dùng file mẫu ngắn (vài giây) và transcript ngắn; văn bản test giữ dưới 300 ký tự.'
                        }), 400
                    raise
                
                tts.save(audio, str(audio_path))
            
            # Apply speed adjustment if needed
            if speed_adj != 1.0:
                try:
                    import librosa
                    import soundfile as sf
                    print(f"[TEST VOICE] Applying speed adjustment: {speed_adj}x")
                    
                    # Load audio
                    audio_data, sr = librosa.load(audio_path, sr=None)
                    
                    # Change speed
                    audio_adjusted = librosa.effects.time_stretch(audio_data, rate=speed_adj)
                    
                    # Save adjusted audio
                    sf.write(audio_path, audio_adjusted, sr)
                    print(f"[TEST VOICE] Speed adjustment applied successfully")
                except Exception as e:
                    print(f"[WARNING] Could not apply speed adjustment: {e}")
            
            # Apply pitch adjustment if needed
            if pitch_adj != 0:
                try:
                    rvc_processor = get_rvc_processor()
                    if rvc_processor.is_available():
                        adjusted_path = audio_path.replace('.wav', '_adjusted.wav')
                        success, msg, output_path = rvc_processor.adjust_voice(
                            audio_path, adjusted_path, pitch=pitch_adj
                        )
                        if success and output_path:
                            os.remove(audio_path)
                            audio_path = output_path
                            audio_filename = os.path.basename(audio_path)
                except Exception as e:
                    print(f"[WARNING] Could not apply pitch adjustment: {e}")
            
            audio_url = url_for('get_audio', filename=audio_filename)
            
            return jsonify({
                'success': True,
                'message': f'Test thành công với giọng: {voice["voice_name"]}',
                'audio_url': audio_url,
                'is_sample': False
            })
            
        except Exception as e:
            print(f"[ERROR] Test voice generation failed: {e}")
            import traceback
            traceback.print_exc()
            return jsonify({'error': f'Lỗi tạo audio test: {str(e)}'}), 500
        
    except Exception as e:
        print(f"[ERROR] Test custom voice failed: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/custom-voice/<int:voice_id>/delete', methods=['DELETE'])
@login_required
def delete_custom_voice(voice_id):
    """Delete custom voice"""
    user_id = session.get('user_id')
    
    try:
        conn = get_db_connection()
        if not conn:
            return jsonify({'error': 'Database connection failed'}), 500
        
        cursor = conn.cursor()
        
        # Verify ownership
        cursor.execute("""
            SELECT sample_audio_path, model_file_path, index_file_path
            FROM custom_voices 
            WHERE id = %s AND user_id = %s
        """, (voice_id, user_id))
        voice = cursor.fetchone()
        
        if not voice:
            conn.close()
            return jsonify({'error': 'Voice not found'}), 404
        
        # Delete files
        for file_path in [voice.get('sample_audio_path'), voice.get('model_file_path'), voice.get('index_file_path')]:
            if file_path and os.path.exists(file_path):
                try:
                    os.remove(file_path)
                except Exception as e:
                    print(f"[WARNING] Could not delete file {file_path}: {e}")
        
        # Delete from database
        cursor.execute("DELETE FROM custom_voices WHERE id = %s", (voice_id,))
        conn.commit()
        conn.close()
        
        return jsonify({'success': True, 'message': 'Voice deleted successfully'})
        
    except Exception as e:
        print(f"[ERROR] Delete custom voice failed: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/custom-voices/list')
@login_required
def list_custom_voices():
    """List user's custom voices"""
    user_id = session.get('user_id')
    
    try:
        conn = get_db_connection()
        if not conn:
            return jsonify({'error': 'Database connection failed'}), 500
        
        cursor = conn.cursor()
        try:
            cursor.execute("""
                SELECT id, voice_name, status, quality_score, created_at, usage_count, voice_type
                FROM custom_voices 
                WHERE user_id = %s AND status = 'completed'
                ORDER BY created_at DESC
            """, (user_id,))
        except Exception:
            cursor.execute("""
                SELECT id, voice_name, status, quality_score, created_at, usage_count
                FROM custom_voices 
                WHERE user_id = %s AND status = 'completed'
                ORDER BY created_at DESC
            """, (user_id,))
        voices = cursor.fetchall()
        conn.close()
        
        # Convert to JSON-serializable format
        voices_list = []
        for voice in voices:
            voices_list.append({
                'id': voice['id'],
                'name': voice['voice_name'],
                'status': voice['status'],
                'quality_score': float(voice['quality_score']) if voice['quality_score'] else 0,
                'created_at': voice['created_at'].isoformat() if voice['created_at'] else None,
                'usage_count': voice['usage_count'],
                'voice_type': (voice.get('voice_type') or 'rvc').strip().lower()
            })
        
        return jsonify({'success': True, 'voices': voices_list})
        
    except Exception as e:
        print(f"[ERROR] List custom voices failed: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/worker/status')
@login_required
def worker_status():
    """Get background worker status (admin only)"""
    user_role = session.get('user_role')
    
    if user_role != 'admin':
        return jsonify({'error': 'Unauthorized'}), 403
    
    try:
        if not CUSTOM_VOICE_AVAILABLE:
            return jsonify({'error': 'Custom voice feature not available'}), 503
        
        status = get_worker_status()
        return jsonify(status)
        
    except Exception as e:
        print(f"[ERROR] Get worker status failed: {e}")
        return jsonify({'error': str(e)}), 500

# ==================== END CUSTOM VOICE ROUTES ====================

# ==================== PAYMENT STATUS & NOTIFICATION ROUTES ====================

@app.route('/api/payment/status/<int:payment_id>')
def check_payment_status(payment_id):
    """
    Kiểm tra trạng thái thanh toán - có auto-verify qua SePay API nếu còn pending.
    Frontend polls endpoint này mỗi 5 giây.
    """
    if not is_logged_in():
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': 'Database connection failed'}), 500
    
    try:
        with conn.cursor() as cursor:
            cursor.execute("""
                SELECT p.*, sp.package_name, sp.characters_limit, sp.price_vnd, sp.duration_days,
                       u.username, u.full_name
                FROM payments p
                JOIN subscription_packages sp ON p.package_id = sp.id
                JOIN users u ON p.user_id = u.id
                WHERE p.id = %s AND p.user_id = %s
            """, (payment_id, session['user_id']))
            
            payment = cursor.fetchone()
            if not payment:
                return jsonify({'success': False, 'message': 'Không tìm thấy thanh toán'}), 404
            
            # Nếu PENDING → thử auto-verify qua SePay API ngay
            if payment['payment_status'] == 'pending':
                verification = verify_sepay_transaction(
                    payment['transaction_id'],
                    payment['amount_vnd']
                )
                
                if verification.get('verified'):
                    # Tìm thấy giao dịch trên SePay → auto-approve
                    cursor.execute("""
                        UPDATE payments
                        SET payment_status = 'completed',
                            bank_transaction_id = %s,
                            description = 'Auto-verified via SePay API polling',
                            completed_at = NOW()
                        WHERE id = %s AND payment_status = 'pending'
                    """, (payment['transaction_id'], payment_id))
                    
                    conn.commit()
                    
                    # Cộng ký tự cho user
                    update_user_subscription(
                        session['user_id'],
                        payment['characters_limit'],
                        payment['duration_days'],
                        payment.get('package_id'),
                    )
                    
                    print(f"[POLL-AUTO] Payment {payment_id} auto-approved for user {session.get('username')}")
                    
                    try:
                        _notify_payment_success_from_row(payment)
                    except Exception as ne:
                        print(f"[WARN] payment notify: {ne}")
                    
                    # Re-fetch updated status
                    payment['payment_status'] = 'completed'
            
            # Lấy thông tin ký tự nếu đã completed
            characters_info = None
            if payment['payment_status'] == 'completed':
                cursor.execute("""
                    SELECT characters_limit, characters_used,
                           (characters_limit - COALESCE(characters_used, 0)) AS characters_remaining,
                           end_date
                    FROM user_subscriptions
                    WHERE user_id = %s AND is_active = 1
                    ORDER BY created_at DESC LIMIT 1
                """, (session['user_id'],))
                sub_info = cursor.fetchone()
                
                if sub_info:
                    characters_info = {
                        'characters_remaining': int(sub_info['characters_remaining'] or 0),
                        'subscription_expires_at': sub_info['end_date'].isoformat() if sub_info['end_date'] else None
                    }
            
            return jsonify({
                'success': True,
                'payment': {
                    'id': payment['id'],
                    'transaction_id': payment['transaction_id'],
                    'status': payment['payment_status'],
                    'amount': payment['amount_vnd'],
                    'created_at': payment['created_at'].isoformat(),
                    'package_info': {
                        'name': payment['package_name'],
                        'characters': payment['characters_limit'],
                        'price': payment['price_vnd'],
                        'duration': payment['duration_days']
                    }
                },
                'user_characters': characters_info
            })
            
    except Exception as e:
        print(f"[ERROR] Check payment status error: {e}")
        return jsonify({'success': False, 'message': f'Lỗi: {str(e)}'}), 500
    finally:
        conn.close()


@app.route('/api/user/characters')
def get_user_characters():
    """Lấy thông tin ký tự còn lại của user"""
    if not is_logged_in():
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': 'Database connection failed'}), 500
    
    try:
        with conn.cursor() as cursor:
            cursor.execute("""
                SELECT characters_limit, characters_used,
                       (characters_limit - COALESCE(characters_used, 0)) AS characters_remaining,
                       end_date, created_at, updated_at
                FROM user_subscriptions
                WHERE user_id = %s AND is_active = 1
                ORDER BY created_at DESC LIMIT 1
            """, (session['user_id'],))
            
            subscription = cursor.fetchone()
            
            if subscription:
                return jsonify({
                    'success': True,
                    'characters_remaining': int(subscription['characters_remaining'] or 0),
                    'subscription_expires_at': subscription['end_date'].isoformat() if subscription['end_date'] else None,
                    'last_updated': subscription['updated_at'].isoformat() if subscription['updated_at'] else None
                })
            else:
                return jsonify({
                    'success': True,
                    'characters_remaining': 0,
                    'subscription_expires_at': None,
                    'last_updated': None
                })
                
    except Exception as e:
        print(f"[ERROR] Get user characters error: {e}")
        return jsonify({'success': False, 'message': f'Lỗi: {str(e)}'}), 500
    finally:
        conn.close()


@app.route('/api/payment/verify/<transaction_id>')
def manual_verify_payment(transaction_id):
    """Manual verification endpoint for payment"""
    if not is_logged_in():
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': 'Database connection failed'}), 500
    
    try:
        with conn.cursor() as cursor:
            # Kiểm tra payment thuộc về user hiện tại
            cursor.execute("""
                SELECT id, payment_status, user_id, package_id, amount_vnd, transaction_id, payment_method
                FROM payments
                WHERE transaction_id = %s AND user_id = %s
            """, (transaction_id, session['user_id']))
            
            payment = cursor.fetchone()
            if not payment:
                return jsonify({'success': False, 'message': 'Không tìm thấy thanh toán'}), 404
            
            if payment['payment_status'] == 'completed':
                return jsonify({
                    'success': True,
                    'already_verified': True,
                    'message': 'Thanh toán đã được xác nhận thành công'
                })
            
            verify_result = verify_sepay_transaction(transaction_id, payment['amount_vnd'])
            
            if verify_result['verified']:
                cursor.execute("""
                    SELECT p.*, sp.characters_limit, sp.duration_days, sp.package_name
                    FROM payments p
                    LEFT JOIN subscription_packages sp ON p.package_id = sp.id
                    WHERE p.id = %s
                """, (payment['id'],))
                payment_full = cursor.fetchone()

                cursor.execute("""
                    UPDATE payments
                    SET payment_status = 'completed',
                        description = 'Manual verify SePay',
                        completed_at = NOW()
                    WHERE id = %s AND payment_status != 'completed'
                """, (payment['id'],))
                conn.commit()

                sub_ok = update_user_subscription(
                    payment['user_id'],
                    payment_full['characters_limit'],
                    payment_full['duration_days'],
                    payment_full.get('package_id'),
                )
                
                if sub_ok:
                    try:
                        _notify_payment_success_from_row(payment_full)
                    except Exception as ne:
                        print(f"[WARN] payment notify: {ne}")
                    return jsonify({
                        'success': True,
                        'verified': True,
                        'message': '🎉 Thanh toán đã được xác nhận! Bạn đã nhận thêm ký tự vào tài khoản.',
                    })
                return jsonify({
                    'success': False,
                    'message': 'Thanh toán đã ghi nhận nhưng cập nhật gói thất bại. Liên hệ hỗ trợ.',
                })
            else:
                return jsonify({
                    'success': True,
                    'verified': False,
                    'message': 'Thanh toán chưa được xác nhận. Vui lòng kiểm tra lại hoặc liên hệ hỗ trợ.'
                })
                
    except Exception as e:
        print(f"[ERROR] Manual verify payment error: {e}")
        return jsonify({'success': False, 'message': f'Lỗi: {str(e)}'}), 500
    finally:
        conn.close()

# ==================== END PAYMENT STATUS & NOTIFICATION ROUTES ====================

def run_db_migrations():
    """Tự động thêm các cột mới vào DB nếu chưa tồn tại (idempotent)."""
    migrations = [
        # (table, column, definition)
        ('conversions', 'display_name', 'VARCHAR(200) NULL DEFAULT NULL AFTER voice_name'),
        ('conversions', 'is_public',    'TINYINT(1) NOT NULL DEFAULT 0'),
        ('conversions', 'share_token',  'VARCHAR(64) NULL DEFAULT NULL'),
        ('users', 'status', 'VARCHAR(20) NOT NULL DEFAULT \'active\''),
        ('users', 'delete_requested', 'TINYINT(1) NOT NULL DEFAULT 0'),
        ('users', 'delete_requested_at', 'DATETIME NULL DEFAULT NULL'),
        ('users', 'delete_reason', 'TEXT NULL'),
        ('users', 'delete_status', 'VARCHAR(20) NOT NULL DEFAULT \'none\''),
        ('users', 'deleted_at', 'DATETIME NULL DEFAULT NULL'),
        ('users', 'admin_delete_note', 'TEXT NULL'),
        ('users', 'deletion_effective_at', 'DATETIME NULL DEFAULT NULL'),
        ('users', 'restore_requested', 'TINYINT(1) NOT NULL DEFAULT 0'),
        ('users', 'restore_requested_at', 'DATETIME NULL DEFAULT NULL'),
    ]
    conn = get_db_connection()
    if not conn:
        print("[MIGRATION] ⚠️  Không kết nối được DB — bỏ qua migration")
        return
    try:
        with conn.cursor() as cursor:
            db_name = DB_CONFIG['database']
            for table, column, definition in migrations:
                cursor.execute("""
                    SELECT COUNT(*) as cnt FROM INFORMATION_SCHEMA.COLUMNS
                    WHERE TABLE_SCHEMA = %s AND TABLE_NAME = %s AND COLUMN_NAME = %s
                """, (db_name, table, column))
                exists = cursor.fetchone()['cnt'] > 0
                if not exists:
                    cursor.execute(f"ALTER TABLE `{table}` ADD COLUMN `{column}` {definition}")
                    conn.commit()
                    print(f"[MIGRATION] ✅ Đã thêm cột '{column}' vào bảng '{table}'")
                else:
                    print(f"[MIGRATION] ✓  Cột '{column}' trong '{table}' đã tồn tại")

            # Thêm index cho share_token nếu chưa có
            cursor.execute("""
                SELECT COUNT(*) as cnt FROM INFORMATION_SCHEMA.STATISTICS
                WHERE TABLE_SCHEMA = %s AND TABLE_NAME = 'conversions' AND INDEX_NAME = 'idx_share_token'
            """, (db_name,))
            if cursor.fetchone()['cnt'] == 0:
                cursor.execute("ALTER TABLE conversions ADD INDEX idx_share_token (share_token)")
                conn.commit()
                print("[MIGRATION] ✅ Đã thêm index 'idx_share_token'")

            # Backfill grace period cho tài khoản đã duyệt xóa trước khi có deletion_effective_at
            cursor.execute("""
                UPDATE users SET
                    deletion_effective_at = DATE_ADD(deleted_at, INTERVAL %s DAY),
                    status = 'deactivated'
                WHERE delete_status = 'approved'
                  AND deleted_at IS NOT NULL
                  AND deletion_effective_at IS NULL
                  AND status = 'deleted'
                  AND DATE_ADD(deleted_at, INTERVAL %s DAY) > NOW()
            """, (ACCOUNT_DELETION_GRACE_DAYS, ACCOUNT_DELETION_GRACE_DAYS))
            conn.commit()

            _finalize_expired_account_deletions(cursor)
            conn.commit()
            print("[MIGRATION] ✓  Đã kiểm tra grace period xóa tài khoản")
    except Exception as e:
        print(f"[MIGRATION] ❌ Lỗi migration: {e}")
    finally:
        conn.close()


def warmup_legal_en_cache():
    """Build EN legal caches at startup so first page load is instant."""
    for key in LEGAL_CONTENT_PAGE_KEYS:
        try:
            get_legal_for_display(key, 'en')
        except Exception as e:
            print(f'[i18n] Legal EN cache warmup failed for {key}: {e}')


def warmup_support_en_cache():
    """Build EN support cache at startup."""
    try:
        get_support_for_display('en')
    except Exception as e:
        print(f'[i18n] Support EN cache warmup failed: {e}')


if __name__ == '__main__':
    print("=" * 60)
    print("[TTS] TTS Web Application dang khoi dong...")
    print("[TTS] URL: http://localhost:5000")
    print("[TTS] Port: 5000")
    print("=" * 60)

    # Chạy DB migration tự động
    print("[MIGRATION] Kiểm tra và cập nhật cấu trúc database...")
    run_db_migrations()
    warmup_legal_en_cache()
    warmup_support_en_cache()
    
    # Start background worker for custom voice training
    try:
        if CUSTOM_VOICE_AVAILABLE:
            from background_worker import start_worker
            success = start_worker()
            if success:
                print("[WORKER] ✅ Background training worker started")
            else:
                print("[WORKER] ⚠️ Worker already running")
    except Exception as e:
        print(f"[WORKER] ❌ Failed to start worker: {e}")
    
    # Pre-load viXTTS Emotional TTS để user không phải chờ lần đầu
    if VIXTTS_EMOTIONAL_AVAILABLE:
        try:
            print("\n" + "=" * 60)
            print("[viXTTS] 🚀 ĐANG LOAD EMOTIONAL TTS MODEL...")
            print("[viXTTS] Vui lòng đợi, server sẽ sẵn sàng sau 30-45 giây")
            print("[viXTTS] (Lần đầu tiên sẽ download model ~2GB)")
            print("=" * 60 + "\n")
            
            VIXTTS_INSTANCE = get_vixtts_emotional_instance()
            VIXTTS_INSTANCE.load_model()  # Load model ngay
            
            print("\n" + "=" * 60)
            print("[viXTTS] ✅ MODEL ĐÃ SẴN SÀNG!")
            print("[viXTTS] User có thể sử dụng ngay không cần chờ")
            print("[viXTTS] Model sẽ được giữ trong RAM cho đến khi restart")
            print("=" * 60 + "\n")
        except Exception as e:
            import traceback
            VIXTTS_INSTANCE = None
            print(f"\n[viXTTS] ❌ Pre-load FAILED: {e}")
            print(f"[viXTTS] Traceback:\n{traceback.format_exc()}")
            print("[viXTTS] Model sẽ được load khi có request đầu tiên\n")
    
    # Check readiness status
    is_ready = VIXTTS_INSTANCE is not None and VIXTTS_INSTANCE.model is not None if VIXTTS_EMOTIONAL_AVAILABLE else False
    
    print("=" * 60)
    print("[TTS] 🎉 SERVER READY - SẴN SÀNG PHỤC VỤ!")
    print("[TTS] URL: http://127.0.0.1:5000")
    print("[TTS] Emotional TTS: " + ("✅ Sẵn sàng" if is_ready else "❌ Không khả dụng"))
    if SMTP_HOST and SMTP_USER:
        print(f"[EMAIL] SMTP: ✅ {SMTP_USER} @ {SMTP_HOST}:{SMTP_PORT}")
    else:
        print("[EMAIL] SMTP: ❌ Chưa cấu hình (thêm SMTP_* vào .env.local để gửi email thật)")
    print("[TTS] Nhấn Ctrl+C để dừng server")
    print("=" * 60)
    print()
    
    # Chạy Flask với cấu hình tối ưu
    # use_reloader=False để tránh reload loop với PyTorch
    port = int(os.environ.get('PORT', 5000))
    is_production = os.environ.get('FLASK_ENV') == 'production'
    app.run(debug=not is_production, host='0.0.0.0', port=port, use_reloader=False, threaded=True)
